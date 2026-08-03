from __future__ import annotations

import os
import shutil
import sys
import tempfile
import unittest
import uuid
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import AsyncMock, MagicMock, patch

from docx import Document
from openpyxl import Workbook
from PyPDF2 import PdfReader, PdfWriter

try:
    import fitz

    _HAS_FITZ = True
except Exception:  # pragma: no cover - PyMuPDF is a declared dependency
    _HAS_FITZ = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas as _reportlab_canvas

    _HAS_REPORTLAB = True
except Exception:  # pragma: no cover - reportlab is a declared dependency
    _HAS_REPORTLAB = False

os.environ.setdefault("DATABASE_URL", "sqlite:///:memory:")
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from models.db_models import FormFillRun
from services.form_fill_service import FormFillService, _normalize_repeat_mode, _parse_date


class RetryableGeminiInvalidArgument(Exception):
    status_code = 400

    def __str__(self) -> str:
        return "400 INVALID_ARGUMENT. {'error': {'code': 400, 'message': 'Request contains an invalid argument.', 'status': 'INVALID_ARGUMENT'}}"


class RetryableGeminiCancelled(Exception):
    status_code = 499

    def __str__(self) -> str:
        return "499 CANCELLED. {'error': {'code': 499, 'message': 'The operation was cancelled.', 'status': 'CANCELLED'}}"


class FormFillOutputEnqueueTests(unittest.IsolatedAsyncioTestCase):
    async def test_enqueue_output_units_uses_staggered_delays(self) -> None:
        service = FormFillService()
        outputs = [SimpleNamespace(id=f"output-{index}") for index in range(5)]
        run = SimpleNamespace(
            id="run-id",
            user_id="user-id",
            outputs=outputs,
            total_outputs=0,
            completed_outputs=0,
            failed_outputs=0,
            fill_plan=None,
            usage_basis=None,
        )
        query = MagicMock()
        query.filter.return_value.order_by.return_value.all.return_value = outputs
        db = MagicMock()
        db.query.return_value = query

        with patch.object(service, "_check_usage_limit_or_raise"), patch.object(
            service,
            "_sync_run_output_counts",
        ), patch.object(
            service,
            "_finalize_run_if_ready",
            new=AsyncMock(),
        ), patch.dict(
            os.environ,
            {
                "FORM_FILL_OUTPUT_ENQUEUE_BATCH_SIZE": "2",
                "FORM_FILL_OUTPUT_ENQUEUE_BATCH_DELAY_SECONDS": "15",
                "FORM_FILL_OUTPUT_ENQUEUE_MAX_DELAY_SECONDS": "900",
                "FORM_FILL_OUTPUT_ENQUEUE_JITTER_SECONDS": "0",
            },
        ), patch("services.form_fill_service.cloud_run_task_service.enqueue_form_fill_output_task", new_callable=AsyncMock) as enqueue:
            enqueue.side_effect = [f"cloud-task-{index}" for index in range(5)]

            result = await service._enqueue_output_units(
                db=db,
                run=run,
                target_page_count=1,
                units=[{"record_index": index, "record_label": str(index), "record_payload": {}} for index in range(5)],
                strategy="test",
            )

        delays = [call.kwargs["delay_seconds"] for call in enqueue.await_args_list]
        self.assertEqual(delays, [0, 0, 15, 15, 30])
        self.assertEqual(result["enqueued_outputs"], 5)


class FormFillCreateRunTests(unittest.IsolatedAsyncioTestCase):
    async def test_create_run_defers_docx_target_page_count(self) -> None:
        service = FormFillService()
        added: list[object] = []
        run_id = uuid.UUID("11111111-1111-1111-1111-111111111111")

        db = MagicMock()
        db.add.side_effect = added.append

        def flush() -> None:
            for item in added:
                if item.__class__.__name__ == "FormFillRun" and getattr(item, "id", None) is None:
                    item.id = run_id

        db.flush.side_effect = flush

        class Upload:
            filename = "target.docx"
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

            async def read(self) -> bytes:
                return b"docx-bytes"

        with patch.object(service, "_get_session", return_value=db), patch.object(
            service,
            "_load_extraction_source_payload",
            return_value={"kind": "extraction_result", "columns": ["Name"], "rows": [["Acme"]], "source_files": []},
        ), patch.object(
            service,
            "_count_target_pages_from_bytes",
            new=AsyncMock(side_effect=AssertionError("DOCX page counting should be deferred")),
        ) as count_pages, patch.object(
            service,
            "_upload_bytes",
            new=AsyncMock(),
        ), patch.object(
            service,
            "_serialize_run",
            side_effect=lambda run: {"id": str(run.id), "target_page_count": run.target_page_count},
        ), patch("services.form_fill_service.cloud_run_task_service.enqueue_form_fill_task", new_callable=AsyncMock):
            result = await service.create_run(
                user_id="user-id",
                target_file=Upload(),
                source_job_id="22222222-2222-2222-2222-222222222222",
                source_run_id="33333333-3333-3333-3333-333333333333",
                source_task_id="44444444-4444-4444-4444-444444444444",
            )

        count_pages.assert_not_awaited()
        run = next(item for item in added if item.__class__.__name__ == "FormFillRun")
        self.assertIsNone(run.target_page_count)
        self.assertEqual(result["target_page_count"], None)


class FormFillServiceDocxEditPlanTests(unittest.TestCase):
    def setUp(self) -> None:
        self.service = FormFillService()

    def _create_table_docx(self) -> str:
        document = Document()
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header A"
        table.cell(0, 1).text = "Header B"
        table.cell(1, 0).text = "Value A1"
        table.cell(1, 1).text = "Value B1"

        handle = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        handle.close()
        document.save(handle.name)
        self.addCleanup(lambda: os.path.exists(handle.name) and os.unlink(handle.name))
        return handle.name

    def _create_merged_table_docx(self) -> str:
        document = Document()
        table = document.add_table(rows=5, cols=3)

        title_cell = table.cell(0, 0).merge(table.cell(0, 2))
        title_cell.text = "Inventory"

        header_label = table.cell(1, 0).merge(table.cell(1, 1))
        header_label.text = "DESCRIPTION"
        table.cell(1, 2).text = "AMOUNT/VALUE"

        item_label = table.cell(2, 0).merge(table.cell(2, 1))
        item_label.text = "Oak dining table"
        table.cell(2, 2).text = "850.00"

        item_label = table.cell(3, 0).merge(table.cell(3, 1))
        item_label.text = "2018 Toyota Camry"
        table.cell(3, 2).text = "14250.00"

        total_cell = table.cell(4, 0).merge(table.cell(4, 2))
        total_cell.text = "TOTAL INVENTORY: $ 15100.00"

        handle = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        handle.close()
        document.save(handle.name)
        self.addCleanup(lambda: os.path.exists(handle.name) and os.unlink(handle.name))
        return handle.name

    def _temp_output_path(self) -> str:
        handle = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        handle.close()
        self.addCleanup(lambda: os.path.exists(handle.name) and os.unlink(handle.name))
        return handle.name

    def test_apply_docx_edit_plan_inserts_row_after_selected_row(self) -> None:
        source_path = self._create_table_docx()
        output_path = self._temp_output_path()

        warnings = self.service._apply_docx_edit_plan(
            source_path,
            [
                {
                    "action": "insert_table_row_after",
                    "table_id": "table.0",
                    "row_index": 0,
                    "cells": ["Header C", "Header D"],
                }
            ],
            output_path,
            allow_table_expansion=True,
        )

        self.assertEqual(warnings, [])
        document = Document(output_path)
        table = document.tables[0]
        self.assertEqual(len(table.rows), 3)
        self.assertEqual([cell.text for cell in table.rows[1].cells], ["Header C", "Header D"])
        self.assertEqual([cell.text for cell in table.rows[2].cells], ["Value A1", "Value B1"])

    def test_apply_docx_edit_plan_accepts_generated_operation_alias_and_cell_text_objects(self) -> None:
        source_path = self._create_table_docx()
        output_path = self._temp_output_path()

        warnings = self.service._apply_docx_edit_plan(
            source_path,
            [
                {
                    "operation": "insert_table_row_after",
                    "table_id": "table.0",
                    "row_index": 0,
                    "cells": [[{"text": "Generated A"}], [{"text": "Generated B"}]],
                }
            ],
            output_path,
            allow_table_expansion=True,
        )

        self.assertEqual(warnings, [])
        document = Document(output_path)
        table = document.tables[0]
        self.assertEqual([cell.text for cell in table.rows[1].cells], ["Generated A", "Generated B"])

    def test_validate_generated_operations_normalizes_action_alias(self) -> None:
        result = self.service._validate_generated_transform_result(
            {
                "operations": [
                    {
                        "operation": "insert_table_row_after",
                        "table_id": "table.0",
                        "row_index": 0,
                        "cells": [[{"text": "Date"}], [{"text": "Amount"}]],
                    }
                ],
                "warnings": [],
            },
            expected_key="operations",
        )

        self.assertEqual(
            result["operations"],
            [
                {
                    "operation": "insert_table_row_after",
                    "table_id": "table.0",
                    "row_index": 0,
                    "cells": ["Date", "Amount"],
                    "action": "insert_table_row_after",
                }
            ],
        )

    def test_apply_docx_edit_plan_inserts_column_after_selected_column(self) -> None:
        source_path = self._create_table_docx()
        output_path = self._temp_output_path()

        warnings = self.service._apply_docx_edit_plan(
            source_path,
            [
                {
                    "action": "insert_table_column_after",
                    "table_id": "table.0",
                    "column_index": 0,
                    "cells": ["Header Inserted", "Value Inserted"],
                }
            ],
            output_path,
            allow_table_expansion=True,
        )

        self.assertEqual(warnings, [])
        document = Document(output_path)
        table = document.tables[0]
        self.assertEqual(len(table.rows[0].cells), 3)
        self.assertEqual([cell.text for cell in table.rows[0].cells], ["Header A", "Header Inserted", "Header B"])
        self.assertEqual([cell.text for cell in table.rows[1].cells], ["Value A1", "Value Inserted", "Value B1"])

    def test_apply_docx_edit_plan_blocks_table_growth_when_not_allowed(self) -> None:
        source_path = self._create_table_docx()
        output_path = self._temp_output_path()

        warnings = self.service._apply_docx_edit_plan(
            source_path,
            [
                {
                    "action": "insert_table_row_after",
                    "table_id": "table.0",
                    "row_index": 0,
                    "cells": ["Blocked A", "Blocked B"],
                }
            ],
            output_path,
            allow_table_expansion=False,
        )

        self.assertEqual(len(warnings), 1)
        self.assertIn("not permitted", warnings[0])
        document = Document(output_path)
        table = document.tables[0]
        self.assertEqual(len(table.rows), 2)
        self.assertEqual([cell.text for cell in table.rows[1].cells], ["Value A1", "Value B1"])

    def test_collect_docx_tables_uses_logical_cells_for_merged_rows(self) -> None:
        source_path = self._create_merged_table_docx()

        document = Document(source_path)
        tables, _table_map = self.service._collect_docx_tables(document)

        self.assertEqual(len(tables), 1)
        self.assertEqual(tables[0]["columns"], 2)
        self.assertEqual(
            tables[0]["preview_rows"],
            [
                "row 0: Inventory",
                "row 1: DESCRIPTION | AMOUNT/VALUE",
                "row 2: Oak dining table | 850.00",
                "row 3: 2018 Toyota Camry | 14250.00",
                "row 4: TOTAL INVENTORY: $ 15100.00",
            ],
        )

    def test_extract_docx_text_does_not_repeat_merged_cell_text(self) -> None:
        source_path = self._create_merged_table_docx()

        text = self.service._extract_docx_text(source_path)

        self.assertEqual(text.count("TOTAL INVENTORY: $ 15100.00"), 1)
        self.assertIn("DESCRIPTION | AMOUNT/VALUE", text)

    def test_apply_docx_edit_plan_inserts_row_into_merged_table_using_logical_cells(self) -> None:
        source_path = self._create_merged_table_docx()
        output_path = self._temp_output_path()

        warnings = self.service._apply_docx_edit_plan(
            source_path,
            [
                {
                    "action": "insert_table_row_after",
                    "table_id": "table.0",
                    "row_index": 2,
                    "cells": ["Savings account ending 4421", "6325.47"],
                }
            ],
            output_path,
            allow_table_expansion=True,
        )

        self.assertEqual(warnings, [])
        document = Document(output_path)
        table = document.tables[0]
        inserted_cells = self.service._docx_row_cells(table.rows[3])
        self.assertEqual([cell.text for cell in inserted_cells], ["Savings account ending 4421", "6325.47"])

    def test_apply_docx_edit_plan_blocks_column_insertion_for_merged_table(self) -> None:
        source_path = self._create_merged_table_docx()
        output_path = self._temp_output_path()

        warnings = self.service._apply_docx_edit_plan(
            source_path,
            [
                {
                    "action": "insert_table_column_after",
                    "table_id": "table.0",
                    "column_index": 0,
                    "cells": ["Inserted"] * 5,
                }
            ],
            output_path,
            allow_table_expansion=True,
        )

        self.assertEqual(warnings, ["Could not insert a table column because the table contains merged cells."])
        document = Document(output_path)
        table = document.tables[0]
        self.assertEqual(len(self.service._docx_row_cells(table.rows[1])), 2)


class FormFillServiceSourceContextTests(unittest.IsolatedAsyncioTestCase):
    def setUp(self) -> None:
        self.service = FormFillService()

    def test_combine_extraction_payloads_aligns_all_rows_to_unified_columns(self) -> None:
        payload = self.service._combine_extraction_payloads(
            [
                {
                    "columns": ["Name", "Amount"],
                    "rows": [["Acme", 100], ["Beta", 250]],
                    "source_files": ["first.pdf"],
                },
                {
                    "columns": ["Tax ID", "Name"],
                    "rows": [["12-3456789", "Gamma"]],
                    "source_files": ["second.pdf", "first.pdf"],
                },
            ],
            job_id="job-id",
            run_id="run-id",
        )

        self.assertEqual(payload["scope"], "all")
        self.assertIsNone(payload["task_id"])
        self.assertEqual(payload["columns"], ["Name", "Amount", "Tax ID"])
        self.assertEqual(
            payload["rows"],
            [
                ["Acme", 100, None],
                ["Beta", 250, None],
                ["Gamma", None, "12-3456789"],
            ],
        )
        self.assertEqual(payload["source_files"], ["first.pdf", "second.pdf"])
        self.assertEqual(len(payload["task_groups"]), 2)
        self.assertEqual(payload["task_groups"][0]["source_files"], ["first.pdf"])
        self.assertEqual(payload["task_groups"][0]["rows"], [["Acme", 100], ["Beta", 250]])
        self.assertEqual(payload["task_groups"][1]["source_files"], ["second.pdf", "first.pdf"])
        self.assertEqual(payload["task_groups"][1]["rows"], [["12-3456789", "Gamma"]])

    def test_source_file_units_create_one_unit_per_uploaded_source_file(self) -> None:
        run = SimpleNamespace(
            source_files=[
                SimpleNamespace(
                    id="source-1",
                    original_filename="first client.pdf",
                    file_type="application/pdf",
                    gcs_object_name="first-object",
                    display_order=0,
                ),
                SimpleNamespace(
                    id="source-2",
                    original_filename="second client.pdf",
                    file_type="application/pdf",
                    gcs_object_name="second-object",
                    display_order=1,
                ),
            ]
        )

        units = self.service._source_file_units(run)

        self.assertEqual([unit["record_label"] for unit in units], ["first_client", "second_client"])
        self.assertEqual([unit["record_payload"]["kind"] for unit in units], ["source_file", "source_file"])
        self.assertEqual([unit["record_payload"]["gcs_object_name"] for unit in units], ["first-object", "second-object"])

    def test_extraction_task_units_use_task_groups_not_flat_rows(self) -> None:
        payload = {
            "kind": "extraction_result",
            "scope": "all",
            "columns": ["Name"],
            "rows": [["first"], ["second"]],
            "task_groups": [
                {
                    "task_id": "task-1",
                    "source_files": ["source-one.pdf"],
                    "columns": ["Name"],
                    "rows": [["first"]],
                },
                {
                    "task_id": "task-2",
                    "source_files": ["source-two.pdf"],
                    "columns": ["Name"],
                    "rows": [["second"]],
                },
            ],
        }

        units = self.service._extraction_task_units_from_payload(payload)

        self.assertEqual([unit["record_label"] for unit in units], ["source-one", "source-two"])
        self.assertEqual(units[0]["record_payload"]["rows"], [["first"]])
        self.assertEqual(units[1]["record_payload"]["rows"], [["second"]])

    async def test_build_output_source_context_uses_one_extraction_task_group(self) -> None:
        unit = self.service._extraction_task_units_from_payload(
            {
                "kind": "extraction_result",
                "scope": "all",
                "task_groups": [
                    {
                        "task_id": "task-1",
                        "source_files": ["source-one.pdf"],
                        "columns": ["Name"],
                        "rows": [["first"]],
                    }
                ],
            }
        )[0]

        source_parts, source_text = await self.service._build_output_source_context(
            run=SimpleNamespace(),
            record_payload=unit["record_payload"],
            record_index=0,
        )

        self.assertEqual(source_parts, [])
        self.assertIn("source-one.pdf", source_text)
        self.assertIn("| first |", source_text)

    def test_sync_run_output_counts_counts_completed_and_failed_outputs(self) -> None:
        completed_query = MagicMock()
        completed_query.filter.return_value.count.return_value = 2
        failed_query = MagicMock()
        failed_query.filter.return_value.count.return_value = 1
        db = MagicMock()
        db.query.side_effect = [completed_query, failed_query]
        run = SimpleNamespace(id="run-id", completed_outputs=0, failed_outputs=0)

        self.service._sync_run_output_counts(db, run)

        self.assertEqual(run.completed_outputs, 2)
        self.assertEqual(run.failed_outputs, 1)

    async def test_enqueue_output_units_requeues_pending_existing_outputs(self) -> None:
        run_id = "11111111-1111-1111-1111-111111111111"
        output_id = "22222222-2222-2222-2222-222222222222"
        output = SimpleNamespace(id=output_id, status="pending")
        run = SimpleNamespace(
            id=run_id,
            user_id="user-id",
            outputs=[output],
            total_outputs=1,
            fill_plan=None,
        )
        pending_query = MagicMock()
        pending_query.filter.return_value.order_by.return_value.all.return_value = [output]
        db = MagicMock()
        db.query.return_value = pending_query

        with patch.object(self.service, "_check_usage_limit_or_raise") as check_limit:
            with patch.object(self.service, "_sync_run_output_counts") as sync_counts:
                with patch.object(self.service, "_finalize_run_if_ready", new=AsyncMock(return_value={"finalized": False})):
                    with patch("services.form_fill_service.cloud_run_task_service") as task_service:
                        task_service.calculate_stagger_delay.return_value = 0
                        task_service.enqueue_form_fill_output_task = AsyncMock(return_value="task-name")

                        result = await self.service._enqueue_output_units(
                            db=db,
                            run=run,
                            target_page_count=3,
                            units=[{"record_index": 0, "record_label": "first", "record_payload": {}}],
                            strategy="source_files",
                        )

        check_limit.assert_called_once()
        self.assertEqual(sync_counts.call_count, 2)
        sync_counts.assert_called_with(db, run)
        task_service.enqueue_form_fill_output_task.assert_awaited_once_with(run_id, output_id, delay_seconds=0)
        self.assertEqual(result["enqueued_outputs"], 1)
        self.assertEqual(result["failed_to_enqueue"], 0)

    def _create_csv(self, content: str) -> str:
        handle = tempfile.NamedTemporaryFile(suffix=".csv", delete=False, mode="w", encoding="utf-8")
        try:
            handle.write(content)
        finally:
            handle.close()
        self.addCleanup(lambda: os.path.exists(handle.name) and os.unlink(handle.name))
        return handle.name

    def _create_xlsx(self, rows: list[list[object]]) -> str:
        workbook = Workbook()
        worksheet = workbook.active
        for row in rows:
            worksheet.append(row)

        handle = tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False)
        handle.close()
        workbook.save(handle.name)
        workbook.close()
        self.addCleanup(lambda: os.path.exists(handle.name) and os.unlink(handle.name))
        return handle.name

    async def test_build_source_context_labels_multiple_uploaded_csv_sources(self) -> None:
        first_path = self._create_csv("Name,Amount\nChecking,100\n")
        second_path = self._create_csv("Name,Amount\nSavings,250\n")

        async def fake_download(object_name: str, local_path: str) -> None:
            shutil.copyfile({"first": first_path, "second": second_path}[object_name], local_path)

        self.service._download_to_local = fake_download
        run = SimpleNamespace(
            id="run-id",
            user_id="user-id",
            source_payload=None,
            source_files=[
                SimpleNamespace(
                    id="source-1",
                    original_filename="first.csv",
                    file_type="text/csv",
                    gcs_object_name="first",
                    display_order=0,
                ),
                SimpleNamespace(
                    id="source-2",
                    original_filename="second.csv",
                    file_type="text/csv",
                    gcs_object_name="second",
                    display_order=1,
                ),
            ],
            source_gcs_object_name=None,
            source_file_type=None,
        )

        source_parts, source_text = await self.service._build_source_context(
            run=run,
            source_parts=[],
            source_text_sections=[],
        )

        self.assertEqual(source_parts, [])
        self.assertIn("Source file 1: first.csv", source_text)
        self.assertIn("| Checking | 100 |", source_text)
        self.assertIn("Source file 2: second.csv", source_text)
        self.assertIn("| Savings | 250 |", source_text)

    async def test_build_source_context_supports_legacy_single_source_upload(self) -> None:
        csv_path = self._create_csv("Field,Value\nOwner,Jane\n")

        async def fake_download(object_name: str, local_path: str) -> None:
            self.assertEqual(object_name, "legacy-object")
            shutil.copyfile(csv_path, local_path)

        self.service._download_to_local = fake_download
        run = SimpleNamespace(
            id="run-id",
            user_id="user-id",
            source_payload=None,
            source_files=[],
            source_gcs_object_name="legacy-object",
            source_file_type="text/csv",
            source_filename="legacy.csv",
        )

        _source_parts, source_text = await self.service._build_source_context(
            run=run,
            source_parts=[],
            source_text_sections=[],
        )

        self.assertIn("Source file 1: legacy.csv", source_text)
        self.assertIn("| Owner | Jane |", source_text)

    async def test_build_source_context_does_not_cap_extraction_result_rows(self) -> None:
        self.service.max_spreadsheet_rows = 1
        run = SimpleNamespace(
            id="run-id",
            user_id="user-id",
            source_payload={
                "kind": "extraction_result",
                "columns": ["Name"],
                "rows": [["first"], ["second"], ["third"]],
                "source_files": ["source.pdf"],
            },
            source_files=[],
            source_gcs_object_name=None,
            source_file_type=None,
        )

        _source_parts, source_text = await self.service._build_source_context(
            run=run,
            source_parts=[],
            source_text_sections=[],
        )

        self.assertIn("| first |", source_text)
        self.assertIn("| second |", source_text)
        self.assertIn("| third |", source_text)

    async def test_extract_repeat_records_does_not_cap_extraction_result_rows(self) -> None:
        self.service.max_repeat_records = 1
        run = SimpleNamespace(
            source_payload={
                "kind": "extraction_result",
                "columns": ["Name"],
                "rows": [["first"], ["second"], ["third"]],
            },
            source_files=[],
        )

        records = await self.service._extract_repeat_records(run)

        self.assertEqual([record["record_label"] for record in records], ["first", "second", "third"])

    def test_load_csv_text_does_not_cap_input_rows_or_characters(self) -> None:
        self.service.max_spreadsheet_rows = 1
        self.service.max_sheet_chars = 20
        csv_path = self._create_csv("Name,Amount\nfirst,1\nsecond,2\nthird,3\n")

        source_text = self.service._load_csv_text(csv_path)

        self.assertIn("| first | 1 |", source_text)
        self.assertIn("| second | 2 |", source_text)
        self.assertIn("| third | 3 |", source_text)

    def test_load_xlsx_text_does_not_cap_input_rows_or_characters(self) -> None:
        self.service.max_spreadsheet_rows = 1
        self.service.max_sheet_chars = 20
        xlsx_path = self._create_xlsx([
            ["Name", "Amount"],
            ["first", 1],
            ["second", 2],
            ["third", 3],
        ])

        source_text = self.service._load_xlsx_text(xlsx_path)

        self.assertIn("| first | 1 |", source_text)
        self.assertIn("| second | 2 |", source_text)
        self.assertIn("| third | 3 |", source_text)

    def test_load_csv_records_does_not_cap_output_rows(self) -> None:
        self.service.max_repeat_records = 1
        csv_path = self._create_csv("Name,Amount\nfirst,1\nsecond,2\nthird,3\n")

        records = self.service._load_csv_records(csv_path)

        self.assertEqual([record["record_label"] for record in records], ["first", "second", "third"])

    def test_load_xlsx_records_does_not_cap_output_rows(self) -> None:
        self.service.max_repeat_records = 1
        xlsx_path = self._create_xlsx([
            ["Name", "Amount"],
            ["first", 1],
            ["second", 2],
            ["third", 3],
        ])

        records = self.service._load_xlsx_records(xlsx_path)

        self.assertEqual([record["record_label"] for record in records], ["first", "second", "third"])

    def test_load_csv_records_uses_header_row_and_name_label(self) -> None:
        csv_path = self._create_csv("Participant Name,Email\nJane Doe,jane@example.com\nJohn Smith,john@example.com\n")

        records = self.service._load_csv_records(csv_path)

        self.assertEqual(len(records), 2)
        self.assertEqual(records[0]["record_index"], 0)
        self.assertEqual(records[0]["record_label"], "Jane Doe")
        self.assertEqual(records[0]["record_payload"]["Email"], "jane@example.com")
        self.assertEqual(records[1]["record_label"], "John Smith")

    def test_record_source_text_scopes_prompt_to_one_record(self) -> None:
        source_text = self.service._record_source_text(
            {
                "record_payload": {
                    "Participant Name": "Jane Doe",
                    "Email": "jane@example.com",
                }
            }
        )

        self.assertIn("single source record only", source_text)
        self.assertIn("Participant Name: Jane Doe", source_text)
        self.assertIn("Email: jane@example.com", source_text)

    def test_record_placeholder_replacements_match_normalized_row_keys(self) -> None:
        replacements = self.service._record_placeholder_replacements(
            ["{{client_name}}", "{{Tax Year}}", "[[primary-contact]]", "<<missing>>"],
            {
                "Client Name": "Arbor & Finch LLC",
                "tax_year": 2025,
                "Primary Contact": "Maya Patel",
            },
        )

        self.assertEqual(replacements["{{client_name}}"], "Arbor & Finch LLC")
        self.assertEqual(replacements["{{Tax Year}}"], "2025")
        self.assertEqual(replacements["[[primary-contact]]"], "Maya Patel")
        self.assertNotIn("<<missing>>", replacements)

    async def test_generate_docx_placeholders_never_runs_post_placeholder_table_expansion(self) -> None:
        handle = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        handle.close()
        self.addCleanup(lambda: os.path.exists(handle.name) and os.unlink(handle.name))
        target_path = handle.name
        temp_dir = tempfile.mkdtemp(prefix="form_fill_test_")
        self.addCleanup(lambda: shutil.rmtree(temp_dir, ignore_errors=True))
        run = SimpleNamespace(
            id="run-id",
            user_id="user-id",
            target_file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            target_filename="target.docx",
            target_gcs_object_name="target-object",
            allow_docx_table_expansion=True,
            fill_chronologically=True,
            output_format="docx",
        )
        document = Document()
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Client"
        table.cell(0, 1).text = "{{client_name}}"
        table.cell(1, 0).text = "Contact"
        table.cell(1, 1).text = "{{primary_contact}}"
        document.save(target_path)

        converter = SimpleNamespace(convert_docx_gcs_to_pdf_gcs=AsyncMock())
        self.service.storage_service.construct_gcs_uri_for_object = MagicMock(return_value="gs://bucket/target")

        with patch.object(self.service, "_part_from_storage_object", return_value=object()):
            with patch("services.form_fill_service.get_document_conversion_service", return_value=converter):
                with patch.object(self.service, "_generate_mapping_payload") as generate_mapping:
                    with patch.object(self.service, "_generate_collection_json_response") as generate_collection:
                        with patch.object(self.service, "_generate_and_execute_tabular_transform") as generate_tabular:
                            with patch.object(self.service, "_apply_docx_edit_plan") as apply_edit_plan:
                                generate_collection.side_effect = AssertionError("post-placeholder DOCX table expansion should not run")
                                generate_tabular.side_effect = AssertionError("DOCX placeholders should use direct matching")
                                apply_edit_plan.side_effect = AssertionError("post-placeholder DOCX table expansion should not be applied")
                                result = await self.service._generate_filled_document(
                                    run=run,
                                    temp_dir=temp_dir,
                                    target_local_path=target_path,
                                    source_parts=[],
                                    source_text="Fill the DOCX placeholders from the source data.",
                                    tabular_context={
                                        "columns": ["client_name", "primary_contact"],
                                        "rows": [{"client_name": "Beta LLC", "primary_contact": "Jane Doe"}],
                                        "row_count": 1,
                                    },
                                )

        generate_mapping.assert_not_called()
        generate_collection.assert_not_called()
        generate_tabular.assert_not_called()
        apply_edit_plan.assert_not_called()
        output = Document(result["local_path"])
        output_table = output.tables[0]
        self.assertEqual(len(self.service._docx_row_cells(output_table.rows[0])), 2)
        self.assertEqual(output_table.cell(0, 1).text, "Beta LLC")
        self.assertEqual(output_table.cell(1, 1).text, "Jane Doe")
        self.assertNotIn("table_operations", result["fill_plan"])

    async def test_generate_docx_without_placeholders_uses_generated_code_for_single_row_context(self) -> None:
        handle = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        handle.close()
        self.addCleanup(lambda: os.path.exists(handle.name) and os.unlink(handle.name))
        target_path = handle.name
        temp_dir = tempfile.mkdtemp(prefix="form_fill_test_")
        self.addCleanup(lambda: shutil.rmtree(temp_dir, ignore_errors=True))
        run = SimpleNamespace(
            id="run-id",
            user_id="user-id",
            target_file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            target_filename="target.docx",
            target_gcs_object_name="target-object",
            allow_docx_table_expansion=False,
            fill_chronologically=True,
            output_format="docx",
        )
        document = Document()
        document.add_paragraph("Client: [blank]")
        document.save(target_path)

        converter = SimpleNamespace(convert_docx_gcs_to_pdf_gcs=AsyncMock())
        self.service.storage_service.construct_gcs_uri_for_object = MagicMock(return_value="gs://bucket/target")
        row_payload = {"client_name": "Beta LLC"}

        with patch.object(self.service, "_part_from_storage_object", return_value=object()):
            with patch("services.form_fill_service.get_document_conversion_service", return_value=converter):
                with patch.object(
                    self.service,
                    "_generate_and_execute_tabular_transform",
                    return_value={
                        "operations": [
                            {
                                "action": "replace_block_text",
                                "block_id": "body.paragraph.0",
                                "text": "Client: Beta LLC",
                            }
                        ],
                        "warnings": [],
                        "code_hash": "code-hash",
                    },
                ) as generate_tabular:
                    result = await self.service._generate_filled_document(
                        run=run,
                        temp_dir=temp_dir,
                        target_local_path=target_path,
                        source_parts=[],
                        source_text="single record",
                        tabular_context={
                            "columns": ["client_name"],
                            "rows": [row_payload],
                            "row_count": 1,
                            "single_record": True,
                        },
                    )

        generate_tabular.assert_called_once()
        self.assertEqual(generate_tabular.call_args.kwargs["rows"], [row_payload])
        self.assertEqual(result["strategy"], "docx_edit_generated_code")
        self.assertEqual(result["fill_plan"]["source_rows"], 1)
        self.assertEqual(result["fill_plan"]["code_hash"], "code-hash")
        output = Document(result["local_path"])
        self.assertEqual(output.paragraphs[0].text, "Client: Beta LLC")

    def test_normalize_repeat_mode_accepts_all_sources(self) -> None:
        self.assertEqual(_normalize_repeat_mode("all_sources"), "all_sources")

    def test_normalize_repeat_mode_defaults_to_all_sources(self) -> None:
        self.assertEqual(_normalize_repeat_mode(None), "all_sources")
        self.assertEqual(_normalize_repeat_mode("  "), "all_sources")
        self.assertEqual(_normalize_repeat_mode("single"), "single")

    async def test_process_run_all_sources_uses_combined_source_context_once(self) -> None:
        run_id = uuid.UUID("11111111-1111-1111-1111-111111111111")
        run = SimpleNamespace(
            id=run_id,
            user_id="user-id",
            status="pending",
            repeat_mode="all_sources",
            target_filename="target.pdf",
            target_gcs_object_name="target-object",
            target_file_type="application/pdf",
            target_page_count=2,
        )
        query = MagicMock()
        query.filter.return_value.first.return_value = run
        db = MagicMock()
        db.query.return_value = query
        self.service._get_session = MagicMock(return_value=db)
        self.service.storage_service.upload_file = AsyncMock()

        with patch.object(self.service, "_download_to_local", new=AsyncMock()):
            with patch.object(self.service, "_ensure_run_target_page_count", new=AsyncMock(return_value=2)):
                with patch.object(self.service, "_source_units_for_run", return_value=[{"record_index": 0}, {"record_index": 1}]) as source_units:
                    with patch.object(self.service, "_check_usage_limit_or_raise") as check_limit:
                        with patch.object(self.service, "_build_source_context", new=AsyncMock(return_value=([], "combined source"))) as build_context:
                            with patch.object(
                                self.service,
                                "_generate_filled_document",
                                new=AsyncMock(
                                    return_value={
                                        "local_path": "/tmp/filled.pdf",
                                        "filename": "target_filled.pdf",
                                        "mime_type": "application/pdf",
                                        "strategy": "fillable_pdf",
                                        "warnings": [],
                                        "fill_plan": {"strategy": "fillable_pdf"},
                                    }
                                ),
                            ) as generate:
                                with patch.object(self.service, "_record_usage_for_run") as record_usage:
                                    result = await self.service.process_run(str(run_id))

        source_units.assert_not_called()
        check_limit.assert_called_once_with(db, user_id="user-id", page_count=2)
        build_context.assert_awaited_once()
        generate.assert_awaited_once()
        record_usage.assert_called_once_with(db, run)
        self.assertTrue(result["success"])
        self.assertEqual(run.status, "completed")
        self.assertEqual(run.total_outputs, 1)
        self.assertEqual(run.completed_outputs, 1)
        self.assertEqual(run.usage_pages, 2)

    async def test_process_run_default_still_enqueues_multiple_source_units(self) -> None:
        run_id = uuid.UUID("11111111-1111-1111-1111-111111111111")
        run = SimpleNamespace(
            id=run_id,
            user_id="user-id",
            status="pending",
            repeat_mode="single",
            target_filename="target.pdf",
            target_gcs_object_name="target-object",
            target_file_type="application/pdf",
            target_page_count=2,
        )
        units = [
            {"record_index": 0, "record_label": "first", "record_payload": {"kind": "source_file"}},
            {"record_index": 1, "record_label": "second", "record_payload": {"kind": "source_file"}},
        ]
        query = MagicMock()
        query.filter.return_value.first.return_value = run
        db = MagicMock()
        db.query.return_value = query
        self.service._get_session = MagicMock(return_value=db)

        with patch.object(self.service, "_download_to_local", new=AsyncMock()):
            with patch.object(self.service, "_ensure_run_target_page_count", new=AsyncMock(return_value=2)):
                with patch.object(self.service, "_source_units_for_run", return_value=units) as source_units:
                    with patch.object(self.service, "_enqueue_output_units", new=AsyncMock(return_value={"success": True, "outputs": 2})) as enqueue:
                        with patch.object(self.service, "_build_source_context", new=AsyncMock()) as build_context:
                            result = await self.service.process_run(str(run_id))

        source_units.assert_called_once_with(run)
        enqueue.assert_awaited_once_with(
            db=db,
            run=run,
            target_page_count=2,
            units=units,
            strategy="source_files",
        )
        build_context.assert_not_called()
        self.assertEqual(result["outputs"], 2)

    async def test_process_output_source_rows_uses_single_record_context(self) -> None:
        run_id = uuid.UUID("11111111-1111-1111-1111-111111111111")
        output_id = uuid.UUID("22222222-2222-2222-2222-222222222222")
        row_payload = {"client_name": "Northstar Dental PLLC", "email": "northstar@example.com"}
        run = SimpleNamespace(
            id=run_id,
            user_id="user-id",
            repeat_mode="source_rows",
            target_filename="target.docx",
            target_gcs_object_name="target-object",
            target_file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        output = SimpleNamespace(
            id=output_id,
            run_id=run_id,
            status="pending",
            record_index=1,
            record_label="Northstar Dental PLLC",
            record_payload=row_payload,
            result_gcs_object_name=None,
        )
        run_query = MagicMock()
        run_query.filter.return_value.first.return_value = run
        output_query = MagicMock()
        output_query.filter.return_value.first.return_value = output
        db = MagicMock()
        db.query.side_effect = [run_query, output_query]
        self.service._get_session = MagicMock(return_value=db)
        self.service.storage_service.upload_file = AsyncMock()

        generated = {
            "local_path": "/tmp/filled.docx",
            "filename": "target_002_Northstar_Dental_PLLC_filled.docx",
            "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "strategy": "docx_placeholders",
            "warnings": [],
            "fill_plan": {"strategy": "docx_placeholders"},
        }

        with patch.object(self.service, "_try_advisory_lock", return_value=True):
            with patch.object(self.service, "_advisory_unlock"):
                with patch.object(self.service, "_download_to_local", new=AsyncMock()):
                    with patch.object(self.service, "_ensure_run_target_page_count", new=AsyncMock()):
                        with patch.object(self.service, "_build_tabular_source_context", new=AsyncMock()) as build_tabular:
                            with patch.object(self.service, "_build_output_source_context", new=AsyncMock(return_value=([], "single record"))) as build_output:
                                with patch.object(self.service, "_generate_filled_document", new=AsyncMock(return_value=generated)) as generate:
                                    with patch.object(self.service, "_sync_run_output_counts"):
                                        with patch.object(self.service, "_finalize_run_if_ready", new=AsyncMock()):
                                            result = await self.service.process_output(str(run_id), str(output_id))

        build_tabular.assert_not_awaited()
        build_output.assert_awaited_once_with(run=run, record_payload=row_payload, record_index=1)
        generate.assert_awaited_once()
        generate_kwargs = generate.await_args.kwargs
        tabular_context = generate_kwargs["tabular_context"]
        self.assertEqual(tabular_context["columns"], ["client_name", "email"])
        self.assertEqual(tabular_context["rows"], [row_payload])
        self.assertTrue(tabular_context["single_record"])
        self.assertEqual(generate_kwargs["source_text"], "single record")
        self.assertTrue(result["success"])


class FormFillServiceContinuationTests(unittest.TestCase):
    def setUp(self) -> None:
        self.service = FormFillService()

    def _response(
        self,
        *,
        text: str | None = None,
        parsed: dict[str, object] | None = None,
        finish_reason: str = "STOP",
        output_tokens: int = 10,
    ) -> SimpleNamespace:
        return SimpleNamespace(
            text=text,
            parsed=parsed,
            candidates=[SimpleNamespace(finish_reason=finish_reason)],
            usage_metadata=SimpleNamespace(candidates_token_count=output_tokens),
        )

    def test_salvages_complete_operations_from_truncated_json(self) -> None:
        text = '{"operations":[{"action":"replace_block_text","block_id":"a"},{"action":"insert_after_block","block_id":"b"},'

        operations = self.service._salvage_collection_from_text(text, "operations")

        self.assertEqual(
            operations,
            [
                {"action": "replace_block_text", "block_id": "a"},
                {"action": "insert_after_block", "block_id": "b"},
            ],
        )

    def test_collection_continuation_merges_overlap_without_duplication(self) -> None:
        self.service.client = SimpleNamespace(
            models=SimpleNamespace(
                generate_content=MagicMock(
                    side_effect=[
                        self._response(
                            text='{"operations":[{"action":"replace_block_text","block_id":"a"},{"action":"insert_after_block","block_id":"b"},',
                            finish_reason="MAX_TOKENS",
                        ),
                        self._response(
                            text=(
                                '{"operations":['
                                '{"action":"insert_after_block","block_id":"b"},'
                                '{"action":"append_to_block","block_id":"c"}'
                                '],"warnings":["continued"]}'
                            )
                        ),
                    ]
                )
            )
        )

        payload = self.service._generate_collection_json_response(
            [],
            prompt="Edit the DOCX.",
            schema=self.service._docx_edit_schema(),
            collection_key="operations",
            label="test",
        )

        self.assertEqual(
            payload["operations"],
            [
                {"action": "replace_block_text", "block_id": "a"},
                {"action": "insert_after_block", "block_id": "b"},
                {"action": "append_to_block", "block_id": "c"},
            ],
        )
        self.assertEqual(payload["warnings"], ["continued"])
        self.assertEqual(self.service.client.models.generate_content.call_count, 2)
        continuation_contents = self.service.client.models.generate_content.call_args_list[1].kwargs["contents"]
        self.assertIn(
            'prior_entries (all entries already returned, in order): [{"action":"replace_block_text","block_id":"a"},{"action":"insert_after_block","block_id":"b"}]',
            continuation_contents[-1],
        )
        self.assertIn("Do not summarize, collapse, or replace remaining entries", continuation_contents[-1])
        self.assertIn("Do not write 'see attached'", continuation_contents[-1])

    def test_collection_continuation_continues_on_full_batch_without_truncation(self) -> None:
        self.service.batch_items_per_call = 2
        self.service.client = SimpleNamespace(
            models=SimpleNamespace(
                generate_content=MagicMock(
                    side_effect=[
                        self._response(
                            parsed={
                                "operations": [
                                    {"action": "replace_block_text", "block_id": "a"},
                                    {"action": "insert_after_block", "block_id": "b"},
                                ],
                                "warnings": [],
                            },
                            finish_reason="STOP",
                            output_tokens=10,
                        ),
                        self._response(
                            parsed={
                                "operations": [
                                    {"action": "append_to_block", "block_id": "c"},
                                ],
                                "warnings": [],
                            },
                            finish_reason="STOP",
                            output_tokens=10,
                        ),
                    ]
                )
            )
        )

        payload = self.service._generate_collection_json_response(
            [],
            prompt="Edit the DOCX.",
            schema=self.service._docx_edit_schema(),
            collection_key="operations",
            label="test",
            continue_on_full_batch=True,
        )

        self.assertEqual(
            payload["operations"],
            [
                {"action": "replace_block_text", "block_id": "a"},
                {"action": "insert_after_block", "block_id": "b"},
                {"action": "append_to_block", "block_id": "c"},
            ],
        )
        self.assertEqual(self.service.client.models.generate_content.call_count, 2)

    def test_collection_batching_does_not_continue_on_clean_non_full_batch(self) -> None:
        self.service.batch_items_per_call = 5
        self.service.client = SimpleNamespace(
            models=SimpleNamespace(
                generate_content=MagicMock(
                    return_value=self._response(
                        parsed={
                            "operations": [{"action": "replace_block_text", "block_id": "a"}],
                            "warnings": [],
                        },
                        finish_reason="STOP",
                        output_tokens=10,
                    )
                )
            )
        )

        payload = self.service._generate_collection_json_response(
            [],
            prompt="Edit the DOCX.",
            schema=self.service._docx_edit_schema(),
            collection_key="operations",
            label="test",
            continue_on_full_batch=True,
        )

        self.assertEqual(
            payload["operations"],
            [
                {"action": "replace_block_text", "block_id": "a"},
            ],
        )
        self.assertEqual(self.service.client.models.generate_content.call_count, 1)

    def test_collection_filters_output_limit_warnings_for_continuable_outputs(self) -> None:
        self.service.client = SimpleNamespace(
            models=SimpleNamespace(
                generate_content=MagicMock(
                    side_effect=[
                        self._response(
                            parsed={
                                "operations": [{"action": "append_to_block", "block_id": "a"}],
                                "warnings": [
                                    "The source material contains over 500 transactions and would exceed output limits.",
                                    "Unable to calculate exact totals due to the large volume of transactions.",
                                    "Due to the large number of transactions, only the first 100 chronological transactions have been inserted. Please request continuation to insert the remaining rows.",
                                    "Missing date for one transaction.",
                                ],
                            }
                        ),
                        self._response(parsed={"operations": [], "warnings": []}),
                    ]
                )
            )
        )

        payload = self.service._generate_collection_json_response(
            [],
            prompt="Edit the DOCX.",
            schema=self.service._docx_edit_schema(),
            collection_key="operations",
            label="test",
            continue_on_full_batch=True,
        )

        self.assertEqual(payload["warnings"], ["Missing date for one transaction."])

    def test_collection_continuation_does_not_continue_full_batch_when_disabled(self) -> None:
        self.service.batch_items_per_call = 2
        self.service.client = SimpleNamespace(
            models=SimpleNamespace(
                generate_content=MagicMock(
                    return_value=self._response(
                        parsed={
                            "items": [
                                {"name": "A", "value": "1"},
                                {"name": "B", "value": "2"},
                            ],
                            "warnings": [],
                        },
                        finish_reason="STOP",
                        output_tokens=10,
                    )
                )
            )
        )

        payload = self.service._generate_collection_json_response(
            [],
            prompt="Map known fields.",
            schema=self.service._field_mapping_schema(),
            collection_key="items",
            label="test",
        )

        self.assertEqual(payload["items"], [{"name": "A", "value": "1"}, {"name": "B", "value": "2"}])
        self.assertEqual(self.service.client.models.generate_content.call_count, 1)

    def test_mapping_payload_chunks_known_names(self) -> None:
        self.service.mapping_chunk_size = 2
        self.service.client = SimpleNamespace(
            models=SimpleNamespace(
                generate_content=MagicMock(
                    side_effect=[
                        self._response(parsed={"items": [{"name": "A", "value": "1"}, {"name": "B", "value": "2"}], "warnings": []}),
                        self._response(parsed={"items": [{"name": "C", "value": "3"}], "warnings": ["check C"]}),
                    ]
                )
            )
        )

        payload = self.service._generate_mapping_payload(
            [],
            source_text="Source",
            mapping_items=["A", "B", "C"],
            mapping_label="Names",
            target_hint="target",
            label="mapping",
        )

        self.assertEqual(payload["items"], [{"name": "A", "value": "1"}, {"name": "B", "value": "2"}, {"name": "C", "value": "3"}])
        self.assertEqual(payload["warnings"], ["check C"])
        self.assertEqual(self.service.client.models.generate_content.call_count, 2)
        first_prompt = self.service.client.models.generate_content.call_args_list[0].kwargs["contents"][-1]
        second_prompt = self.service.client.models.generate_content.call_args_list[1].kwargs["contents"][-1]
        self.assertIn("- A", first_prompt)
        self.assertIn("- B", first_prompt)
        self.assertNotIn("- C", first_prompt)
        self.assertIn("- C", second_prompt)

    def test_form_fill_prompts_forbid_output_limit_workarounds(self) -> None:
        pdf_prompt = self.service._build_pdf_overlay_prompt(source_text="Rows", target_preview_text="Target")
        docx_prompt = self.service._build_docx_edit_prompt(
            source_text="Rows",
            block_summary="Blocks",
            table_summary="Tables",
            target_preview_text="Target",
            allow_table_expansion=True,
        )

        self.assertIn("Do not summarize, collapse, or omit entries", pdf_prompt)
        self.assertIn("Do not write \"see attached\"", pdf_prompt)
        self.assertIn("Do not ask the user to request continuation", pdf_prompt)
        self.assertIn("relative_position: auto, center, right, left, below, or above", pdf_prompt)
        self.assertIn("Auto prefers center, then start, then end", pdf_prompt)
        self.assertIn("emit one insert_table_row_after operation per source row", docx_prompt)
        self.assertIn("Do not use operation as the key name", docx_prompt)
        self.assertIn("Calculate totals", docx_prompt)
        self.assertIn("Do not insert only the first N rows", docx_prompt)
        self.assertIn("Do not claim totals cannot be calculated", docx_prompt)

    def _generated_code_prompt(self, *, fill_chronologically: bool) -> str:
        return self.service._build_generated_code_prompt(
            tabular_context={"columns": ["Date", "Amount"], "rows": [{"Date": "2024-01-01", "Amount": "10"}]},
            target_kind="DOCX edit in place",
            target_context={"target_kind": "DOCX edit in place"},
            output_contract="Return {'operations': [...], 'warnings': [...]}.",
            fill_chronologically=fill_chronologically,
        )

    def test_generated_code_prompt_includes_chronological_instruction_when_enabled(self) -> None:
        prompt = self._generated_code_prompt(fill_chronologically=True)
        self.assertIn("Order the emitted entries chronologically by date, oldest first", prompt)
        self.assertIn("parse_date(", prompt)
        self.assertIn("no date column exists", prompt)
        self.assertIn("do not emit the literal token DATE_COL", prompt)

    def test_generated_code_prompt_requires_target_driven_conditional_mapping(self) -> None:
        prompt = self._generated_code_prompt(fill_chronologically=False)
        self.assertIn("printed labels and instructions as authoritative business rules", prompt)
        self.assertIn("Generate conditional logic when a target instruction makes field selection depend", prompt)
        self.assertIn("do not map fields solely by source-column name similarity", prompt)

    def test_generated_code_prompt_omits_chronological_instruction_when_disabled(self) -> None:
        prompt = self._generated_code_prompt(fill_chronologically=False)
        self.assertNotIn("chronologically", prompt)

    def test_docx_edit_prompt_chronological_instruction_is_gated(self) -> None:
        kwargs = dict(
            source_text="Rows",
            block_summary="Blocks",
            table_summary="Tables",
            target_preview_text="Target",
            allow_table_expansion=True,
        )
        enabled = self.service._build_docx_edit_prompt(**kwargs, fill_chronologically=True)
        disabled = self.service._build_docx_edit_prompt(**kwargs, fill_chronologically=False)
        self.assertIn("Emit table rows and repeated entries in chronological order", enabled)
        self.assertNotIn("chronological order", disabled)

    def test_pdf_overlay_prompt_chronological_instruction_is_gated(self) -> None:
        enabled = self.service._build_pdf_overlay_prompt(
            source_text="Rows", target_preview_text="Target", fill_chronologically=True
        )
        disabled = self.service._build_pdf_overlay_prompt(
            source_text="Rows", target_preview_text="Target", fill_chronologically=False
        )
        self.assertIn("order them chronologically", enabled)
        self.assertNotIn("chronologically", disabled)

    def test_parse_date_handles_common_formats(self) -> None:
        self.assertEqual(_parse_date("2024-03-05"), (2024, 3, 5))
        self.assertEqual(_parse_date("2024/03/05"), (2024, 3, 5))
        self.assertEqual(_parse_date("3/5/2024"), (2024, 3, 5))
        self.assertEqual(_parse_date("03-05-2024"), (2024, 3, 5))
        self.assertEqual(_parse_date("1/2/99"), (1999, 1, 2))
        self.assertEqual(_parse_date("1/2/24"), (2024, 1, 2))
        self.assertEqual(_parse_date("Jan 5, 2024"), (2024, 1, 5))
        self.assertEqual(_parse_date("January 5 2024"), (2024, 1, 5))
        self.assertIsNone(_parse_date("not a date"))
        self.assertIsNone(_parse_date(""))
        self.assertIsNone(_parse_date(None))

    def test_generated_transform_can_sort_rows_with_parse_date(self) -> None:
        code = """
def transform(rows, context):
    ordered = sorted(
        rows,
        key=lambda r: (parse_date(r.get("Date")) is None, parse_date(r.get("Date")) or (0, 0, 0)),
    )
    return {
        "operations": [
            {"action": "insert_table_row_after", "table_id": "table.0", "row_index": i, "cells": [as_text(r.get("Date"))]}
            for i, r in enumerate(ordered)
        ],
        "warnings": [],
    }
"""
        result = self.service._execute_generated_transform(
            code,
            rows=[
                {"Date": "3/15/2024"},
                {"Date": "1/2/2024"},
                {"Date": ""},
                {"Date": "2/1/2024"},
            ],
            context={"target_kind": "DOCX edit in place"},
        )
        ordered_dates = [op["cells"][0] for op in result["operations"]]
        self.assertEqual(ordered_dates, ["1/2/2024", "2/1/2024", "3/15/2024", ""])

    def test_generated_transform_allows_lambda_keys(self) -> None:
        code = """
def transform(rows, context):
    ordered = sorted(rows, key=lambda r: parse_number(r.get("Amount")) or 0)
    return {"operations": [{"action": "x", "cells": [as_text(r.get("Amount"))]} for r in ordered], "warnings": []}
"""
        result = self.service._execute_generated_transform(
            code,
            rows=[{"Amount": "30"}, {"Amount": "10"}, {"Amount": "20"}],
            context={},
        )
        self.assertEqual([op["cells"][0] for op in result["operations"]], ["10", "20", "30"])

    def test_generated_transform_blocks_dunder_access_inside_lambda(self) -> None:
        code = """
def transform(rows, context):
    return {"operations": sorted(rows, key=lambda r: r.__class__), "warnings": []}
"""
        with self.assertRaisesRegex(ValueError, "dunder"):
            self.service._execute_generated_transform(code, rows=[], context={})

    def test_generated_transform_blocks_banned_name_inside_lambda(self) -> None:
        code = """
def transform(rows, context):
    return {"operations": sorted(rows, key=lambda r: getattr(r, "amount")), "warnings": []}
"""
        with self.assertRaisesRegex(ValueError, "getattr"):
            self.service._execute_generated_transform(code, rows=[], context={})

    def test_compact_fill_plan_stores_counts_and_samples(self) -> None:
        compact = self.service._compact_fill_plan(
            {
                "strategy": "docx_edit_in_place",
                "operations": [{"action": str(index)} for index in range(25)],
                "field_values": {str(index): index for index in range(60)},
            }
        )

        self.assertEqual(compact["strategy"], "docx_edit_in_place")
        self.assertEqual(compact["operations_count"], 25)
        self.assertEqual(len(compact["operations_sample"]), 20)
        self.assertEqual(compact["field_values_count"], 60)
        self.assertEqual(len(compact["field_values_sample"]), 50)

    def test_extraction_tabular_context_converts_rows_to_dicts(self) -> None:
        context = self.service._extraction_tabular_context(
            {
                "columns": ["Name", "Amount"],
                "rows": [["Checking", "$100.50"], ["Savings", "250"]],
                "source_files": ["source.pdf"],
            }
        )

        self.assertIsNotNone(context)
        self.assertEqual(context["row_count"], 2)
        self.assertEqual(context["columns"], ["Name", "Amount"])
        self.assertEqual(context["rows"][0]["Name"], "Checking")
        self.assertEqual(context["rows"][0]["_source_file"], "extraction results")
        self.assertEqual(context["source_files"], ["source.pdf"])

    def test_generated_transform_executes_against_all_rows(self) -> None:
        code = """
def transform(rows, context):
    operations = []
    for index, row in enumerate(rows):
        operations.append({
            "action": "insert_table_row_after",
            "table_id": "table.0",
            "row_index": index,
            "cells": [as_text(row.get("Name")), as_text(row.get("Amount"))],
        })
    return {"operations": operations, "warnings": []}
"""

        result = self.service._execute_generated_transform(
            code,
            rows=[{"Name": "Checking", "Amount": "100"}, {"Name": "Savings", "Amount": "250"}],
            context={"target_kind": "DOCX edit in place"},
        )

        self.assertEqual(len(result["operations"]), 2)
        self.assertEqual(result["operations"][1]["cells"], ["Savings", "250"])

    def test_generated_transform_returns_large_operation_list_without_queue_deadlock(self) -> None:
        code = """
def transform(rows, context):
    operations = []
    for index, row in enumerate(rows):
        operations.append({
            "action": "insert_table_row_after",
            "table_id": "table.0",
            "row_index": index,
            "cells": [as_text(row.get("Name")), as_text(row.get("Amount")), "x" * 500],
        })
    return {"operations": operations, "warnings": []}
"""

        result = self.service._execute_generated_transform(
            code,
            rows=[{"Name": f"Row {index}", "Amount": index} for index in range(1500)],
            context={"target_kind": "DOCX edit in place"},
        )

        self.assertEqual(len(result["operations"]), 1500)
        self.assertEqual(result["operations"][-1]["cells"][0], "Row 1499")

    def test_generated_transform_blocks_imports(self) -> None:
        with self.assertRaisesRegex(ValueError, "may not import"):
            self.service._execute_generated_transform(
                "import os\ndef transform(rows, context):\n    return {'operations': [], 'warnings': []}",
                rows=[],
                context={},
            )

    def test_generate_and_execute_tabular_transform_uses_one_code_generation_call(self) -> None:
        code = """
def transform(rows, context):
    return {
        "operations": [
            {"action": "insert_after_block", "block_id": "body.paragraph.0", "text": as_text(row.get("Name"))}
            for row in rows
        ],
        "warnings": [],
    }
"""
        self.service.client = SimpleNamespace(
            models=SimpleNamespace(
                generate_content=MagicMock(
                    return_value=self._response(parsed={"language": "python", "code": code, "warnings": []})
                )
            )
        )

        payload = self.service._generate_and_execute_tabular_transform(
            [],
            prompt="Generate code.",
            rows=[{"Name": "first"}, {"Name": "second"}, {"Name": "third"}],
            context={
                "target_kind": "DOCX edit in place",
                "source_columns": ["Name"],
                "output_contract": "Return operations.",
            },
            expected_key="operations",
            label="test_generated_code",
        )

        self.assertEqual([item["text"] for item in payload["operations"]], ["first", "second", "third"])
        self.assertIn("code_hash", payload)
        self.assertEqual(self.service.client.models.generate_content.call_count, 1)


class FormFillServiceUsageTests(unittest.IsolatedAsyncioTestCase):
    def setUp(self) -> None:
        self.service = FormFillService()

    def _pdf_bytes(self, pages: int) -> bytes:
        handle = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        handle.close()
        self.addCleanup(lambda: os.path.exists(handle.name) and os.unlink(handle.name))
        writer = PdfWriter()
        for _ in range(pages):
            writer.add_blank_page(width=72, height=72)
        with open(handle.name, "wb") as output:
            writer.write(output)
        with open(handle.name, "rb") as source:
            return source.read()

    async def test_count_target_pages_from_pdf_bytes(self) -> None:
        page_count = await self.service._count_target_pages_from_bytes(
            content=self._pdf_bytes(3),
            filename="target.pdf",
            mime_type="application/pdf",
        )

        self.assertEqual(page_count, 3)

    def test_check_usage_limit_raises_with_clear_message(self) -> None:
        billing_service = MagicMock()
        billing_service.check_page_limit.return_value = False
        billing_service.get_billing_info.return_value = {
            "plan_display_name": "Free",
            "pages_used": 8,
            "pages_included": 10,
        }

        with patch("services.billing_service.get_billing_service", return_value=billing_service):
            with self.assertRaisesRegex(ValueError, "processing 3 target pages"):
                self.service._check_usage_limit_or_raise(MagicMock(), user_id="user-id", page_count=3)

    def test_record_usage_for_run_uses_form_fill_run_source(self) -> None:
        billing_service = MagicMock()
        billing_service.record_usage.return_value = "event-id"
        run = SimpleNamespace(
            id="11111111-1111-1111-1111-111111111111",
            user_id="user-id",
            usage_pages=4,
            target_filename="target.pdf",
        )

        with patch("services.billing_service.get_billing_service", return_value=billing_service):
            self.service._record_usage_for_run(MagicMock(), run)

        billing_service.record_usage.assert_called_once_with(
            user_id="user-id",
            pages=4,
            source="form_fill_run",
            form_fill_run_id="11111111-1111-1111-1111-111111111111",
            notes="Form Fill run for target target.pdf",
        )


class FormFillOutputRetryTests(unittest.IsolatedAsyncioTestCase):
    def setUp(self) -> None:
        self.service = FormFillService()
        self.run_id = uuid.UUID("11111111-1111-1111-1111-111111111111")
        self.output_id = uuid.UUID("22222222-2222-2222-2222-222222222222")

    def _db_for_output(self, run: SimpleNamespace, output: SimpleNamespace) -> MagicMock:
        run_query = MagicMock()
        run_query.filter.return_value.first.return_value = run
        output_query = MagicMock()
        output_query.filter.return_value.first.return_value = output
        db = MagicMock()
        db.query.side_effect = [run_query, output_query, output_query, run_query]
        return db

    def test_generic_gemini_invalid_argument_is_retryable_before_final_attempt(self) -> None:
        self.assertTrue(
            self.service._should_retry_output_error(
                RetryableGeminiInvalidArgument(),
                task_retry_count=0,
                task_execution_count=1,
            )
        )

    def test_generic_gemini_invalid_argument_is_terminal_on_final_attempt(self) -> None:
        self.assertFalse(
            self.service._should_retry_output_error(
                RetryableGeminiInvalidArgument(),
                task_retry_count=self.service.output_max_attempts - 1,
                task_execution_count=self.service.output_max_attempts,
            )
        )

    def test_gemini_cancelled_is_retryable_before_final_attempt(self) -> None:
        self.assertTrue(
            self.service._should_retry_output_error(
                RetryableGeminiCancelled(),
                task_retry_count=0,
                task_execution_count=1,
            )
        )

    async def test_process_output_retries_retryable_error_before_final_attempt(self) -> None:
        run = SimpleNamespace(
            id=self.run_id,
            target_filename="target.docx",
            target_gcs_object_name="target-object",
        )
        output = SimpleNamespace(
            id=self.output_id,
            run_id=self.run_id,
            status="pending",
            record_payload={},
            record_index=0,
            record_label="record",
            error_message=None,
            completed_at=None,
        )
        db = self._db_for_output(run, output)

        with patch.object(self.service, "_get_session", return_value=db):
            with patch.object(self.service, "_try_advisory_lock", return_value=True):
                with patch.object(self.service, "_advisory_unlock"):
                    with patch.object(self.service, "_download_to_local", new=AsyncMock()):
                        with patch.object(self.service, "_ensure_run_target_page_count", new=AsyncMock(return_value=1)):
                            with patch.object(self.service, "_build_output_source_context", new=AsyncMock(return_value=([], "source"))):
                                with patch.object(self.service, "_generate_filled_document", new=AsyncMock(side_effect=RetryableGeminiInvalidArgument())):
                                    with patch.object(self.service, "_sync_run_output_counts") as sync_counts:
                                        with patch.object(self.service, "_finalize_run_if_ready", new=AsyncMock()) as finalize:
                                            with self.assertRaises(RetryableGeminiInvalidArgument):
                                                await self.service.process_output(
                                                    str(self.run_id),
                                                    str(self.output_id),
                                                    task_retry_count=0,
                                                    task_execution_count=1,
                                                    task_name="task-name",
                                                    task_queue_name="extract-tasks",
                                                )

        self.assertEqual(output.status, "pending")
        self.assertIn("Retrying after attempt 1 failed", output.error_message)
        self.assertIsNone(output.completed_at)
        sync_counts.assert_called_once_with(db, run)
        finalize.assert_not_awaited()

    async def test_process_output_marks_failed_on_final_retryable_attempt(self) -> None:
        run = SimpleNamespace(
            id=self.run_id,
            target_filename="target.docx",
            target_gcs_object_name="target-object",
        )
        output = SimpleNamespace(
            id=self.output_id,
            run_id=self.run_id,
            status="pending",
            record_payload={},
            record_index=0,
            record_label="record",
            error_message=None,
            completed_at=None,
        )
        db = self._db_for_output(run, output)

        with patch.object(self.service, "_get_session", return_value=db):
            with patch.object(self.service, "_try_advisory_lock", return_value=True):
                with patch.object(self.service, "_advisory_unlock"):
                    with patch.object(self.service, "_download_to_local", new=AsyncMock()):
                        with patch.object(self.service, "_ensure_run_target_page_count", new=AsyncMock(return_value=1)):
                            with patch.object(self.service, "_build_output_source_context", new=AsyncMock(return_value=([], "source"))):
                                with patch.object(self.service, "_generate_filled_document", new=AsyncMock(side_effect=RetryableGeminiInvalidArgument())):
                                    with patch.object(self.service, "_sync_run_output_counts"):
                                        with patch.object(self.service, "_finalize_run_if_ready", new=AsyncMock(return_value={"finalized": False})) as finalize:
                                            result = await self.service.process_output(
                                                str(self.run_id),
                                                str(self.output_id),
                                                task_retry_count=self.service.output_max_attempts - 1,
                                                task_execution_count=self.service.output_max_attempts,
                                            )

        self.assertEqual(result["status"], "failed")
        self.assertEqual(output.status, "failed")
        self.assertIn("INVALID_ARGUMENT", output.error_message)
        self.assertIsNotNone(output.completed_at)
        finalize.assert_awaited_once_with(str(self.run_id))


class FormFillFinalizeLockContentionTests(unittest.IsolatedAsyncioTestCase):
    """Losing the run-level advisory lock must never drop finalization: the
    lock holder may have read output statuses before the last output's commit
    and concluded "not ready", leaving the run in_progress forever."""

    def setUp(self) -> None:
        self.service = FormFillService()
        self.service.finalize_lock_wait_seconds = 1
        self.service.finalize_lock_poll_seconds = 0.01
        self.run_id = uuid.UUID("11111111-1111-1111-1111-111111111111")

    def _make_run(self, status: str = "in_progress", result: str | None = None) -> SimpleNamespace:
        return SimpleNamespace(
            id=self.run_id,
            user_id="user-1",
            status=status,
            result_gcs_object_name=result,
            result_filename=None,
            result_file_type=None,
            target_filename="fw9.pdf",
            target_page_count=6,
            total_outputs=2,
            completed_outputs=2,
            failed_outputs=0,
            usage_basis=None,
            usage_pages=None,
            warnings=None,
            fill_plan=None,
            processing_strategy=None,
            error_message=None,
            completed_at=None,
        )

    def _db(self, run: SimpleNamespace, outputs: list[SimpleNamespace]) -> MagicMock:
        run_query = MagicMock()
        run_query.filter.return_value.first.return_value = run
        output_query = MagicMock()
        output_query.filter.return_value.order_by.return_value.all.return_value = outputs
        db = MagicMock()
        db.query.side_effect = lambda model: run_query if model is FormFillRun else output_query
        return db

    async def test_finalize_waits_for_contended_lock_then_finalizes(self) -> None:
        run = self._make_run()
        outputs = [
            SimpleNamespace(
                id=uuid.uuid4(),
                record_index=index,
                record_label=f"row {index}",
                status="completed",
                result_gcs_object_name=f"outputs/{index}.pdf",
                result_filename=f"filled-{index}.pdf",
                warnings=None,
                error_message=None,
                fill_plan={"strategy": "fillable_pdf"},
            )
            for index in range(2)
        ]
        db = self._db(run, outputs)

        async def fake_download(object_name: str, local_path: str) -> None:
            Path(local_path).write_bytes(b"pdf-bytes")

        with patch.object(self.service, "_get_session", return_value=db), \
                patch.object(self.service, "_try_advisory_lock", side_effect=[False, True]) as try_lock, \
                patch.object(self.service, "_advisory_unlock"), \
                patch.object(self.service, "_sync_run_output_counts"), \
                patch.object(self.service, "_download_to_local", new=AsyncMock(side_effect=fake_download)), \
                patch.object(self.service.storage_service, "upload_file", new=AsyncMock()), \
                patch.object(self.service, "_record_usage_for_run"):
            result = await self.service._finalize_run_if_ready(str(self.run_id))

        self.assertTrue(result["finalized"])
        self.assertEqual(run.status, "completed")
        self.assertEqual(run.processing_strategy, "fillable_pdf")
        self.assertEqual(try_lock.call_count, 2)

    async def test_finalize_skips_when_other_worker_finalized_while_waiting(self) -> None:
        run = self._make_run(status="completed", result="form-fill/user-1/runs/r/result.zip")
        db = self._db(run, [])

        with patch.object(self.service, "_get_session", return_value=db), \
                patch.object(self.service, "_try_advisory_lock", return_value=False), \
                patch.object(self.service, "_advisory_unlock"), \
                patch("services.form_fill_service.cloud_run_task_service.enqueue_form_fill_task", new_callable=AsyncMock) as enqueue:
            result = await self.service._finalize_run_if_ready(str(self.run_id))

        self.assertFalse(result["finalized"])
        self.assertTrue(result.get("skipped"))
        enqueue.assert_not_awaited()

    async def test_finalize_enqueues_retry_when_lock_stays_contended(self) -> None:
        run = self._make_run()
        db = self._db(run, [])
        self.service.finalize_lock_wait_seconds = 0.05

        with patch.object(self.service, "_get_session", return_value=db), \
                patch.object(self.service, "_try_advisory_lock", return_value=False), \
                patch.object(self.service, "_advisory_unlock"), \
                patch("services.form_fill_service.cloud_run_task_service.enqueue_form_fill_task", new_callable=AsyncMock) as enqueue:
            result = await self.service._finalize_run_if_ready(str(self.run_id))

        self.assertFalse(result["finalized"])
        self.assertTrue(result.get("locked"))
        enqueue.assert_awaited_once_with(
            str(self.run_id),
            delay_seconds=self.service.finalize_retry_delay_seconds,
        )


FW9_PDF_PATH = Path(__file__).resolve().parents[2] / "examples" / "form-fill" / "targets" / "fw9.pdf"


@unittest.skipUnless(FW9_PDF_PATH.exists(), f"fixture not found: {FW9_PDF_PATH}")
class FormFillFillablePdfCheckboxTests(unittest.TestCase):
    """Checkbox fill behavior on the IRS W-9 (fw9.pdf).

    fw9 has /Btn fields (c1_1[0..6], c1_2[0]) whose only valid states are /Off or
    a single on-state (/1../7). These tests verify that a valid on-state checks
    the box, an empty/Off/wrong-state value leaves it /Off, and that only the
    targeted box is checked. See _apply_fillable_pdf in form_fill_service.py.
    """

    def setUp(self) -> None:
        self.service = FormFillService()

    def _apply(self, field_values: dict) -> str:
        handle = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        handle.close()
        self.addCleanup(os.unlink, handle.name)
        self.service._apply_fillable_pdf(str(FW9_PDF_PATH), dict(field_values), handle.name)
        return handle.name

    @staticmethod
    def _widget(path: str, name: str):
        # Read the filled value/state directly off the widget annotation.
        reader = PdfReader(path)
        for page in reader.pages:
            annots = page.get("/Annots")
            if not annots:
                continue
            for ref in annots.get_object():
                obj = ref.get_object()
                if str(obj.get("/T")) == name:
                    return obj
        return None

    def test_empty_checkbox_does_not_crash_and_stays_off(self) -> None:
        out = self._apply({"f1_01[0]": "Acme LLC", "c1_1[0]": ""})
        self.assertEqual(str(self._widget(out, "c1_1[0]").get("/AS")), "/Off")
        self.assertEqual(str(self._widget(out, "f1_01[0]").get("/V")), "Acme LLC")

    def test_valid_on_state_checks_box(self) -> None:
        for value in ("1", "/1"):
            with self.subTest(value=value):
                out = self._apply({"c1_1[0]": value})
                widget = self._widget(out, "c1_1[0]")
                self.assertEqual(str(widget.get("/AS")), "/1")
                self.assertEqual(str(widget.get("/V")), "/1")

    def test_wrong_on_state_for_widget_is_dropped(self) -> None:
        # "1" is the on-state for c1_1[0], not c1_1[1] (whose on-state is "2").
        out = self._apply({"c1_1[1]": "1"})
        self.assertEqual(str(self._widget(out, "c1_1[1]").get("/AS")), "/Off")

    def test_off_value_is_dropped(self) -> None:
        out = self._apply({"c1_1[0]": "Off"})
        self.assertEqual(str(self._widget(out, "c1_1[0]").get("/AS")), "/Off")

    def test_setting_one_checkbox_leaves_others_off(self) -> None:
        out = self._apply({"c1_1[2]": "3"})
        self.assertEqual(str(self._widget(out, "c1_1[2]").get("/AS")), "/3")
        self.assertEqual(str(self._widget(out, "c1_1[0]").get("/AS")), "/Off")


class FormFillProcessRunRetryTests(unittest.IsolatedAsyncioTestCase):
    """process_run should record deterministic failures once (2xx, no Cloud Tasks
    retry) while still re-raising transient errors so they retry."""

    def setUp(self) -> None:
        self.service = FormFillService()
        self.run_id = uuid.UUID("11111111-1111-1111-1111-111111111111")

    def _run_and_db(self):
        run = SimpleNamespace(
            id=self.run_id,
            user_id="user-id",
            status="pending",
            repeat_mode="all_sources",
            target_filename="target.pdf",
            target_gcs_object_name="target-object",
            target_file_type="application/pdf",
            target_page_count=2,
            error_message=None,
            completed_at=None,
        )
        query = MagicMock()
        query.filter.return_value.first.return_value = run
        db = MagicMock()
        db.query.return_value = query
        return run, db

    async def _process_with_generate_error(self, exc: Exception, **task_kwargs):
        run, db = self._run_and_db()
        self.service._get_session = MagicMock(return_value=db)
        with patch.object(self.service, "_download_to_local", new=AsyncMock()):
            with patch.object(self.service, "_ensure_run_target_page_count", new=AsyncMock(return_value=2)):
                with patch.object(self.service, "_check_usage_limit_or_raise"):
                    with patch.object(self.service, "_build_tabular_source_context", new=AsyncMock(return_value=None)):
                        with patch.object(self.service, "_build_source_context", new=AsyncMock(return_value=([], "source"))):
                            with patch.object(self.service, "_generate_filled_document", new=AsyncMock(side_effect=exc)):
                                return run, await self.service.process_run(str(self.run_id), **task_kwargs)

    async def test_deterministic_failure_is_terminal_without_reraise(self) -> None:
        run, result = await self._process_with_generate_error(
            ValueError("bad mapping"), task_retry_count=0, task_execution_count=1
        )
        self.assertFalse(result["success"])
        self.assertTrue(result["terminal"])
        self.assertEqual(result["status"], "failed")
        self.assertEqual(run.status, "failed")
        self.assertEqual(run.error_message, "bad mapping")

    async def test_transient_failure_reraises_before_final_attempt(self) -> None:
        with self.assertRaises(RetryableGeminiInvalidArgument):
            await self._process_with_generate_error(
                RetryableGeminiInvalidArgument(), task_retry_count=0, task_execution_count=1
            )

    async def test_transient_failure_is_terminal_on_final_attempt(self) -> None:
        run, result = await self._process_with_generate_error(
            RetryableGeminiInvalidArgument(),
            task_retry_count=self.service.output_max_attempts - 1,
            task_execution_count=self.service.output_max_attempts,
        )
        self.assertTrue(result["terminal"])
        self.assertEqual(run.status, "failed")


def _build_radio_pdf(path: str, *, selected: str | None = None) -> None:
    """A single radio group ``entity_type`` with three mutually-exclusive options."""
    c = _reportlab_canvas.Canvas(path, pagesize=letter)
    form = c.acroForm
    c.drawString(100, 720, "Federal tax classification:")
    for value, label, y in (("individual", "Individual", 670), ("corp", "Corporation", 645), ("partnership", "Partnership", 620)):
        form.radio(name="entity_type", value=value, selected=(value == selected), x=100, y=y, size=15, buttonStyle="check")
        c.drawString(120, y + 2, label)
    c.save()


def _build_choice_pdf(path: str, *, multi: bool = False) -> None:
    """A choice field ``state`` with (export, display) options."""
    doc = fitz.open()
    page = doc.new_page()
    widget = fitz.Widget()
    widget.field_name = "state"
    widget.field_type = fitz.PDF_WIDGET_TYPE_LISTBOX if multi else fitz.PDF_WIDGET_TYPE_COMBOBOX
    widget.rect = fitz.Rect(100, 100, 260, 170 if multi else 122)
    widget.choice_values = [("AL", "Alabama"), ("SC", "South Carolina"), ("TX", "Texas")]
    if multi:
        widget.field_flags = 1 << 21
    page.add_widget(widget)
    doc.save(path)
    doc.close()


def _widget_by_name(path: str, name: str):
    reader = PdfReader(path)
    for page in reader.pages:
        annots = page.get("/Annots")
        if not annots:
            continue
        for ref in annots.get_object():
            obj = ref.get_object()
            if str(obj.get("/T")) == name:
                return obj
    return None


def _fitz_field_value(path: str, leaf: str):
    """Read a field's current value with an independent reader (PyMuPDF). For a
    radio group this is the export of the selected option (or None if unset)."""
    doc = fitz.open(path)
    try:
        selected = None
        for page in doc:
            for widget in page.widgets() or []:
                if widget.field_name.split(".")[-1] != leaf:
                    continue
                value = str(widget.field_value)
                if value and value.lower() != "off":
                    selected = value
        return selected
    finally:
        doc.close()


def _mapping_response(items: list[dict]) -> SimpleNamespace:
    return SimpleNamespace(
        text=None,
        parsed={"items": items, "warnings": []},
        candidates=[SimpleNamespace(finish_reason="STOP")],
        usage_metadata=SimpleNamespace(candidates_token_count=5),
    )


@unittest.skipUnless(FW9_PDF_PATH.exists(), f"fixture not found: {FW9_PDF_PATH}")
class FormFillPdfFieldMetadataTests(unittest.TestCase):
    """The fitz-based extractor must emit leaf names + type + valid values so the
    LLM can fill non-text fields correctly."""

    def setUp(self) -> None:
        self.service = FormFillService()
        self.metadata = self.service._extract_pdf_form_field_metadata(str(FW9_PDF_PATH))
        self.by_name = {m["name"]: m for m in self.metadata}

    def test_emits_leaf_names_that_join_to_write_path(self) -> None:
        self.assertIn("c1_1[0]", self.by_name)
        for meta in self.metadata:
            # Must be the leaf name (no dotted prefix) to match PyPDF2 field keys.
            self.assertNotIn(".", meta["name"])

    def test_checkbox_type_and_on_states(self) -> None:
        self.assertEqual(self.by_name["c1_1[0]"]["type"], "checkbox")
        self.assertEqual(self.by_name["c1_1[0]"]["on_states"], ["1"])
        self.assertEqual(self.by_name["c1_1[6]"]["on_states"], ["7"])

    def test_text_fields_have_no_value_constraints(self) -> None:
        text_fields = [m for m in self.metadata if m["type"] == "text"]
        self.assertTrue(text_fields)
        for meta in text_fields:
            self.assertNotIn("on_states", meta)
            self.assertNotIn("options", meta)

    def test_label_is_first_line_without_bleed(self) -> None:
        label = self.by_name["c1_1[0]"].get("label", "")
        self.assertTrue(label.startswith("Individual/sole proprietor"))
        self.assertNotIn("\n", label)
        self.assertNotIn("C corporation", label)


class FormFillMappingMetadataTests(unittest.TestCase):
    """Field metadata must reach the mapping prompt, while the no-metadata path
    stays byte-compatible with the previous behavior."""

    def setUp(self) -> None:
        self.service = FormFillService()

    def test_prompt_renders_nontext_constraints(self) -> None:
        meta = {
            "cb": {"name": "cb", "type": "checkbox", "on_states": ["1"], "label": "Individual"},
            "rd": {"name": "rd", "type": "radio", "on_states": ["a", "b"], "options": [{"export": "a", "label": "A"}, {"export": "b", "label": "B"}]},
            "dd": {"name": "dd", "type": "combobox", "options": [{"export": "SC", "label": "South Carolina"}]},
            "tx": {"name": "tx", "type": "text"},
        }
        prompt = self.service._build_mapping_prompt(
            source_text="src",
            mapping_items=list(meta),
            mapping_label="Fields",
            target_hint="fillable PDF form",
            field_metadata=meta,
        )
        self.assertIn('to check this box set value to "1"', prompt)
        self.assertIn('choose exactly one value: "a" = A; "b" = B', prompt)
        self.assertIn('"SC" = South Carolina', prompt)
        self.assertIn("never return the human-readable label", prompt)
        self.assertIn("printed labels and instructions as authoritative business rules", prompt)
        self.assertIn("do not map fields solely by source-name similarity", prompt)
        self.assertIn("- tx", prompt)
        self.assertNotIn("- tx\n    type", prompt)

    def test_prompt_without_metadata_is_unchanged(self) -> None:
        prompt = self.service._build_mapping_prompt(
            source_text="src", mapping_items=["A", "B"], mapping_label="Fields", target_hint="t"
        )
        self.assertIn("- A\n- B", prompt)
        self.assertNotIn("type:", prompt)
        self.assertNotIn("never return the human-readable label", prompt)

    def test_mapping_payload_applies_metadata_per_chunk(self) -> None:
        self.service.mapping_chunk_size = 2
        self.service.client = SimpleNamespace(
            models=SimpleNamespace(
                generate_content=MagicMock(
                    side_effect=[
                        _mapping_response([{"name": "c1_1[0]", "value": "1"}, {"name": "X", "value": "y"}]),
                        _mapping_response([{"name": "Z", "value": "z"}]),
                    ]
                )
            )
        )
        meta = {"c1_1[0]": {"name": "c1_1[0]", "type": "checkbox", "on_states": ["1"], "label": "Individual"}}
        self.service._generate_mapping_payload(
            [],
            source_text="src",
            mapping_items=["c1_1[0]", "X", "Z"],
            mapping_label="Fields",
            target_hint="fillable PDF form",
            label="m",
            field_metadata=meta,
        )
        first_prompt = self.service.client.models.generate_content.call_args_list[0].kwargs["contents"][-1]
        second_prompt = self.service.client.models.generate_content.call_args_list[1].kwargs["contents"][-1]
        self.assertIn('to check this box set value to "1"', first_prompt)
        self.assertNotIn("to check this box", second_prompt)


@unittest.skipUnless(_HAS_REPORTLAB, "reportlab not installed")
class FormFillRadioFieldTests(unittest.TestCase):
    """Radio groups must group into one logical field and fill the chosen kid."""

    def setUp(self) -> None:
        self.service = FormFillService()
        handle = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        handle.close()
        self.addCleanup(os.unlink, handle.name)
        self.path = handle.name
        _build_radio_pdf(self.path)

    def _apply(self, field_values: dict) -> str:
        out = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        out.close()
        self.addCleanup(os.unlink, out.name)
        self.service._apply_fillable_pdf(self.path, dict(field_values), out.name)
        return out.name

    def test_metadata_groups_widgets_into_one_radio_field(self) -> None:
        by_name = {m["name"]: m for m in self.service._extract_pdf_form_field_metadata(self.path)}
        field = by_name["entity_type"]
        self.assertEqual(field["type"], "radio")
        self.assertEqual(set(field["on_states"]), {"individual", "corp", "partnership"})
        self.assertEqual({opt["export"] for opt in field["options"]}, {"individual", "corp", "partnership"})

    def test_selecting_option_sets_radio_value(self) -> None:
        self.assertEqual(_fitz_field_value(self._apply({"entity_type": "corp"}), "entity_type"), "corp")

    def test_invalid_option_leaves_radio_unset(self) -> None:
        self.assertIsNone(_fitz_field_value(self._apply({"entity_type": "bogus"}), "entity_type"))


@unittest.skipUnless(_HAS_FITZ, "PyMuPDF not installed")
class FormFillChoiceFieldTests(unittest.TestCase):
    """Dropdown / list-box values must be constrained to real option exports."""

    def setUp(self) -> None:
        self.service = FormFillService()
        handle = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        handle.close()
        self.addCleanup(os.unlink, handle.name)
        self.path = handle.name
        _build_choice_pdf(self.path)

    def _apply_value(self, value, *, path: str | None = None):
        out = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        out.close()
        self.addCleanup(os.unlink, out.name)
        self.service._apply_fillable_pdf(path or self.path, {"state": value}, out.name)
        widget = _widget_by_name(out.name, "state")
        value = None if widget is None else widget.get("/V")
        return None if value is None else str(value)

    def test_metadata_extracts_options(self) -> None:
        by_name = {m["name"]: m for m in self.service._extract_pdf_form_field_metadata(self.path)}
        self.assertEqual(by_name["state"]["type"], "combobox")
        self.assertEqual(
            [(opt["export"], opt["label"]) for opt in by_name["state"]["options"]],
            [("AL", "Alabama"), ("SC", "South Carolina"), ("TX", "Texas")],
        )

    def test_export_value_is_written(self) -> None:
        self.assertEqual(self._apply_value("SC"), "SC")

    def test_display_value_maps_to_export(self) -> None:
        self.assertEqual(self._apply_value("Texas"), "TX")

    def test_value_not_in_options_is_dropped(self) -> None:
        self.assertIsNone(self._apply_value("ZZ"))

    def test_multi_select_uses_only_first_value(self) -> None:
        handle = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        handle.close()
        self.addCleanup(os.unlink, handle.name)
        _build_choice_pdf(handle.name, multi=True)
        self.assertEqual(self._apply_value("AL;TX", path=handle.name), "AL")


BLANK_RUN = "____________________"
OVERLAY_LABELS = ["Vendor legal name:", "Street address:", "Primary contact:"]


def _build_flat_blanks_pdf(path: str) -> None:
    """A flat (non-AcroForm) page with identical underscore blanks next to each label."""
    doc = fitz.open()
    page = doc.new_page()
    for index, label in enumerate(OVERLAY_LABELS):
        y = 100 + index * 40
        page.insert_text((72, y), label, fontsize=11, fontname="helv")
        page.insert_text((300, y), BLANK_RUN, fontsize=11, fontname="helv")
    doc.save(path)
    doc.close()


@unittest.skipUnless(_HAS_FITZ, "PyMuPDF not installed")
class FormFillPdfOverlayAnchorTests(unittest.TestCase):
    """Overlay items anchored to repeated identical blanks must land on distinct lines."""

    def setUp(self) -> None:
        self.service = FormFillService()
        handle = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        handle.close()
        self.addCleanup(os.unlink, handle.name)
        self.path = handle.name
        _build_flat_blanks_pdf(self.path)

    def _apply_items(self, items: list[dict]) -> tuple[list[str], dict[str, float]]:
        out = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        out.close()
        self.addCleanup(os.unlink, out.name)
        warnings = self.service._apply_pdf_overlay_plan(self.path, items, out.name)
        doc = fitz.open(out.name)
        try:
            centers: dict[str, float] = {}
            for block in doc[0].get_text("dict")["blocks"]:
                for line in block.get("lines", []):
                    for span in line["spans"]:
                        centers[span["text"]] = (span["bbox"][1] + span["bbox"][3]) / 2
        finally:
            doc.close()
        return warnings, centers

    def _assert_on_label_line(self, centers: dict[str, float], overlay_text: str, label: str) -> None:
        self.assertIn(overlay_text, centers)
        self.assertIn(label, centers)
        self.assertLess(
            abs(centers[overlay_text] - centers[label]),
            6.0,
            f"{overlay_text!r} not on the {label!r} line",
        )

    def test_overlay_contract_supports_explicit_anchor_occurrences(self) -> None:
        schema = self.service._pdf_overlay_schema().model_dump(exclude_none=True, mode="json")
        item_properties = schema["properties"]["items"]["items"]["properties"]
        self.assertIn("match_index", item_properties)
        prompt = self.service._build_pdf_overlay_prompt(source_text="Source", target_preview_text="Target")
        self.assertIn("zero-based reading-order occurrence", prompt)
        self.assertIn("skipped rather than guessed", prompt)

    def test_anchor_before_selects_the_matching_blank(self) -> None:
        # Deliberately out of document order: the label must pick the blank, not emission order.
        items = [
            {
                "page_number": 1,
                "anchor_text": BLANK_RUN,
                "anchor_before": "Street address:",
                "overlay_text": "480 Skyline Terrace",
                "placement_hint": "replace_anchor",
            },
            {
                "page_number": 1,
                "anchor_text": BLANK_RUN,
                "anchor_before": "Primary contact:",
                "overlay_text": "Jordan Ellis",
                "placement_hint": "replace_anchor",
            },
            {
                "page_number": 1,
                "anchor_text": BLANK_RUN,
                "anchor_before": "Vendor legal name:",
                "overlay_text": "Brightpath Consulting LLC",
                "placement_hint": "replace_anchor",
            },
        ]

        warnings, centers = self._apply_items(items)

        self.assertEqual(warnings, [])
        self._assert_on_label_line(centers, "480 Skyline Terrace", "Street address:")
        self._assert_on_label_line(centers, "Jordan Ellis", "Primary contact:")
        self._assert_on_label_line(centers, "Brightpath Consulting LLC", "Vendor legal name:")

    def test_anchor_after_next_label_still_selects_preceding_blank(self) -> None:
        # Regression: the model sometimes disambiguates a blank with the next
        # field's label in anchor_after instead of the field's own label in anchor_before.
        items = [
            {
                "page_number": 1,
                "anchor_text": BLANK_RUN,
                "anchor_after": "Street address:",
                "overlay_text": "Cedar Bookkeeping Services",
                "placement_hint": "replace_anchor",
            },
            {
                "page_number": 1,
                "anchor_text": BLANK_RUN,
                "anchor_after": "Primary contact:",
                "overlay_text": "915 Alder Street",
                "placement_hint": "replace_anchor",
            },
        ]

        warnings, centers = self._apply_items(items)

        self.assertEqual(warnings, [])
        self._assert_on_label_line(centers, "Cedar Bookkeeping Services", "Vendor legal name:")
        self._assert_on_label_line(centers, "915 Alder Street", "Street address:")

    def test_repeated_anchor_without_context_is_skipped(self) -> None:
        items = [
            {
                "page_number": 1,
                "anchor_text": BLANK_RUN,
                "overlay_text": f"Value {index}",
                "placement_hint": "replace_anchor",
            }
            for index in range(len(OVERLAY_LABELS))
        ]

        warnings, centers = self._apply_items(items)

        self.assertEqual(len(warnings), len(OVERLAY_LABELS))
        self.assertTrue(all("appears more than once" in warning for warning in warnings))
        for index in range(len(OVERLAY_LABELS)):
            self.assertNotIn(f"Value {index}", centers)

    def test_match_indexes_place_repeated_anchors_in_reading_order(self) -> None:
        items = [
            {
                "page_number": 1,
                "anchor_text": BLANK_RUN,
                "match_index": index,
                "overlay_text": f"Value {index}",
                "placement_hint": "replace_anchor",
            }
            for index in range(len(OVERLAY_LABELS))
        ]

        warnings, centers = self._apply_items(items)

        self.assertEqual(warnings, [])
        for index, label in enumerate(OVERLAY_LABELS):
            self._assert_on_label_line(centers, f"Value {index}", label)

    def test_conflicting_context_and_match_index_is_skipped(self) -> None:
        items = [{
            "page_number": 1,
            "anchor_text": BLANK_RUN,
            "anchor_before": "Street address:",
            "match_index": 0,
            "overlay_text": "Wrong row",
            "placement_hint": "replace_anchor",
        }]

        warnings, centers = self._apply_items(items)

        self.assertEqual(len(warnings), 1)
        self.assertIn("conflicts", warnings[0])
        self.assertNotIn("Wrong row", centers)

    def test_case_sensitive_anchor_filter_matches_esign_semantics(self) -> None:
        doc = fitz.open(self.path)
        try:
            page = doc[0]
            insensitive = self.service._resolve_pdf_anchor_rect(
                page,
                {"anchor_text": "vendor legal name:"},
            )
            sensitive = self.service._resolve_pdf_anchor_rect(
                page,
                {"anchor_text": "vendor legal name:", "case_sensitive": True},
            )
        finally:
            doc.close()

        self.assertIsNotNone(insensitive)
        self.assertIsNone(sensitive)

    def test_repeated_context_anchors_use_the_nearest_valid_pair(self) -> None:
        doc = fitz.open()
        page = doc.new_page()
        for y in (100, 180):
            page.insert_text((72, y), "Account", fontsize=11)
            page.insert_text((170, y), BLANK_RUN, fontsize=11)
        try:
            matches = page.search_for(BLANK_RUN)
            chosen = self.service._pick_pdf_anchor_match(
                matches,
                page.search_for("Account"),
                [fitz.Rect(400, 170, 450, 190)],
            )
        finally:
            doc.close()

        self.assertAlmostEqual(chosen.y0, matches[1].y0)


@unittest.skipUnless(_HAS_FITZ, "PyMuPDF not installed")
class FormFillPdfOverlayRelativeAnchorTests(unittest.TestCase):
    def setUp(self) -> None:
        self.service = FormFillService()

    def test_geometry_has_esign_auto_placement_and_edge_fallback(self) -> None:
        small = self.service._relative_anchor_field_position(
            0.75,
            0.4,
            0.05,
            0.05,
            relative_position="auto",
            cross_axis_alignment="auto",
            field_width=0.1,
            field_height=0.1,
        )
        large = self.service._relative_anchor_field_position(
            0.75,
            0.4,
            0.05,
            0.05,
            relative_position="auto",
            cross_axis_alignment="auto",
            field_width=0.25,
            field_height=0.1,
        )

        self.assertAlmostEqual(small[0], 0.8)
        self.assertAlmostEqual(large[0], 0.5)

    def test_relative_offsets_and_units_are_applied(self) -> None:
        doc = fitz.open()
        page = doc.new_page(width=612, height=792)
        page.insert_text((72, 200), "Anchor", fontsize=11)
        anchor = page.search_for("Anchor")[0]
        base = self.service._relative_target_rect_from_anchor(
            page,
            anchor,
            {
                "overlay_text": "Value",
                "relative_position": "below",
                "cross_axis_alignment": "start",
            },
        )
        offset = self.service._relative_target_rect_from_anchor(
            page,
            anchor,
            {
                "overlay_text": "Value",
                "relative_position": "below",
                "cross_axis_alignment": "start",
                "offset_x": 1,
                "offset_y": 0.5,
                "offset_unit": "inch",
            },
        )
        doc.close()

        self.assertAlmostEqual(offset.x0 - base.x0, 72.0)
        self.assertAlmostEqual(offset.y0 - base.y0, 36.0)

    def test_rotated_page_uses_display_coordinates_and_stays_bounded(self) -> None:
        doc = fitz.open()
        page = doc.new_page(width=612, height=792)
        page.insert_text((500, 700), "EDGE_ANCHOR", fontsize=11)
        anchor = page.search_for("EDGE_ANCHOR")[0]
        page.set_rotation(90)
        target = self.service._relative_target_rect_from_anchor(
            page,
            anchor,
            {
                "overlay_text": "Value",
                "relative_position": "auto",
                "cross_axis_alignment": "auto",
                "field_width": 100,
                "field_height": 30,
            },
        )
        display_target = fitz.Rect(target) * page.rotation_matrix
        display_target.normalize()
        page_rect = fitz.Rect(page.rect)
        doc.close()

        self.assertGreaterEqual(display_target.x0, page_rect.x0)
        self.assertGreaterEqual(display_target.y0, page_rect.y0)
        self.assertLessEqual(display_target.x1, page_rect.x1)
        self.assertLessEqual(display_target.y1, page_rect.y1)

    def test_relative_overlay_renders_at_every_page_rotation(self) -> None:
        for rotation in (0, 90, 180, 270):
            with self.subTest(rotation=rotation):
                source = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
                source.close()
                output = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
                output.close()
                self.addCleanup(os.unlink, source.name)
                self.addCleanup(os.unlink, output.name)
                doc = fitz.open()
                page = doc.new_page(width=612, height=792)
                page.insert_text((72, 200), "ANCHOR", fontsize=11)
                page.set_rotation(rotation)
                doc.save(source.name)
                doc.close()

                warnings = self.service._apply_pdf_overlay_plan(
                    source.name,
                    [{
                        "page_number": 1,
                        "anchor_text": "ANCHOR",
                        "overlay_text": "VALUE",
                        "relative_position": "auto",
                        "cross_axis_alignment": "auto",
                    }],
                    output.name,
                )

                self.assertEqual(warnings, [])
                rendered = fitz.open(output.name)
                try:
                    self.assertEqual(len(rendered[0].search_for("VALUE")), 1)
                finally:
                    rendered.close()


@unittest.skipUnless(_HAS_FITZ, "PyMuPDF not installed")
class FormFillPdfOverlayReplaceAnchorPlacementTests(unittest.TestCase):
    """replace_anchor must sit on the blank and its baseline, not drift onto
    the label to the left of it or float above the line."""

    CHECKBOX_BASELINE_Y = 100.0

    def setUp(self) -> None:
        self.service = FormFillService()
        handle = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        handle.close()
        self.addCleanup(os.unlink, handle.name)
        self.path = handle.name
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text(
            (72, self.CHECKBOX_BASELINE_Y),
            "Subject to backup withholding?  Yes ____  No ____",
            fontsize=11,
            fontname="helv",
        )
        doc.save(self.path)
        doc.close()

    def test_x_on_short_blank_stays_on_blank_and_baseline(self) -> None:
        base = fitz.open(self.path)
        no_label = base[0].search_for("No")[0]
        no_blank = [rect for rect in base[0].search_for("____") if rect.x0 > no_label.x1][0]
        base.close()

        out = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        out.close()
        self.addCleanup(os.unlink, out.name)
        items = [
            {
                "page_number": 1,
                "anchor_text": "____",
                "anchor_before": "No",
                "overlay_text": "X",
                "placement_hint": "replace_anchor",
            }
        ]

        warnings = self.service._apply_pdf_overlay_plan(self.path, items, out.name)

        self.assertEqual(warnings, [])
        doc = fitz.open(out.name)
        try:
            x_span = None
            for block in doc[0].get_text("dict")["blocks"]:
                for line in block.get("lines", []):
                    for span in line["spans"]:
                        if span["text"] == "X":
                            x_span = span
            self.assertIsNotNone(x_span, "overlay X was not rendered")
            self.assertGreaterEqual(x_span["bbox"][0], no_blank.x0, "X overlaps the 'No' label")
            self.assertLess(
                abs(x_span["origin"][1] - self.CHECKBOX_BASELINE_Y),
                1.5,
                "X baseline does not match the line's baseline",
            )
            white_fills = [
                drawing["rect"] for drawing in doc[0].get_drawings() if drawing.get("fill") == (1.0, 1.0, 1.0)
            ]
            self.assertTrue(white_fills, "cover_anchor white-out was not drawn")
            self.assertGreaterEqual(
                min(rect.x0 for rect in white_fills),
                no_blank.x0 - 2.0,
                "white-out extends onto the 'No' label",
            )
        finally:
            doc.close()


@unittest.skipUnless(FW9_PDF_PATH.exists(), f"fixture not found: {FW9_PDF_PATH}")
class FormFillNonTextEndToEndTests(unittest.TestCase):
    """extract metadata -> mapping prompt -> (mocked) model -> write: the chosen
    checkbox export value must actually check the box on fw9."""

    def setUp(self) -> None:
        self.service = FormFillService()

    def test_mapped_checkbox_export_value_checks_the_box(self) -> None:
        field_metadata = {m["name"]: m for m in self.service._extract_pdf_form_field_metadata(str(FW9_PDF_PATH))}
        self.service.client = SimpleNamespace(
            models=SimpleNamespace(generate_content=MagicMock(return_value=_mapping_response([{"name": "c1_1[2]", "value": "3"}])))
        )
        payload = self.service._generate_mapping_payload(
            [],
            source_text="The entity is an S corporation.",
            mapping_items=["c1_1[2]"],
            mapping_label="Fillable PDF field names",
            target_hint="fillable PDF form",
            label="m",
            field_metadata=field_metadata,
        )
        field_values = {str(item["name"]): str(item["value"]) for item in payload["items"]}
        out = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        out.close()
        self.addCleanup(os.unlink, out.name)
        self.service._apply_fillable_pdf(str(FW9_PDF_PATH), field_values, out.name)
        self.assertEqual(str(_widget_by_name(out.name, "c1_1[2]").get("/AS")), "/3")
        self.assertEqual(str(_widget_by_name(out.name, "c1_1[0]").get("/AS")), "/Off")


class FormFillServiceDownloadMetadataTests(unittest.TestCase):
    def setUp(self) -> None:
        self.service = FormFillService()

    def _session_returning(self, query: MagicMock) -> MagicMock:
        db = MagicMock()
        db.query.return_value = query
        self.service._get_session = MagicMock(return_value=db)
        return db

    def test_get_source_file_metadata_returns_owned_file(self) -> None:
        source_file = SimpleNamespace(
            id=uuid.uuid4(),
            gcs_object_name="form-fill/u/runs/r/sources/1.pdf",
            original_filename="source.pdf",
            file_type="application/pdf",
        )
        query = MagicMock()
        query.join.return_value.filter.return_value.first.return_value = source_file
        db = self._session_returning(query)

        result = self.service.get_source_file_metadata("user-id", str(uuid.uuid4()), str(source_file.id))

        self.assertIs(result, source_file)
        db.expunge.assert_called_once_with(source_file)
        db.close.assert_called_once()

    def test_get_source_file_metadata_raises_when_missing(self) -> None:
        query = MagicMock()
        query.join.return_value.filter.return_value.first.return_value = None
        db = self._session_returning(query)

        with self.assertRaises(ValueError):
            self.service.get_source_file_metadata("user-id", str(uuid.uuid4()), str(uuid.uuid4()))
        db.close.assert_called_once()

    def test_get_source_file_metadata_raises_when_gcs_object_missing(self) -> None:
        source_file = SimpleNamespace(
            id=uuid.uuid4(),
            gcs_object_name=None,
            original_filename="source.pdf",
            file_type="application/pdf",
        )
        query = MagicMock()
        query.join.return_value.filter.return_value.first.return_value = source_file
        self._session_returning(query)

        with self.assertRaises(ValueError):
            self.service.get_source_file_metadata("user-id", str(uuid.uuid4()), str(source_file.id))

    def test_get_target_metadata_returns_owned_run(self) -> None:
        run = SimpleNamespace(
            id=uuid.uuid4(),
            target_gcs_object_name="form-fill/u/runs/r/target.pdf",
            target_filename="target.pdf",
            target_file_type="application/pdf",
        )
        query = MagicMock()
        query.filter.return_value.first.return_value = run
        db = self._session_returning(query)

        result = self.service.get_target_metadata("user-id", str(run.id))

        self.assertIs(result, run)
        db.expunge.assert_called_once_with(run)
        db.close.assert_called_once()

    def test_get_target_metadata_raises_when_run_missing(self) -> None:
        query = MagicMock()
        query.filter.return_value.first.return_value = None
        db = self._session_returning(query)

        with self.assertRaises(ValueError):
            self.service.get_target_metadata("user-id", str(uuid.uuid4()))
        db.close.assert_called_once()

    def test_get_target_metadata_raises_when_target_object_missing(self) -> None:
        run = SimpleNamespace(
            id=uuid.uuid4(),
            target_gcs_object_name=None,
            target_filename="target.pdf",
            target_file_type="application/pdf",
        )
        query = MagicMock()
        query.filter.return_value.first.return_value = run
        self._session_returning(query)

        with self.assertRaises(ValueError):
            self.service.get_target_metadata("user-id", str(run.id))


_W9_EXAMPLE_PATH = FW9_PDF_PATH
_FORM_FILL_EXAMPLE_DIR = FW9_PDF_PATH.parent.parent
_CONTRACTORS_EXAMPLE_PATH = _FORM_FILL_EXAMPLE_DIR / "sources" / "contractors.csv"
_W9_REQUEST_LETTER_PATH = _FORM_FILL_EXAMPLE_DIR / "targets" / "w9-request-letter.docx"


@unittest.skipUnless(_HAS_FITZ, "PyMuPDF is required for PDF field metadata tests")
class FormFillFieldLabelHintTests(unittest.TestCase):
    """Label-hint extraction against the bundled W-9, whose fields have no /TU
    tooltips and whose dense layout previously produced swapped/garbled hints
    (the LLC classification box labelled 'Exempt payee code', etc.)."""

    def setUp(self) -> None:
        self.service = FormFillService()

    @unittest.skipUnless(_W9_EXAMPLE_PATH.exists(), "bundled example fw9.pdf not found")
    def test_w9_label_hints_disambiguate_dense_form_fields(self) -> None:
        metadata = self.service._extract_pdf_form_field_metadata(str(_W9_EXAMPLE_PATH))
        by_name = {meta["name"]: meta for meta in metadata}

        def label(name: str) -> str:
            return str(by_name[name].get("label") or "").lower()

        # Line 3a: f1_03 is the LLC tax-classification code box; f1_05 is the
        # exempt payee code. These two used to receive each other's labels.
        self.assertIn("llc", label("f1_03[0]"))
        self.assertNotIn("exempt", label("f1_03[0]"))
        self.assertIn("exempt payee", label("f1_05[0]"))

        self.assertIn("address", label("f1_07[0]"))
        self.assertIn("city", label("f1_08[0]"))
        self.assertIn("requester", label("f1_09[0]"))
        self.assertNotIn("apt. or suite", label("f1_09[0]"))
        self.assertIn("account number", label("f1_10[0]"))
        self.assertIn("social security", label("f1_11[0]"))
        self.assertIn("social security", label("f1_12[0]"))
        self.assertIn("employer identification", label("f1_14[0]"))
        self.assertIn("employer identification", label("f1_15[0]"))
        self.assertIn("individual/sole proprietor", label("c1_1[0]"))
        self.assertIn("trust/estate", label("c1_1[4]"))

        line_one = label("f1_01[0]")
        self.assertIn("name of entity/individual", line_one)
        self.assertIn("business/disregarded", line_one)
        self.assertIn("line 2", line_one)

    @unittest.skipUnless(_W9_EXAMPLE_PATH.exists(), "bundled example fw9.pdf not found")
    def test_w9_field_metadata_includes_widget_geometry(self) -> None:
        metadata = self.service._extract_pdf_form_field_metadata(str(_W9_EXAMPLE_PATH))
        by_name = {meta["name"]: meta for meta in metadata}
        for meta in metadata:
            rect = meta.get("rect")
            self.assertIsInstance(rect, list, meta["name"])
            self.assertEqual(len(rect), 4, meta["name"])
        # Line 7 (account numbers) is full width; the SSN digit-group boxes are
        # narrow. Geometry is what lets the model tell them apart.
        account_width = by_name["f1_10[0]"]["rect"][2] - by_name["f1_10[0]"]["rect"][0]
        ssn_width = by_name["f1_11[0]"]["rect"][2] - by_name["f1_11[0]"]["rect"][0]
        self.assertGreater(account_width, 400)
        self.assertLess(ssn_width, 100)

    def test_field_tooltip_takes_precedence_over_geometric_hint(self) -> None:
        widget = SimpleNamespace(
            field_name="topmostSubform[0].f1_07[0]",
            field_type=fitz.PDF_WIDGET_TYPE_TEXT,
            rect=fitz.Rect(0, 0, 100, 12),
            field_label="Line 5. Address",
        )
        fields: dict = {}
        order: list = []
        with patch.object(self.service, "_widget_label_hint") as hint:
            self.service._merge_widget_metadata(
                widget=widget,
                page=None,
                page_index=1,
                row_widgets=[widget],
                label_lines=[],
                type_map={fitz.PDF_WIDGET_TYPE_TEXT: "text"},
                fields=fields,
                order=order,
            )
        hint.assert_not_called()
        self.assertEqual(fields["f1_07[0]"]["label"], "Line 5. Address")
        self.assertEqual(fields["f1_07[0]"]["rect"], [0.0, 0.0, 100.0, 12.0])

    def test_render_field_descriptor_includes_text_field_label_and_rect(self) -> None:
        descriptor = self.service._render_field_descriptor(
            "f1_08[0]",
            {"name": "f1_08[0]", "type": "text", "page": 1, "label": "City, state, and ZIP code", "rect": [58.6, 310.0, 388.0, 324.0]},
        )
        self.assertIn("f1_08[0]", descriptor)
        self.assertIn('label: "City, state, and ZIP code"', descriptor)
        self.assertIn("rect [58.6, 310.0, 388.0, 324.0]", descriptor)


class FormFillExampleFixtureTests(unittest.TestCase):
    def setUp(self) -> None:
        self.service = FormFillService()

    def test_contractor_names_distinguish_legal_entities_from_dbas(self) -> None:
        records = self.service._load_csv_records(str(_CONTRACTORS_EXAMPLE_PATH))
        names = {
            record["record_payload"]["Name"]: (
                record["record_payload"]["Legal Entity Name"],
                record["record_payload"]["DBA/Disregarded Entity Name"],
            )
            for record in records
        }

        self.assertEqual(names["Alice Monroe"], ("", "Monroe Design Studio"))
        self.assertEqual(names["Daniel Okafor"], ("Okafor Drafting LLC", ""))
        self.assertEqual(names["Priya Raman"], ("", ""))
        self.assertEqual(names["Marcus Bell"], ("Bell & Sons Electrical", ""))
        self.assertEqual(names["Sofia Herrera"], ("Herrera Translations Inc.", ""))

    def test_request_letter_replaces_only_matching_contractor_columns(self) -> None:
        records = self.service._load_csv_records(str(_CONTRACTORS_EXAMPLE_PATH))
        placeholders = self.service._extract_docx_placeholders(str(_W9_REQUEST_LETTER_PATH))
        replacements = self.service._record_placeholder_replacements(placeholders, records[0]["record_payload"])

        self.assertEqual(set(placeholders) - set(replacements), {"{{Business Name}}"})


class FormFillDroppedValueWarningTests(unittest.TestCase):
    def setUp(self) -> None:
        self.service = FormFillService()

    def test_flags_digit_bearing_value_missing_from_all_fields(self) -> None:
        record = {
            "Name": "Alice Monroe",
            "Address": "742 Juniper Lane",
            "TIN": "900-55-0101",
            "Federal Tax Classification": "Individual/sole proprietor",
        }
        field_values = {
            "f1_01[0]": "Alice Monroe",
            "f1_11[0]": "900",
            "f1_12[0]": "55",
            "f1_13[0]": "0101",
        }
        warnings = self.service._dropped_source_value_warnings(record, field_values)
        self.assertEqual(len(warnings), 1)
        self.assertIn("'Address'", warnings[0])

    def test_accepts_values_split_across_fields_and_ignores_text_only_columns(self) -> None:
        record = {
            "Name": "Sofia Herrera",
            "City, State, ZIP": "Portside, OR 97205",
            "TIN": "98-7000105",
            "TIN Type": "EIN",
        }
        field_values = {
            "f1_01[0]": "Sofia Herrera",
            "f1_08[0]": "Portside, OR 97205",
            "f1_14[0]": "98",
            "f1_15[0]": "7000105",
        }
        self.assertEqual(self.service._dropped_source_value_warnings(record, field_values), [])


def _build_text_pdf(path: str, *, pages: int = 1) -> None:
    """A flat digital PDF whose every page carries plenty of extractable text."""
    doc = fitz.open()
    for index in range(pages):
        page = doc.new_page()
        page.insert_text((72, 100), f"Page {index + 1}: taxpayer name, address, and identification details.", fontsize=11)
    doc.save(path)
    doc.close()


def _build_textless_pdf(path: str, *, pages: int = 1) -> None:
    """A PDF with no text layer, standing in for a scanned document."""
    doc = fitz.open()
    for _ in range(pages):
        doc.new_page()
    doc.save(path)
    doc.close()


@unittest.skipUnless(_HAS_FITZ, "PyMuPDF not installed")
class FormFillTargetOcrTests(unittest.IsolatedAsyncioTestCase):
    """Scanned PDF targets must be OCR-normalized before the overlay fill runs."""

    def setUp(self) -> None:
        self.service = FormFillService()
        self.temp_dir = tempfile.mkdtemp(prefix="form_fill_ocr_test_")
        self.addCleanup(shutil.rmtree, self.temp_dir, True)

    def _path(self, name: str) -> str:
        return os.path.join(self.temp_dir, name)

    def _run(self, *, target_file_type: str = "application/pdf") -> SimpleNamespace:
        return SimpleNamespace(
            id="run-1",
            user_id="user-1",
            target_file_type=target_file_type,
            target_gcs_object_name="form-fill/user-1/runs/run-1/target.pdf",
        )

    def test_needs_ocr_false_for_text_pdf(self) -> None:
        path = self._path("text.pdf")
        _build_text_pdf(path, pages=2)
        self.assertFalse(self.service._target_pdf_needs_ocr(path))

    def test_needs_ocr_true_when_any_page_lacks_text(self) -> None:
        mixed = self._path("mixed.pdf")
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 100), "This first page is digital and has plenty of extractable text on it.", fontsize=11)
        doc.new_page()
        doc.save(mixed)
        doc.close()
        self.assertTrue(self.service._target_pdf_needs_ocr(mixed))

    def test_needs_ocr_false_for_fillable_pdf(self) -> None:
        path = self._path("blank.pdf")
        _build_textless_pdf(path)
        with patch.object(self.service, "_pdf_has_form_fields", return_value=True):
            self.assertFalse(self.service._target_pdf_needs_ocr(path))

    async def test_ensure_ocr_target_returns_none_for_docx_target(self) -> None:
        run = self._run(target_file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        result = await self.service._ensure_ocr_target(run, self._path("target.docx"), self.temp_dir)
        self.assertIsNone(result)

    async def test_ensure_ocr_target_returns_none_for_text_pdf(self) -> None:
        path = self._path("target.pdf")
        _build_text_pdf(path)
        result = await self.service._ensure_ocr_target(self._run(), path, self.temp_dir)
        self.assertIsNone(result)

    async def test_ensure_ocr_target_returns_none_when_disabled(self) -> None:
        path = self._path("target.pdf")
        _build_textless_pdf(path)
        self.service.target_ocr_enabled = False
        result = await self.service._ensure_ocr_target(self._run(), path, self.temp_dir)
        self.assertIsNone(result)

    async def test_ensure_ocr_target_runs_ocr_uploads_and_returns_paths(self) -> None:
        path = self._path("target.pdf")
        _build_textless_pdf(path)

        def fake_run_ocr(*, input_pdf_path: str, output_pdf_path: str, **_kwargs) -> str:
            _build_text_pdf(output_pdf_path)
            return output_pdf_path

        upload = AsyncMock()
        with patch.object(self.service, "_download_to_local", AsyncMock(side_effect=Exception("not found"))), \
                patch.object(self.service.target_ocr_service, "run_ocr", side_effect=fake_run_ocr) as run_ocr, \
                patch.object(self.service.storage_service, "upload_file", upload):
            result = await self.service._ensure_ocr_target(self._run(), path, self.temp_dir)

        self.assertIsNotNone(result)
        ocr_local_path, ocr_object_name = result
        self.assertEqual(ocr_object_name, "form-fill/user-1/runs/run-1/target-ocr.pdf")
        self.assertTrue(os.path.exists(ocr_local_path))
        run_ocr.assert_called_once()
        upload.assert_awaited_once_with(ocr_local_path, ocr_object_name)

    async def test_ensure_ocr_target_reuses_cached_copy(self) -> None:
        path = self._path("target.pdf")
        _build_textless_pdf(path)

        async def fake_download(object_name: str, local_path: str) -> None:
            _build_text_pdf(local_path)

        with patch.object(self.service, "_download_to_local", AsyncMock(side_effect=fake_download)), \
                patch.object(self.service.target_ocr_service, "run_ocr") as run_ocr:
            result = await self.service._ensure_ocr_target(self._run(), path, self.temp_dir)

        self.assertIsNotNone(result)
        run_ocr.assert_not_called()

    async def test_ensure_ocr_target_raises_when_ocr_recovers_no_text(self) -> None:
        path = self._path("target.pdf")
        _build_textless_pdf(path)

        def fake_run_ocr(*, input_pdf_path: str, output_pdf_path: str, **_kwargs) -> str:
            _build_textless_pdf(output_pdf_path)
            return output_pdf_path

        with patch.object(self.service, "_download_to_local", AsyncMock(side_effect=Exception("not found"))), \
                patch.object(self.service.target_ocr_service, "run_ocr", side_effect=fake_run_ocr):
            with self.assertRaises(ValueError):
                await self.service._ensure_ocr_target(self._run(), path, self.temp_dir)


if __name__ == "__main__":
    unittest.main()
