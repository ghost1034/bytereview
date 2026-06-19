from __future__ import annotations

import io
import unittest

import docx

from inkwise.services.exporter import render_docx, render_pdf


def _rich_document() -> dict:
    return {
        "type": "doc",
        "content": [
            {"type": "heading", "attrs": {"level": 1}, "content": [{"type": "text", "text": "Section & Title"}]},
            {
                "type": "paragraph",
                "content": [
                    {"type": "text", "text": "Plain "},
                    {"type": "text", "text": "bold", "marks": [{"type": "bold"}]},
                    {"type": "text", "text": " "},
                    {"type": "text", "text": "italic", "marks": [{"type": "italic"}]},
                    {"type": "text", "text": " "},
                    {"type": "text", "text": "struck", "marks": [{"type": "strike"}]},
                    {"type": "text", "text": " "},
                    {"type": "text", "text": "snippet", "marks": [{"type": "code"}]},
                    {"type": "inkwiseNoteRef", "attrs": {"noteNumber": 1, "noteKind": "footnote"}},
                ],
            },
            {"type": "blockquote", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "A < B & C quote"}]}]},
            {"type": "codeBlock", "content": [{"type": "text", "text": "def f():\n    return 1 < 2"}]},
            {
                "type": "bulletList",
                "content": [
                    {
                        "type": "listItem",
                        "content": [
                            {"type": "paragraph", "content": [{"type": "text", "text": "Top bullet"}]},
                            {
                                "type": "bulletList",
                                "content": [{"type": "listItem", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "Nested bullet"}]}]}],
                            },
                        ],
                    }
                ],
            },
            {"type": "orderedList", "attrs": {"start": 1}, "content": [{"type": "listItem", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "First numbered"}]}]}]},
            {
                "type": "table",
                "content": [
                    {
                        "type": "tableRow",
                        "content": [
                            {"type": "tableHeader", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "Header A"}]}]},
                            {"type": "tableHeader", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "Header B"}]}]},
                        ],
                    },
                    {"type": "tableRow", "content": [{"type": "tableCell", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "cell"}]}]}]},
                ],
            },
            {"type": "horizontalRule"},
            {"type": "inkwisePageBreak"},
            {"type": "paragraph", "content": [{"type": "text", "text": "Second page body"}]},
        ],
    }


def _citation_document() -> dict:
    return {
        "type": "doc",
        "content": [
            {
                "type": "paragraph",
                "content": [
                    {"type": "text", "text": "Supported sentence"},
                    {
                        "type": "inkwiseInlineCitation",
                        "attrs": {"citationStyle": "default", "label": "(Smith p.3)", "citations": []},
                    },
                    {"type": "inkwiseNoteRef", "attrs": {"noteNumber": 1, "noteKind": "footnote"}},
                ],
            },
            {
                "type": "inkwiseNoteDefinition",
                "attrs": {"noteKind": "footnote", "noteNumber": 1, "citationStyle": "none"},
                "content": [{"type": "text", "text": "Explanatory footnote."}],
            },
        ],
    }


class RenderDocxTests(unittest.TestCase):
    def setUp(self) -> None:
        self.doc = docx.Document(io.BytesIO(render_docx(title="My Title", content_html=None, content_json=_rich_document())))

    def _styles(self) -> list[str]:
        return [p.style.name for p in self.doc.paragraphs]

    def test_title_and_heading_present(self) -> None:
        styles = self._styles()
        self.assertEqual(styles[0], "Title")
        self.assertEqual(self.doc.paragraphs[0].text, "My Title")
        self.assertIn("Heading 1", styles)

    def test_inline_marks_render_as_runs(self) -> None:
        runs = [run for paragraph in self.doc.paragraphs for run in paragraph.runs]
        self.assertTrue(any(run.text == "bold" and run.bold for run in runs))
        self.assertTrue(any(run.text == "italic" and run.italic for run in runs))
        self.assertTrue(any(run.text == "struck" and run.font.strike for run in runs))
        self.assertTrue(any(run.text == "snippet" and run.font.name == "Courier New" for run in runs))

    def test_lists_use_list_styles(self) -> None:
        styles = self._styles()
        self.assertIn("List Bullet", styles)
        self.assertIn("List Bullet 2", styles)
        self.assertIn("List Number", styles)

    def test_blockquote_and_code_block(self) -> None:
        styles = self._styles()
        self.assertIn("Quote", styles)
        code_runs = [run for paragraph in self.doc.paragraphs for run in paragraph.runs if run.font.name == "Courier New"]
        self.assertTrue(any("return 1 < 2" in run.text for run in code_runs))

    def test_table_rendered_with_bold_header(self) -> None:
        self.assertEqual(len(self.doc.tables), 1)
        table = self.doc.tables[0]
        self.assertEqual(len(table.columns), 2)  # ragged row padded to 2 columns
        header_runs = [run for paragraph in table.rows[0].cells[0].paragraphs for run in paragraph.runs]
        self.assertTrue(header_runs and all(run.bold for run in header_runs))

    def test_page_break_present(self) -> None:
        self.assertIn('w:type="page"', self.doc.element.xml)

    def test_horizontal_rule_border_present(self) -> None:
        self.assertIn("w:pBdr", self.doc.element.xml)


class RenderPdfTests(unittest.TestCase):
    def test_pdf_is_valid_and_nonempty(self) -> None:
        data = render_pdf(title="My Title", content_html=None, content_json=_rich_document())
        self.assertTrue(data.startswith(b"%PDF"))
        self.assertGreater(len(data), 1000)

    def test_pdf_escapes_special_characters_without_raising(self) -> None:
        document = {"type": "doc", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "A & B < C > D"}]}]}
        data = render_pdf(title="Title & <Tag>", content_html=None, content_json=document)
        self.assertTrue(data.startswith(b"%PDF"))

    def test_pdf_text_contains_expected_content(self) -> None:
        try:
            from pypdf import PdfReader
        except ImportError:
            try:
                from PyPDF2 import PdfReader
            except ImportError:
                self.skipTest("no PDF text extraction library available")
        data = render_pdf(title="My Title", content_html=None, content_json=_citation_document())
        reader = PdfReader(io.BytesIO(data))
        text = "".join(page.extract_text() or "" for page in reader.pages)
        self.assertIn("Supported sentence", text)
        self.assertIn("(Smith p.3)", text)
        self.assertIn("Explanatory footnote.", text)


class FallbackTests(unittest.TestCase):
    def test_html_only_fallback_produces_output(self) -> None:
        html = "<h1>Title</h1><p>First paragraph.</p><p>Second paragraph.</p>"
        docx_bytes = render_docx(title="Doc", content_html=html, content_json=None)
        pdf_bytes = render_pdf(title="Doc", content_html=html, content_json=None)
        self.assertGreater(len(docx_bytes), 1000)
        self.assertTrue(pdf_bytes.startswith(b"%PDF"))
        parsed = docx.Document(io.BytesIO(docx_bytes))
        body_text = "\n".join(p.text for p in parsed.paragraphs)
        self.assertIn("First paragraph.", body_text)

    def test_malformed_json_does_not_raise(self) -> None:
        docx_bytes = render_docx(title="Doc", content_html=None, content_json={"garbage": True})
        pdf_bytes = render_pdf(title="Doc", content_html=None, content_json={"garbage": True})
        self.assertGreater(len(docx_bytes), 0)
        self.assertTrue(pdf_bytes.startswith(b"%PDF"))


class CitationReuseTests(unittest.TestCase):
    def test_docx_preserves_inline_citation_and_note(self) -> None:
        parsed = docx.Document(io.BytesIO(render_docx(title="Doc", content_html=None, content_json=_citation_document())))
        body_text = "\n".join(p.text for p in parsed.paragraphs)
        self.assertIn("(Smith p.3)", body_text)
        self.assertIn("[^1]", body_text)
        self.assertIn("Explanatory footnote.", body_text)


if __name__ == "__main__":
    unittest.main()
