"""Document export helpers for the Inkwise module.

Inkwise documents are authored in TipTap (ProseMirror) and stored as TipTap JSON.
To preserve formatting on export, the JSON is walked once into a small,
backend-neutral intermediate representation (IR) and then rendered to DOCX
(python-docx) and PDF (ReportLab Platypus). A plain-text fallback keeps exports
working when the structured content is missing or malformed.
"""

from __future__ import annotations

import html
import io
import json
import re
from dataclasses import dataclass, field
from typing import Any
from xml.sax.saxutils import escape as xml_escape

from inkwise.services.citation_styles import citation_style_requires_reference_text, format_inline_citation, format_note_citation, normalize_citation_style


class ExportError(RuntimeError):
    pass


_TAG_RE = re.compile(r"<[^>]+>")
_BR_RE = re.compile(r"<\s*br\s*/?\s*>", re.IGNORECASE)
_P_END_RE = re.compile(r"</\s*p\s*>", re.IGNORECASE)
_PAGE_BREAK_SENTINEL = "\f"


def html_to_text(content_html: str | None) -> str:
    if not content_html:
        return ""

    text = _BR_RE.sub("\n", content_html)
    text = _P_END_RE.sub("\n\n", text)
    text = _TAG_RE.sub("", text)
    text = html.unescape(text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _normalize_block_text(value: str) -> str:
    value = (value or "").replace("\r\n", "\n").replace("\r", "\n")
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def _content_json_to_text(content_json: Any) -> str:
    def render_node(node: Any) -> str:
        if isinstance(node, str):
            return node
        if not isinstance(node, dict):
            return ""

        node_type = node.get("type")
        content = node.get("content") or []

        if node_type == "text":
            return str(node.get("text") or "")
        if node_type == "inkwiseCitationAnchor":
            return ""
        if node_type == "inkwiseInlineCitation":
            attrs = node.get("attrs") or {}
            citations = attrs.get("citations") if isinstance(attrs.get("citations"), list) else []
            citation_style = normalize_citation_style(attrs.get("citationStyle"))
            label = str(attrs.get("label") or "").strip()
            if not label and citation_style_requires_reference_text(citation_style):
                label = format_inline_citation(citations, citation_style).strip()
            return label
        if node_type == "inkwiseNoteRef":
            attrs = node.get("attrs") or {}
            note_number = attrs.get("noteNumber", "")
            note_kind = attrs.get("noteKind", "footnote")
            if note_kind == "footnote":
                return f"[^{note_number}]"
            return f"[{note_number}]"
        if node_type == "inkwiseNoteDefinition":
            attrs = node.get("attrs") or {}
            note_number = attrs.get("noteNumber", "")
            note_kind = attrs.get("noteKind", "footnote")
            citations = attrs.get("citations") if isinstance(attrs.get("citations"), list) else []
            citation_style = normalize_citation_style(attrs.get("citationStyle"))
            if citations and citation_style_requires_reference_text(citation_style):
                inner_text = format_note_citation(citations, citation_style)
            else:
                inner_text = "".join(render_node(child) for child in content)
            inner_text = _normalize_block_text(inner_text)
            if note_kind == "footnote":
                return f"[^{note_number}]: {inner_text}"
            return f"{note_number}. {inner_text}"
        if node_type == "inkwisePageBreak":
            return _PAGE_BREAK_SENTINEL

        if node_type in {"doc", "paragraph", "heading", "blockquote", "tableCell", "tableHeader"}:
            rendered = "".join(render_node(child) for child in content)
            return _normalize_block_text(rendered)

        if node_type == "listItem":
            rendered = " ".join(part for part in (render_node(child) for child in content) if part)
            return _normalize_block_text(rendered)

        if node_type == "bulletList":
            return "\n".join(f"- {item}" for item in (render_node(child) for child in content) if item)

        if node_type == "orderedList":
            items = [item for item in (render_node(child) for child in content) if item]
            return "\n".join(f"{index}. {item}" for index, item in enumerate(items, start=1))

        if node_type == "tableRow":
            cells = [render_node(child) for child in content]
            return "\t".join(cell for cell in cells if cell is not None)

        if node_type == "table":
            rows = [render_node(child) for child in content]
            return "\n".join(row for row in rows if row)

        rendered_children = [render_node(child) for child in content]
        return _normalize_block_text("\n".join(part for part in rendered_children if part))

    if not isinstance(content_json, dict):
        return ""

    text = render_node(content_json)
    text = text.replace(f"{_PAGE_BREAK_SENTINEL}\n", _PAGE_BREAK_SENTINEL)
    text = text.replace(f"\n{_PAGE_BREAK_SENTINEL}", _PAGE_BREAK_SENTINEL)
    return _normalize_block_text(text.replace(_PAGE_BREAK_SENTINEL, f"\n{_PAGE_BREAK_SENTINEL}\n"))


def content_to_text(*, content_html: str | None, content_json: Any = None) -> str:
    if isinstance(content_json, dict):
        structured_text = _content_json_to_text(content_json)
        if structured_text:
            return structured_text

    text = html_to_text(content_html)
    if text:
        return text
    if isinstance(content_json, str):
        return content_json.strip()
    if isinstance(content_json, dict):
        try:
            return json.dumps(content_json, ensure_ascii=True, indent=2)
        except Exception:
            return str(content_json)
    return ""


# ---------------------------------------------------------------------------
# Intermediate representation
# ---------------------------------------------------------------------------


@dataclass
class InlineRun:
    text: str
    bold: bool = False
    italic: bool = False
    strike: bool = False
    code: bool = False


@dataclass
class InlineBreak:
    pass


@dataclass
class HeadingBlock:
    level: int
    inlines: list = field(default_factory=list)


@dataclass
class ParagraphBlock:
    inlines: list = field(default_factory=list)


@dataclass
class BlockquoteBlock:
    children: list = field(default_factory=list)


@dataclass
class CodeBlock:
    text: str = ""


@dataclass
class ListBlock:
    ordered: bool = False
    items: list = field(default_factory=list)
    start: int = 1


@dataclass
class ListItemBlock:
    blocks: list = field(default_factory=list)


@dataclass
class TableCellBlock:
    header: bool = False
    blocks: list = field(default_factory=list)


@dataclass
class TableRowBlock:
    cells: list = field(default_factory=list)


@dataclass
class TableBlock:
    rows: list = field(default_factory=list)


@dataclass
class HorizontalRuleBlock:
    pass


@dataclass
class PageBreakBlock:
    pass


@dataclass
class NoteDefinitionBlock:
    inlines: list = field(default_factory=list)


# ---------------------------------------------------------------------------
# TipTap JSON -> IR
# ---------------------------------------------------------------------------

_MARK_FLAGS = ("bold", "italic", "strike", "code")


def _marks_to_flags(marks: Any) -> dict:
    flags = {name: False for name in _MARK_FLAGS}
    if not isinstance(marks, list):
        return flags
    for mark in marks:
        if isinstance(mark, dict) and mark.get("type") in flags:
            flags[mark["type"]] = True
    return flags


def _inline_citation_label(node: dict) -> str:
    attrs = node.get("attrs") or {}
    citations = attrs.get("citations") if isinstance(attrs.get("citations"), list) else []
    citation_style = normalize_citation_style(attrs.get("citationStyle"))
    label = str(attrs.get("label") or "").strip()
    if not label and citation_style_requires_reference_text(citation_style):
        label = format_inline_citation(citations, citation_style).strip()
    return label


def _note_ref_label(node: dict) -> str:
    attrs = node.get("attrs") or {}
    note_number = attrs.get("noteNumber", "")
    note_kind = attrs.get("noteKind", "footnote")
    return f"[^{note_number}]" if note_kind == "footnote" else f"[{note_number}]"


def _inlines(children: Any) -> list:
    runs: list = []
    for node in children or []:
        if isinstance(node, str):
            if node:
                runs.append(InlineRun(text=node))
            continue
        if not isinstance(node, dict):
            continue
        node_type = node.get("type")
        if node_type == "text":
            text = str(node.get("text") or "")
            if text:
                runs.append(InlineRun(text=text, **_marks_to_flags(node.get("marks"))))
        elif node_type == "hardBreak":
            runs.append(InlineBreak())
        elif node_type == "inkwiseCitationAnchor":
            continue
        elif node_type == "inkwiseInlineCitation":
            label = _inline_citation_label(node)
            if label:
                runs.append(InlineRun(text=label))
        elif node_type == "inkwiseNoteRef":
            runs.append(InlineRun(text=_note_ref_label(node)))
        elif node.get("content"):
            runs.extend(_inlines(node.get("content")))
    return runs


def _note_definition_inlines(node: dict) -> list:
    attrs = node.get("attrs") or {}
    note_number = attrs.get("noteNumber", "")
    note_kind = attrs.get("noteKind", "footnote")
    citations = attrs.get("citations") if isinstance(attrs.get("citations"), list) else []
    citation_style = normalize_citation_style(attrs.get("citationStyle"))
    if citations and citation_style_requires_reference_text(citation_style):
        inner_text = format_note_citation(citations, citation_style)
        inner_runs = [InlineRun(text=inner_text)] if inner_text else []
    else:
        inner_runs = _inlines(node.get("content"))
    prefix = f"[^{note_number}]: " if note_kind == "footnote" else f"{note_number}. "
    return [InlineRun(text=prefix), *inner_runs]


def _raw_text(content: Any) -> str:
    parts: list = []
    for child in content or []:
        if isinstance(child, str):
            parts.append(child)
        elif isinstance(child, dict):
            if child.get("type") == "text":
                parts.append(str(child.get("text") or ""))
            elif child.get("type") == "hardBreak":
                parts.append("\n")
            elif child.get("content"):
                parts.append(_raw_text(child.get("content")))
    return "".join(parts)


def _list_items(content: Any) -> list:
    items: list = []
    for child in content or []:
        if isinstance(child, dict) and child.get("type") == "listItem":
            items.append(ListItemBlock(blocks=_blocks(child.get("content"))))
    return items


def _table_block(content: Any) -> TableBlock:
    rows: list = []
    for row in content or []:
        if not isinstance(row, dict) or row.get("type") != "tableRow":
            continue
        cells: list = []
        for cell in row.get("content") or []:
            if not isinstance(cell, dict) or cell.get("type") not in ("tableCell", "tableHeader"):
                continue
            cells.append(TableCellBlock(header=cell.get("type") == "tableHeader", blocks=_blocks(cell.get("content"))))
        rows.append(TableRowBlock(cells=cells))

    max_cols = max((len(row.cells) for row in rows), default=0)
    for row in rows:
        while len(row.cells) < max_cols:
            row.cells.append(TableCellBlock())
    return TableBlock(rows=rows)


def _node_to_blocks(node: Any) -> list:
    if not isinstance(node, dict):
        return []
    node_type = node.get("type")
    content = node.get("content") or []
    attrs = node.get("attrs") if isinstance(node.get("attrs"), dict) else {}

    if node_type == "paragraph":
        return [ParagraphBlock(inlines=_inlines(content))]
    if node_type == "heading":
        level = _coerce_int(attrs.get("level"), default=1)
        level = 1 if level < 1 else 6 if level > 6 else level
        return [HeadingBlock(level=level, inlines=_inlines(content))]
    if node_type == "blockquote":
        return [BlockquoteBlock(children=_blocks(content))]
    if node_type == "codeBlock":
        return [CodeBlock(text=_raw_text(content))]
    if node_type == "bulletList":
        return [ListBlock(ordered=False, items=_list_items(content))]
    if node_type == "orderedList":
        return [ListBlock(ordered=True, items=_list_items(content), start=_coerce_int(attrs.get("start"), default=1))]
    if node_type == "horizontalRule":
        return [HorizontalRuleBlock()]
    if node_type == "inkwisePageBreak":
        return [PageBreakBlock()]
    if node_type == "table":
        return [_table_block(content)]
    if node_type == "inkwiseNoteDefinition":
        return [NoteDefinitionBlock(inlines=_note_definition_inlines(node))]

    # Unknown / structural wrapper: descend into children so new node types still export.
    return _blocks(content)


def _blocks(content: Any) -> list:
    result: list = []
    for child in content or []:
        result.extend(_node_to_blocks(child))
    return result


def _coerce_int(value: Any, *, default: int) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return default


def _content_json_to_ir(content_json: Any) -> list:
    if not isinstance(content_json, dict):
        return []
    if content_json.get("type") == "doc":
        return _blocks(content_json.get("content"))
    return _node_to_blocks(content_json)


def _plain_text_to_ir(text: str) -> list:
    if not text:
        return []
    blocks: list = []
    for page_index, page in enumerate(text.split(_PAGE_BREAK_SENTINEL)):
        if page_index > 0:
            blocks.append(PageBreakBlock())
        for chunk in page.split("\n\n"):
            if not chunk.strip():
                continue
            inlines: list = []
            for line_index, line in enumerate(chunk.split("\n")):
                if line_index > 0:
                    inlines.append(InlineBreak())
                if line:
                    inlines.append(InlineRun(text=line))
            blocks.append(ParagraphBlock(inlines=inlines))
    return blocks


def _build_ir(*, content_html: str | None, content_json: Any) -> list:
    if isinstance(content_json, dict):
        try:
            blocks = _content_json_to_ir(content_json)
        except Exception:
            blocks = []
        if blocks:
            return blocks
    return _plain_text_to_ir(content_to_text(content_html=content_html, content_json=content_json))


def _iter_paragraph_inlines(blocks: list):
    """Yield one inline-run list per paragraph found anywhere within the blocks.

    Used to flatten cell content (which is almost always paragraphs) for table
    rendering in both backends.
    """
    for block in blocks:
        if isinstance(block, (ParagraphBlock, HeadingBlock, NoteDefinitionBlock)):
            yield block.inlines
        elif isinstance(block, BlockquoteBlock):
            yield from _iter_paragraph_inlines(block.children)
        elif isinstance(block, ListBlock):
            for item in block.items:
                yield from _iter_paragraph_inlines(item.blocks)
        elif isinstance(block, CodeBlock):
            yield [InlineRun(text=block.text)]


# ---------------------------------------------------------------------------
# DOCX backend
# ---------------------------------------------------------------------------


def _apply_runs(paragraph: Any, inlines: list) -> None:
    for item in inlines:
        if isinstance(item, InlineBreak):
            paragraph.add_run().add_break()
            continue
        if not isinstance(item, InlineRun) or not item.text:
            continue
        run = paragraph.add_run(item.text)
        if item.bold:
            run.bold = True
        if item.italic:
            run.italic = True
        if item.strike:
            run.font.strike = True
        if item.code:
            run.font.name = "Courier New"


def _docx_has_style(doc: Any, name: str) -> bool:
    try:
        doc.styles[name]
        return True
    except KeyError:
        return False


def _safe_paragraph(doc: Any, *styles: str) -> Any:
    for style in styles:
        if style and _docx_has_style(doc, style):
            return doc.add_paragraph(style=style)
    return doc.add_paragraph()


def _add_horizontal_rule(doc: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph = doc.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    borders.append(bottom)
    p_pr.append(borders)


def _add_list(doc: Any, list_block: ListBlock, level: int = 0) -> None:
    from docx.shared import Inches

    base = "List Number" if list_block.ordered else "List Bullet"
    if level <= 0:
        style_name, manual_indent = base, False
    else:
        leveled = f"{base} {level + 1}"
        style_name, manual_indent = (leveled, False) if _docx_has_style(doc, leveled) else (base, True)

    for item in list_block.items:
        paragraph_blocks = [block for block in item.blocks if isinstance(block, ParagraphBlock)]
        nested_lists = [block for block in item.blocks if isinstance(block, ListBlock)]
        other_blocks = [block for block in item.blocks if not isinstance(block, (ParagraphBlock, ListBlock))]

        targets = paragraph_blocks or ([None] if not (nested_lists or other_blocks) else [])
        for para_block in targets:
            paragraph = _safe_paragraph(doc, style_name, base)
            if manual_indent:
                paragraph.paragraph_format.left_indent = Inches(0.25 * (level + 1))
            if para_block is not None:
                _apply_runs(paragraph, para_block.inlines)

        for block in other_blocks:
            _render_block_docx(doc, block)
        for nested in nested_lists:
            _add_list(doc, nested, level + 1)


def _add_table_docx(doc: Any, table_block: TableBlock) -> None:
    rows = table_block.rows
    max_cols = max((len(row.cells) for row in rows), default=0)
    if not rows or max_cols == 0:
        return

    table = doc.add_table(rows=len(rows), cols=max_cols)
    if _docx_has_style(doc, "Table Grid"):
        table.style = "Table Grid"

    for row_index, row in enumerate(rows):
        for col_index in range(max_cols):
            cell_block = row.cells[col_index] if col_index < len(row.cells) else None
            cell = table.cell(row_index, col_index)
            if cell_block is None:
                continue
            paragraphs = list(_iter_paragraph_inlines(cell_block.blocks))
            for index, inlines in enumerate(paragraphs):
                paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
                _apply_runs(paragraph, inlines)
                if cell_block.header:
                    for run in paragraph.runs:
                        run.bold = True


def _render_block_docx(doc: Any, block: Any) -> None:
    from docx.shared import Inches

    if isinstance(block, HeadingBlock):
        paragraph = doc.add_heading("", level=block.level)
        _apply_runs(paragraph, block.inlines)
    elif isinstance(block, ParagraphBlock):
        paragraph = doc.add_paragraph()
        _apply_runs(paragraph, block.inlines)
    elif isinstance(block, BlockquoteBlock):
        for child in block.children:
            if isinstance(child, ParagraphBlock):
                paragraph = _safe_paragraph(doc, "Quote")
                if paragraph.style is None or paragraph.style.name != "Quote":
                    paragraph.paragraph_format.left_indent = Inches(0.5)
                _apply_runs(paragraph, child.inlines)
            else:
                _render_block_docx(doc, child)
    elif isinstance(block, CodeBlock):
        paragraph = doc.add_paragraph()
        lines = block.text.split("\n")
        for line_index, line in enumerate(lines):
            run = paragraph.add_run(line)
            run.font.name = "Courier New"
            if line_index < len(lines) - 1:
                run.add_break()
    elif isinstance(block, ListBlock):
        _add_list(doc, block, 0)
    elif isinstance(block, TableBlock):
        _add_table_docx(doc, block)
    elif isinstance(block, HorizontalRuleBlock):
        _add_horizontal_rule(doc)
    elif isinstance(block, PageBreakBlock):
        doc.add_page_break()
    elif isinstance(block, NoteDefinitionBlock):
        paragraph = doc.add_paragraph()
        _apply_runs(paragraph, block.inlines)


def render_docx(*, title: str, content_html: str | None, content_json: Any = None) -> bytes:
    try:
        from docx import Document as DocxDocument
    except Exception as exc:
        raise ExportError("python-docx is not installed") from exc

    doc = DocxDocument()
    doc.add_heading(title or "Untitled", level=0)
    for block in _build_ir(content_html=content_html, content_json=content_json):
        _render_block_docx(doc, block)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF backend (ReportLab Platypus)
# ---------------------------------------------------------------------------


def _inline_markup(inlines: list) -> str:
    parts: list = []
    for item in inlines:
        if isinstance(item, InlineBreak):
            parts.append("<br/>")
            continue
        if not isinstance(item, InlineRun) or not item.text:
            continue
        text = xml_escape(item.text)
        if item.code:
            text = f'<font face="Courier">{text}</font>'
        if item.strike:
            text = f"<strike>{text}</strike>"
        if item.italic:
            text = f"<i>{text}</i>"
        if item.bold:
            text = f"<b>{text}</b>"
        parts.append(text)
    return "".join(parts)


def _cell_markup(cell_block: Any, *, header: bool) -> str:
    if cell_block is None:
        return ""
    chunks = [_inline_markup(inlines) for inlines in _iter_paragraph_inlines(cell_block.blocks)]
    markup = "<br/>".join(chunk for chunk in chunks if chunk)
    if header and markup:
        markup = f"<b>{markup}</b>"
    return markup


def _pdf_styles():
    from reportlab.lib.colors import HexColor
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="InkTitle", fontName="Helvetica-Bold", fontSize=18, leading=22, spaceAfter=14))
    styles.add(ParagraphStyle(name="InkBody", fontName="Helvetica", fontSize=11, leading=15, spaceAfter=8))
    for level, size in ((1, 17), (2, 15), (3, 13), (4, 12), (5, 11), (6, 11)):
        styles.add(ParagraphStyle(name=f"InkH{level}", fontName="Helvetica-Bold", fontSize=size, leading=size + 4, spaceBefore=10, spaceAfter=6))
    styles.add(ParagraphStyle(name="InkQuote", fontName="Helvetica-Oblique", fontSize=11, leading=15, leftIndent=18, textColor=HexColor("#555555"), spaceAfter=8))
    styles.add(ParagraphStyle(name="InkCode", fontName="Courier", fontSize=9.5, leading=12, backColor=HexColor("#f2f2f2"), borderPadding=4, spaceAfter=8))
    return styles


def _pdf_list(block: ListBlock, styles, frame_width: float):
    from reportlab.platypus import ListFlowable, ListItem, Paragraph

    items = []
    for item in block.items:
        flowables: list = []
        for child in item.blocks:
            if isinstance(child, ListBlock):
                flowables.append(_pdf_list(child, styles, frame_width))
            else:
                flowables.extend(_pdf_block(child, styles, frame_width))
        if not flowables:
            flowables = [Paragraph("&nbsp;", styles["InkBody"])]
        items.append(ListItem(flowables))

    if block.ordered:
        return ListFlowable(items, bulletType="1", start=block.start)
    return ListFlowable(items, bulletType="bullet")


def _pdf_table(block: TableBlock, styles, frame_width: float):
    from reportlab.lib import colors
    from reportlab.platypus import Paragraph, Table, TableStyle

    rows = block.rows
    max_cols = max((len(row.cells) for row in rows), default=0)
    if not rows or max_cols == 0:
        return None

    data = []
    header_rows: list = []
    for row in rows:
        rendered = []
        is_header = False
        for col_index in range(max_cols):
            cell = row.cells[col_index] if col_index < len(row.cells) else None
            if cell is not None and cell.header:
                is_header = True
            rendered.append(Paragraph(_cell_markup(cell, header=bool(cell and cell.header)) or "&nbsp;", styles["InkBody"]))
        data.append(rendered)
        header_rows.append(is_header)

    col_width = frame_width / max_cols
    table = Table(data, colWidths=[col_width] * max_cols)
    commands = [
        ("GRID", (0, 0), (-1, -1), 0.5, colors.Color(0.6, 0.6, 0.6)),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]
    for row_index, is_header in enumerate(header_rows):
        if is_header:
            commands.append(("BACKGROUND", (0, row_index), (-1, row_index), colors.Color(0.93, 0.93, 0.93)))
    table.setStyle(TableStyle(commands))
    return table


def _pdf_block(block: Any, styles, frame_width: float) -> list:
    from reportlab.lib import colors
    from reportlab.platypus import HRFlowable, PageBreak, Paragraph, Preformatted, Spacer

    if isinstance(block, HeadingBlock):
        return [Paragraph(_inline_markup(block.inlines) or "&nbsp;", styles[f"InkH{block.level}"])]
    if isinstance(block, ParagraphBlock):
        markup = _inline_markup(block.inlines)
        return [Paragraph(markup, styles["InkBody"])] if markup else [Spacer(1, 6)]
    if isinstance(block, BlockquoteBlock):
        out: list = []
        for child in block.children:
            if isinstance(child, ParagraphBlock):
                out.append(Paragraph(_inline_markup(child.inlines) or "&nbsp;", styles["InkQuote"]))
            else:
                out.extend(_pdf_block(child, styles, frame_width))
        return out
    if isinstance(block, CodeBlock):
        return [Preformatted(block.text or " ", styles["InkCode"])]
    if isinstance(block, ListBlock):
        return [_pdf_list(block, styles, frame_width)]
    if isinstance(block, TableBlock):
        table = _pdf_table(block, styles, frame_width)
        return [table] if table is not None else []
    if isinstance(block, HorizontalRuleBlock):
        return [HRFlowable(width="100%", thickness=0.5, color=colors.grey, spaceBefore=6, spaceAfter=6)]
    if isinstance(block, PageBreakBlock):
        return [PageBreak()]
    if isinstance(block, NoteDefinitionBlock):
        return [Paragraph(_inline_markup(block.inlines) or "&nbsp;", styles["InkBody"])]
    return []


def render_pdf(*, title: str, content_html: str | None, content_json: Any = None) -> bytes:
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, SimpleDocTemplate
    except Exception as exc:
        raise ExportError("reportlab is not installed") from exc

    buf = io.BytesIO()
    margin = 0.85 * inch
    doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        leftMargin=margin,
        rightMargin=margin,
        topMargin=margin,
        bottomMargin=margin,
        title=(title or "Untitled")[:120],
    )

    styles = _pdf_styles()
    story = [Paragraph(xml_escape(title or "Untitled"), styles["InkTitle"])]
    for block in _build_ir(content_html=content_html, content_json=content_json):
        story.extend(_pdf_block(block, styles, doc.width))

    doc.build(story)
    return buf.getvalue()
