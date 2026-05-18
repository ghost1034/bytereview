"""Backend service for the Form Fill feature."""

from __future__ import annotations

import ast
import csv
import hashlib
import json
import logging
import mimetypes
import multiprocessing
import os
import queue
import re
import shutil
import tempfile
import uuid
import zipfile
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Optional

import fitz
from google import genai
from google.genai import types
from openpyxl import load_workbook
from sqlalchemy import text
from sqlalchemy.orm import Session

from core.database import db_config
from models.db_models import (
    ExtractionJob,
    ExtractionResult,
    ExtractionTask,
    FormFillOutput,
    FormFillRun,
    FormFillSourceFile,
    FormFillTemplate,
    JobRun,
    SourceFile,
    SourceFileToTask,
)
from models.form_fill import (
    FormFillExtractionSourcePreviewResponse,
    FormFillRunResponse,
    FormFillTemplateResponse,
)
from services.cloud_run_task_service import cloud_run_task_service
from services.document_conversion_service import DOCX_MIME, get_document_conversion_service
from services.gcs_service import get_storage_service
from services.natural_sort import natural_text_key, sort_paths_naturally
from services.page_counting_service import page_counting_service


logger = logging.getLogger(__name__)

PDF_MIME = "application/pdf"
DOCX_PLACEHOLDER_RE = re.compile(r"(\{\{[^{}]+\}\}|\[\[[^\[\]]+\]\]|<<[^<>]+>>)")
DOCX_BLOCK_TEXT_LIMIT = 30000
SUPPORTED_SOURCE_MIME_TYPES = {
    "application/pdf",
    DOCX_MIME,
    "text/csv",
    "application/vnd.ms-excel",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}
TABULAR_SOURCE_MIME_TYPES = {
    "text/csv",
    "application/vnd.ms-excel",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}
SUPPORTED_TARGET_MIME_TYPES = {PDF_MIME, DOCX_MIME}
DEFAULT_MAX_SOURCE_FILES = 100
DEFAULT_MAX_TOTAL_SOURCE_BYTES = 1000 * 1024 * 1024
REPEAT_MODE_SINGLE = "single"
REPEAT_MODE_SOURCE_ROWS = "source_rows"
REPEAT_MODE_ALL_SOURCES = "all_sources"
REPEAT_LABEL_COLUMNS = ("participant", "participant name", "name", "full name", "client", "customer", "employee")
SOURCE_SCOPE_TASK = "task"
SOURCE_SCOPE_ALL = "all"
FORM_FILL_OUTPUT_TERMINAL_STATUSES = {"completed", "failed"}
FORM_FILL_RUN_TERMINAL_STATUSES = {"completed", "completed_with_errors", "failed"}
OUTPUT_LIMIT_WARNING_PATTERNS = (
    "output limit",
    "output limits",
    "too many",
    "large volume",
    "large number",
    "exceed",
    "first 100",
    "only the first",
    "request continuation",
    "remaining rows",
    "see attached",
    "unable to calculate exact totals",
)
GENERATED_CODE_LANGUAGE = "python"
GENERATED_CODE_REQUIRED_FUNCTION = "transform"
GENERATED_CODE_MAX_OUTPUT_BYTES_DEFAULT = 50 * 1024 * 1024
GENERATED_CODE_TIMEOUT_SECONDS_DEFAULT = 120
GENERATED_CODE_ALLOWED_BUILTINS = {
    "abs": abs,
    "all": all,
    "any": any,
    "bool": bool,
    "dict": dict,
    "enumerate": enumerate,
    "float": float,
    "int": int,
    "isinstance": isinstance,
    "len": len,
    "list": list,
    "max": max,
    "min": min,
    "range": range,
    "round": round,
    "set": set,
    "sorted": sorted,
    "str": str,
    "sum": sum,
    "tuple": tuple,
    "zip": zip,
}
GENERATED_CODE_BANNED_CALLS = {
    "__import__",
    "compile",
    "eval",
    "exec",
    "getattr",
    "globals",
    "input",
    "locals",
    "open",
    "setattr",
    "vars",
}


def _guess_mime_type(filename: str, fallback: str = "application/octet-stream") -> str:
    guessed, _ = mimetypes.guess_type(filename or "")
    return guessed or fallback


def _safe_ext(filename: str, fallback: str) -> str:
    ext = Path(filename or "").suffix.lower()
    return ext or fallback


def _normalize_output_format(value: Optional[str]) -> Optional[str]:
    if value is None:
        return None
    lowered = value.strip().lower()
    return lowered or None


def _normalize_repeat_mode(value: Optional[str]) -> str:
    lowered = (value or "").strip().lower()
    if not lowered:
        return REPEAT_MODE_ALL_SOURCES
    if lowered == REPEAT_MODE_ALL_SOURCES:
        return REPEAT_MODE_ALL_SOURCES
    if lowered in {"source_rows", "rows", "repeat"}:
        return REPEAT_MODE_SOURCE_ROWS
    return REPEAT_MODE_SINGLE


def _normalize_source_scope(value: Optional[str]) -> str:
    return SOURCE_SCOPE_ALL if (value or "").strip().lower() == SOURCE_SCOPE_ALL else SOURCE_SCOPE_TASK


def _safe_filename_part(value: str, fallback: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", (value or "").strip()).strip("._-")
    return (cleaned or fallback)[:80]


def _parse_number(value: Any) -> Optional[float]:
    if value is None:
        return None
    if isinstance(value, bool):
        return 1.0 if value else 0.0
    if isinstance(value, (int, float)):
        return float(value)
    text_value = re.sub(r"[^0-9.\-]+", "", str(value))
    if not text_value or text_value in {"-", ".", "-."}:
        return None
    try:
        return float(text_value)
    except ValueError:
        return None


def _as_text(value: Any) -> str:
    return "" if value is None else str(value)


def _validate_generated_transform_code(code: str) -> None:
    if not isinstance(code, str) or not code.strip():
        raise ValueError("Generated code was empty")
    try:
        tree = ast.parse(code)
    except SyntaxError as exc:
        raise ValueError(f"Generated code has invalid Python syntax: {exc}") from exc

    has_transform = False
    for node in ast.walk(tree):
        if isinstance(node, (ast.Import, ast.ImportFrom)):
            raise ValueError("Generated code may not import modules")
        if isinstance(node, (ast.AsyncFunctionDef, ast.ClassDef, ast.Lambda, ast.Global, ast.Nonlocal)):
            raise ValueError("Generated code uses unsupported Python constructs")
        if isinstance(node, ast.FunctionDef) and node.name == GENERATED_CODE_REQUIRED_FUNCTION:
            has_transform = True
        if isinstance(node, ast.Name):
            if node.id.startswith("__") or node.id in GENERATED_CODE_BANNED_CALLS:
                raise ValueError(f"Generated code uses blocked name '{node.id}'")
        if isinstance(node, ast.Attribute):
            if node.attr.startswith("__"):
                raise ValueError("Generated code may not access dunder attributes")
        if isinstance(node, ast.Call) and isinstance(node.func, ast.Name) and node.func.id in GENERATED_CODE_BANNED_CALLS:
            raise ValueError(f"Generated code calls blocked function '{node.func.id}'")

    if not has_transform:
        raise ValueError("Generated code must define transform(rows, context)")


def _run_generated_transform_worker(code: str, rows: list[dict[str, Any]], context: dict[str, Any], result_queue: Any) -> None:
    try:
        _validate_generated_transform_code(code)
        globals_dict = {
            "__builtins__": GENERATED_CODE_ALLOWED_BUILTINS,
            "parse_number": _parse_number,
            "as_text": _as_text,
        }
        locals_dict: dict[str, Any] = {}
        exec(compile(code, "<form-fill-generated-transform>", "exec"), globals_dict, locals_dict)
        transform = locals_dict.get(GENERATED_CODE_REQUIRED_FUNCTION) or globals_dict.get(GENERATED_CODE_REQUIRED_FUNCTION)
        if not callable(transform):
            raise ValueError("Generated code did not define a callable transform(rows, context)")
        result = transform(rows, context)
        result_queue.put({"ok": True, "result": result})
    except Exception as exc:
        result_queue.put({"ok": False, "error": str(exc)})


class FormFillService:
    def __init__(self) -> None:
        self.storage_service = get_storage_service()
        self.max_source_files = int(os.getenv("FORM_FILL_MAX_SOURCE_FILES", str(DEFAULT_MAX_SOURCE_FILES)))
        self.max_total_source_bytes = int(os.getenv("FORM_FILL_MAX_TOTAL_SOURCE_BYTES", str(DEFAULT_MAX_TOTAL_SOURCE_BYTES)))

        project = os.getenv("GOOGLE_CLOUD_PROJECT_ID")
        location = os.getenv("GOOGLE_CLOUD_LOCATION", "global")
        self.model_name = os.getenv("FORM_FILL_GEMINI_MODEL", "gemini-3.1-pro-preview")
        self.max_output_tokens = self._env_int("FORM_FILL_GEMINI_MAX_OUTPUT_TOKENS", 65536)
        self.batch_enabled = os.getenv("FORM_FILL_BATCH_ENABLED", "true").strip().lower() in {"1", "true", "yes", "y"}
        self.batch_max_rounds = self._env_int("FORM_FILL_BATCH_MAX_ROUNDS", 200)
        self.batch_tail_items = self._env_int("FORM_FILL_BATCH_TAIL_ITEMS", 10)
        self.batch_items_per_call = self._env_int("FORM_FILL_BATCH_ITEMS_PER_CALL", 100)
        self.mapping_chunk_size = self._env_int("FORM_FILL_MAPPING_CHUNK_SIZE", self.batch_items_per_call)
        self.output_max_attempts = max(1, self._env_int("FORM_FILL_OUTPUT_MAX_ATTEMPTS", 3))
        self.tabular_code_enabled = os.getenv("FORM_FILL_TABULAR_CODE_ENABLED", "true").strip().lower() in {"1", "true", "yes", "y"}
        self.tabular_code_timeout_seconds = max(
            1,
            self._env_int("FORM_FILL_TABULAR_CODE_TIMEOUT_SECONDS", GENERATED_CODE_TIMEOUT_SECONDS_DEFAULT),
        )
        self.tabular_code_max_output_bytes = max(
            1024,
            self._env_int("FORM_FILL_TABULAR_CODE_MAX_OUTPUT_BYTES", GENERATED_CODE_MAX_OUTPUT_BYTES_DEFAULT),
        )
        self.tabular_code_sample_rows = max(1, self._env_int("FORM_FILL_TABULAR_CODE_SAMPLE_ROWS", 10))
        try:
            self.near_token_ratio = float(os.getenv("FORM_FILL_NEAR_TOKEN_RATIO", "0.98"))
        except Exception:
            self.near_token_ratio = 0.98
        self.client = None
        if project:
            try:
                self.client = genai.Client(vertexai=True, project=project, location=location)
            except Exception as exc:
                logger.error("Failed to initialize Form Fill Vertex AI client: %s", exc)

    def _env_int(self, name: str, default: int) -> int:
        try:
            return int(os.getenv(name, str(default)))
        except (TypeError, ValueError):
            return default

    def _get_session(self) -> Session:
        return db_config.get_session()

    def _is_retryable_output_error(self, exc: Exception) -> bool:
        if isinstance(exc, ValueError):
            return False

        status_code = getattr(exc, "status_code", None)
        try:
            status_int = int(status_code) if status_code is not None else None
        except (TypeError, ValueError):
            status_int = None

        message = str(exc)
        lowered = message.lower()
        if status_int == 429 or (status_int is not None and status_int >= 500):
            return True
        if status_int == 499 and "cancelled" in lowered:
            return True
        if status_int == 400 and "invalid_argument" in lowered and "request contains an invalid argument" in lowered:
            return True
        if any(token in lowered for token in ("timeout", "timed out", "connection", "temporarily unavailable", "service unavailable")):
            return True
        return False

    def _form_fill_output_attempt(self, task_retry_count: Optional[int], task_execution_count: Optional[int]) -> Optional[int]:
        if task_retry_count is not None:
            try:
                return max(1, int(task_retry_count) + 1)
            except (TypeError, ValueError):
                return None
        if task_execution_count is not None:
            try:
                return max(1, int(task_execution_count))
            except (TypeError, ValueError):
                return None
        return None

    def _should_retry_output_error(
        self,
        exc: Exception,
        *,
        task_retry_count: Optional[int] = None,
        task_execution_count: Optional[int] = None,
    ) -> bool:
        attempt = self._form_fill_output_attempt(task_retry_count, task_execution_count)
        if attempt is None or attempt >= self.output_max_attempts:
            return False
        return self._is_retryable_output_error(exc)

    def _serialize_template(self, template: FormFillTemplate) -> FormFillTemplateResponse:
        return FormFillTemplateResponse(
            id=str(template.id),
            name=template.name,
            description=template.description,
            original_filename=template.original_filename,
            file_type=template.file_type,
            allow_docx_table_expansion=bool(template.allow_docx_table_expansion),
            file_size_bytes=int(template.file_size_bytes or 0),
            page_count=template.page_count,
            created_at=template.created_at,
            updated_at=template.updated_at,
        )

    def _serialize_run(self, run: FormFillRun) -> FormFillRunResponse:
        warnings = run.warnings if isinstance(run.warnings, list) else []
        source_payload = run.source_payload if isinstance(run.source_payload, dict) else None
        fill_plan = run.fill_plan if isinstance(run.fill_plan, dict) else None
        return FormFillRunResponse(
            id=str(run.id),
            status=run.status,
            source_mode=run.source_mode,
            source_filename=run.source_filename,
            source_file_type=run.source_file_type,
            source_files=[
                {
                    "id": str(source_file.id),
                    "original_filename": source_file.original_filename,
                    "file_type": source_file.file_type,
                    "file_size_bytes": int(source_file.file_size_bytes or 0),
                    "display_order": int(source_file.display_order or 0),
                }
                for source_file in (run.source_files or [])
            ],
            source_payload=source_payload,
            source_job_id=str(run.source_job_id) if run.source_job_id else None,
            source_run_id=str(run.source_run_id) if run.source_run_id else None,
            source_task_id=str(run.source_task_id) if run.source_task_id else None,
            target_mode=run.target_mode,
            target_template_id=str(run.target_template_id) if run.target_template_id else None,
            target_filename=run.target_filename,
            target_file_type=run.target_file_type,
            target_page_count=run.target_page_count,
            allow_docx_table_expansion=bool(run.allow_docx_table_expansion),
            output_format=run.output_format,
            repeat_mode=run.repeat_mode or REPEAT_MODE_SINGLE,
            total_outputs=int(run.total_outputs) if run.total_outputs is not None else 1,
            completed_outputs=int(run.completed_outputs or 0),
            failed_outputs=int(run.failed_outputs or 0),
            usage_basis=run.usage_basis,
            usage_pages=run.usage_pages,
            processing_strategy=run.processing_strategy,
            warnings=warnings,
            fill_plan=fill_plan,
            outputs=[
                {
                    "id": str(output.id),
                    "record_index": int(output.record_index or 0),
                    "record_label": output.record_label,
                    "status": output.status,
                    "warnings": output.warnings if isinstance(output.warnings, list) else [],
                    "fill_plan": output.fill_plan if isinstance(output.fill_plan, dict) else None,
                    "result_filename": output.result_filename,
                    "result_file_type": output.result_file_type,
                    "error_message": output.error_message,
                    "created_at": output.created_at,
                    "updated_at": output.updated_at,
                    "completed_at": output.completed_at,
                }
                for output in (run.outputs or [])
            ],
            result_filename=run.result_filename,
            result_file_type=run.result_file_type,
            error_message=run.error_message,
            created_at=run.created_at,
            updated_at=run.updated_at,
            completed_at=run.completed_at,
        )

    def _advisory_lock_keys(self, lock_id: str) -> tuple[int, int]:
        lock_uuid = uuid.UUID(str(lock_id))
        value = lock_uuid.int
        key_1 = (value >> 32) & 0xFFFFFFFF
        key_2 = value & 0xFFFFFFFF
        if key_1 >= 2**31:
            key_1 -= 2**32
        if key_2 >= 2**31:
            key_2 -= 2**32
        return int(key_1), int(key_2)

    def _try_advisory_lock(self, db: Session, lock_id: str) -> bool:
        try:
            bind = getattr(db, "bind", None)
            if not bind or getattr(getattr(bind, "dialect", None), "name", "") != "postgresql":
                return True
            key_1, key_2 = self._advisory_lock_keys(lock_id)
            locked = db.execute(text("SELECT pg_try_advisory_lock(:key_1, :key_2)"), {"key_1": key_1, "key_2": key_2}).scalar()
            return bool(locked)
        except Exception as exc:
            logger.warning("Form Fill advisory lock attempt failed; proceeding without lock: %s", exc)
            return True

    def _advisory_unlock(self, db: Session, lock_id: str) -> None:
        try:
            bind = getattr(db, "bind", None)
            if not bind or getattr(getattr(bind, "dialect", None), "name", "") != "postgresql":
                return
            key_1, key_2 = self._advisory_lock_keys(lock_id)
            db.execute(text("SELECT pg_advisory_unlock(:key_1, :key_2)"), {"key_1": key_1, "key_2": key_2})
        except Exception:
            return

    def list_templates(self, user_id: str) -> list[FormFillTemplateResponse]:
        db = self._get_session()
        try:
            templates = db.query(FormFillTemplate).filter(FormFillTemplate.user_id == user_id).order_by(
                FormFillTemplate.updated_at.desc(), FormFillTemplate.created_at.desc()
            ).all()
            return [self._serialize_template(item) for item in templates]
        finally:
            db.close()

    def delete_template(self, user_id: str, template_id: str) -> None:
        db = self._get_session()
        try:
            template = db.query(FormFillTemplate).filter(
                FormFillTemplate.id == uuid.UUID(str(template_id)),
                FormFillTemplate.user_id == user_id,
            ).first()
            if not template:
                raise ValueError("Form Fill template not found")

            db.delete(template)
            db.commit()
        except Exception:
            db.rollback()
            raise
        finally:
            db.close()

    def _task_source_files(self, db: Session, task_id: Any) -> list[str]:
        task_source_files = db.query(SourceFile.original_path).join(
            SourceFileToTask, SourceFile.id == SourceFileToTask.source_file_id
        ).filter(SourceFileToTask.task_id == task_id).all()
        return sort_paths_naturally(item[0] for item in task_source_files)

    def _payload_from_extraction_result(
        self,
        *,
        result: ExtractionResult,
        source_files: list[str],
        job_id: str,
        run_id: str,
        task_id: str,
    ) -> dict[str, Any]:
        if not result or not isinstance(result.extracted_data, dict):
            raise ValueError("Selected extraction result has no saved rows")

        extracted_data = result.extracted_data
        columns = extracted_data.get("columns") or []
        rows = extracted_data.get("results") or []
        if not isinstance(columns, list) or not isinstance(rows, list):
            raise ValueError("Selected extraction result is malformed")

        return {
            "kind": "extraction_result",
            "scope": SOURCE_SCOPE_TASK,
            "columns": columns,
            "rows": rows,
            "source_files": source_files,
            "job_id": str(job_id),
            "run_id": str(run_id),
            "task_id": str(task_id),
            "task_groups": [
                {
                    "task_id": str(task_id),
                    "source_files": source_files,
                    "columns": columns,
                    "rows": rows,
                }
            ],
        }

    def _combine_extraction_payloads(
        self,
        payloads: list[dict[str, Any]],
        *,
        job_id: str,
        run_id: str,
    ) -> dict[str, Any]:
        columns: list[str] = []
        seen_columns: set[str] = set()
        source_files: list[str] = []
        seen_source_files: set[str] = set()

        for payload in payloads:
            for column in payload.get("columns") or []:
                column_name = str(column)
                if column_name not in seen_columns:
                    seen_columns.add(column_name)
                    columns.append(column_name)
            for source_file in payload.get("source_files") or []:
                source_name = str(source_file)
                if source_name not in seen_source_files:
                    seen_source_files.add(source_name)
                    source_files.append(source_name)

        rows: list[list[Any]] = []
        for payload in payloads:
            payload_columns = [str(column) for column in (payload.get("columns") or [])]
            index_by_column = {column: index for index, column in enumerate(payload_columns)}
            for row in payload.get("rows") or []:
                if not isinstance(row, list):
                    continue
                rows.append([
                    row[index_by_column[column]] if column in index_by_column and index_by_column[column] < len(row) else None
                    for column in columns
                ])

        task_groups: list[dict[str, Any]] = []
        for payload in payloads:
            payload_groups = payload.get("task_groups")
            if isinstance(payload_groups, list) and payload_groups:
                for group in payload_groups:
                    if isinstance(group, dict):
                        task_groups.append(
                            {
                                "task_id": str(group.get("task_id") or payload.get("task_id") or ""),
                                "source_files": list(group.get("source_files") or []),
                                "columns": list(group.get("columns") or []),
                                "rows": list(group.get("rows") or []),
                            }
                        )
                continue
            task_groups.append(
                {
                    "task_id": str(payload.get("task_id") or ""),
                    "source_files": list(payload.get("source_files") or []),
                    "columns": list(payload.get("columns") or []),
                    "rows": list(payload.get("rows") or []),
                }
            )

        return {
            "kind": "extraction_result",
            "scope": SOURCE_SCOPE_ALL,
            "columns": columns,
            "rows": rows,
            "source_files": source_files,
            "task_groups": task_groups,
            "job_id": str(job_id),
            "run_id": str(run_id),
            "task_id": None,
        }

    def _load_task_extraction_source_payload(
        self,
        db: Session,
        *,
        user_id: str,
        job_id: str,
        run_id: str,
        task_id: str,
    ) -> dict[str, Any]:
        task = db.query(ExtractionTask).join(JobRun, JobRun.id == ExtractionTask.job_run_id).join(
            ExtractionJob, ExtractionJob.id == JobRun.job_id
        ).filter(
            ExtractionTask.id == uuid.UUID(str(task_id)),
            JobRun.id == uuid.UUID(str(run_id)),
            ExtractionJob.id == uuid.UUID(str(job_id)),
            ExtractionJob.user_id == user_id,
        ).first()
        if not task:
            raise ValueError("Selected extraction result not found")

        result = db.query(ExtractionResult).filter(ExtractionResult.task_id == task.id).first()
        return self._payload_from_extraction_result(
            result=result,
            source_files=self._task_source_files(db, task.id),
            job_id=job_id,
            run_id=run_id,
            task_id=str(task_id),
        )

    def _load_all_extraction_source_payload(
        self,
        db: Session,
        *,
        user_id: str,
        job_id: str,
        run_id: str,
    ) -> dict[str, Any]:
        target_run = db.query(JobRun).join(ExtractionJob, ExtractionJob.id == JobRun.job_id).filter(
            JobRun.id == uuid.UUID(str(run_id)),
            ExtractionJob.id == uuid.UUID(str(job_id)),
            ExtractionJob.user_id == user_id,
        ).first()
        if not target_run:
            raise ValueError("Extraction run not found")

        results_with_tasks = db.query(ExtractionResult, ExtractionTask).join(
            ExtractionTask, ExtractionResult.task_id == ExtractionTask.id
        ).filter(
            ExtractionTask.job_run_id == target_run.id
        ).all()

        source_files_by_task = {
            task.id: self._task_source_files(db, task.id)
            for _, task in results_with_tasks
        }
        results_with_tasks.sort(
            key=lambda item: (
                getattr(item[1], "result_set_index", 0) or 0,
                1 if not source_files_by_task.get(item[1].id) else 0,
                natural_text_key(source_files_by_task[item[1].id][0]) if source_files_by_task.get(item[1].id) else (),
                item[0].processed_at.timestamp() if item[0].processed_at else 0,
                str(item[0].id),
            )
        )

        payloads = [
            self._payload_from_extraction_result(
                result=result,
                source_files=source_files_by_task.get(task.id, []),
                job_id=job_id,
                run_id=run_id,
                task_id=str(task.id),
            )
            for result, task in results_with_tasks
        ]
        if not payloads:
            raise ValueError("Extraction run has no saved rows")
        return self._combine_extraction_payloads(payloads, job_id=job_id, run_id=run_id)

    def _load_extraction_source_payload(
        self,
        db: Session,
        *,
        user_id: str,
        job_id: str,
        run_id: str,
        task_id: Optional[str] = None,
        source_scope: Optional[str] = None,
    ) -> dict[str, Any]:
        if _normalize_source_scope(source_scope) == SOURCE_SCOPE_ALL:
            return self._load_all_extraction_source_payload(db, user_id=user_id, job_id=job_id, run_id=run_id)
        if not task_id:
            raise ValueError("Select an extraction result or all rows")
        return self._load_task_extraction_source_payload(
            db,
            user_id=user_id,
            job_id=job_id,
            run_id=run_id,
            task_id=task_id,
        )

    def get_extraction_source_preview(
        self,
        user_id: str,
        *,
        job_id: str,
        run_id: str,
        task_id: Optional[str] = None,
        source_scope: Optional[str] = None,
    ) -> FormFillExtractionSourcePreviewResponse:
        db = self._get_session()
        try:
            payload = self._load_extraction_source_payload(
                db,
                user_id=user_id,
                job_id=job_id,
                run_id=run_id,
                task_id=task_id,
                source_scope=source_scope,
            )
            return FormFillExtractionSourcePreviewResponse(
                job_id=job_id,
                run_id=run_id,
                task_id=payload.get("task_id"),
                source_scope=payload.get("scope") or _normalize_source_scope(source_scope),
                source_files=list(payload.get("source_files") or []),
                columns=list(payload.get("columns") or []),
                rows=list(payload.get("rows") or []),
            )
        finally:
            db.close()

    async def _upload_bytes(self, object_name: str, content: bytes) -> None:
        if not hasattr(self.storage_service, "upload_file_content"):
            raise RuntimeError("Form Fill requires Google Cloud Storage")
        await self.storage_service.upload_file_content(content, object_name)

    async def _download_to_local(self, object_name: str, local_path: str) -> None:
        if not hasattr(self.storage_service, "download_file"):
            raise RuntimeError("Form Fill requires Google Cloud Storage")
        await self.storage_service.download_file(object_name, local_path)

    async def _count_target_pages_from_local_path(
        self,
        *,
        local_path: str,
        filename: str,
        mime_type: str,
        temp_dir: str,
    ) -> int:
        normalized_mime = (mime_type or "").lower()
        if normalized_mime == PDF_MIME:
            page_count = page_counting_service.count_pages_from_file_path(local_path, filename)
        elif normalized_mime == DOCX_MIME:
            converter = get_document_conversion_service()
            pdf_path = await converter.convert_docx_local_to_pdf_local(local_path, out_dir=temp_dir)
            page_count = page_counting_service.count_pages_from_file_path(pdf_path, f"{Path(filename).stem}.pdf")
        else:
            raise ValueError("Target file must be a PDF or DOCX")

        if page_count is None or page_count <= 0:
            raise ValueError("Could not determine target page count")
        return int(page_count)

    async def _count_target_pages_from_bytes(self, *, content: bytes, filename: str, mime_type: str) -> int:
        normalized_mime = (mime_type or "").lower()
        if normalized_mime == PDF_MIME:
            page_count = page_counting_service.count_pages_from_content(content, filename)
            if page_count is None or page_count <= 0:
                raise ValueError("Could not determine target page count")
            return int(page_count)

        with tempfile.TemporaryDirectory(prefix="form_fill_target_pages_") as temp_dir:
            local_path = os.path.join(temp_dir, f"target{_safe_ext(filename, '.bin')}")
            with open(local_path, "wb") as handle:
                handle.write(content)
            return await self._count_target_pages_from_local_path(
                local_path=local_path,
                filename=filename,
                mime_type=normalized_mime,
                temp_dir=temp_dir,
            )

    async def _ensure_run_target_page_count(self, db: Session, run: FormFillRun, target_local_path: str, temp_dir: str) -> int:
        page_count = int(run.target_page_count or 0)
        if page_count > 0:
            return page_count

        if run.target_template_id:
            template = db.query(FormFillTemplate).filter(FormFillTemplate.id == run.target_template_id).first()
            template_page_count = int(template.page_count or 0) if template else 0
            if template_page_count > 0:
                run.target_page_count = template_page_count
                db.commit()
                return template_page_count

        page_count = await self._count_target_pages_from_local_path(
            local_path=target_local_path,
            filename=run.target_filename,
            mime_type=run.target_file_type,
            temp_dir=temp_dir,
        )
        run.target_page_count = page_count
        if run.target_template_id:
            template = db.query(FormFillTemplate).filter(FormFillTemplate.id == run.target_template_id).first()
            if template and not template.page_count:
                template.page_count = page_count
        db.commit()
        return page_count

    def _check_usage_limit_or_raise(self, db: Session, *, user_id: str, page_count: int) -> None:
        if page_count <= 0:
            return

        from services.billing_service import get_billing_service

        billing_service = get_billing_service(db)
        if billing_service.check_page_limit(user_id, page_count):
            return

        billing_info = billing_service.get_billing_info(user_id)
        plan_name = billing_info.get("plan_display_name") or "current"
        pages_used = int(billing_info.get("pages_used") or 0)
        pages_included = int(billing_info.get("pages_included") or 0)
        pages_remaining = max(0, pages_included - pages_used)
        raise ValueError(
            f"Cannot start Form Fill: processing {page_count} target pages would exceed your {plan_name} plan limit. "
            f"You have {pages_remaining} pages remaining out of {pages_included}. "
            "Please upgrade your plan or reduce the number of target pages."
        )

    def _record_usage_for_run(self, db: Session, run: FormFillRun) -> None:
        usage_pages = int(run.usage_pages or 0)
        if usage_pages <= 0:
            return

        from services.billing_service import PlanLimitExceeded, get_billing_service

        try:
            event_id = get_billing_service(db).record_usage(
                user_id=run.user_id,
                pages=usage_pages,
                source="form_fill_run",
                form_fill_run_id=str(run.id),
                notes=f"Form Fill run for target {run.target_filename}",
            )
            logger.info("Recorded %s Form Fill usage pages for run %s (event %s)", usage_pages, run.id, event_id)
        except PlanLimitExceeded:
            raise
        except Exception as exc:
            logger.error("Failed to record Form Fill usage for run %s: %s", run.id, exc)

    async def create_run(
        self,
        *,
        user_id: str,
        source_files: Optional[list[Any]] = None,
        target_file: Any = None,
        template_id: Optional[str] = None,
        output_format: Optional[str] = None,
        repeat_mode: Optional[str] = None,
        allow_docx_table_expansion: Optional[bool] = None,
        save_template_name: Optional[str] = None,
        save_template_description: Optional[str] = None,
        source_job_id: Optional[str] = None,
        source_run_id: Optional[str] = None,
        source_task_id: Optional[str] = None,
        source_scope: Optional[str] = None,
    ) -> FormFillRunResponse:
        db = self._get_session()
        try:
            uploaded_source_files = [item for item in (source_files or []) if item is not None]
            normalized_source_scope = _normalize_source_scope(source_scope)
            source_from_extraction = bool(
                source_job_id and source_run_id and (source_task_id or normalized_source_scope == SOURCE_SCOPE_ALL)
            )
            normalized_repeat_mode = _normalize_repeat_mode(repeat_mode)
            if bool(uploaded_source_files) == source_from_extraction:
                raise ValueError("Provide either source files or an extraction result source")
            if len(uploaded_source_files) > self.max_source_files:
                raise ValueError(f"Form Fill supports up to {self.max_source_files} source files per run")
            if normalized_repeat_mode == REPEAT_MODE_SOURCE_ROWS and uploaded_source_files and len(uploaded_source_files) != 1:
                raise ValueError("Repeat mode currently supports one CSV or XLSX source file")
            if bool(target_file) == bool(template_id):
                raise ValueError("Provide either a target file or a saved template")

            run = FormFillRun(
                user_id=user_id,
                status="pending",
                source_mode="upload" if uploaded_source_files else "extraction_result",
                target_mode="upload" if target_file else "template",
                target_filename="pending",
                target_file_type="application/octet-stream",
                allow_docx_table_expansion=False,
                target_gcs_object_name="pending",
                target_file_size_bytes=0,
                output_format="pending",
                repeat_mode=normalized_repeat_mode,
                total_outputs=0 if normalized_repeat_mode == REPEAT_MODE_SOURCE_ROWS else 1,
                completed_outputs=0,
                failed_outputs=0,
            )
            db.add(run)
            db.flush()

            if uploaded_source_files:
                total_source_bytes = 0
                source_filenames: list[str] = []
                source_mime_types: list[str] = []
                for index, source_file in enumerate(uploaded_source_files):
                    source_bytes = await source_file.read()
                    if not source_bytes:
                        raise ValueError("Source file is empty")

                    source_filename = source_file.filename or f"source-{index + 1}"
                    source_mime = (source_file.content_type or _guess_mime_type(source_filename)).lower()
                    if source_mime not in SUPPORTED_SOURCE_MIME_TYPES:
                        raise ValueError(f"Unsupported source file type: {source_filename}")
                    if normalized_repeat_mode == REPEAT_MODE_SOURCE_ROWS and source_mime not in TABULAR_SOURCE_MIME_TYPES:
                        raise ValueError("Repeat mode requires a CSV or XLSX source file")

                    total_source_bytes += len(source_bytes)
                    if total_source_bytes > self.max_total_source_bytes:
                        max_mb = self.max_total_source_bytes // (1024 * 1024)
                        raise ValueError(f"Form Fill source files must be {max_mb} MB or less in total")

                    source_object_name = (
                        f"form-fill/{user_id}/runs/{run.id}/sources/{index + 1}-{uuid.uuid4()}"
                        f"{_safe_ext(source_filename, '.bin')}"
                    )
                    await self._upload_bytes(source_object_name, source_bytes)
                    db.add(
                        FormFillSourceFile(
                            run_id=run.id,
                            original_filename=source_filename,
                            file_type=source_mime,
                            gcs_object_name=source_object_name,
                            file_size_bytes=len(source_bytes),
                            display_order=index,
                        )
                    )
                    source_filenames.append(source_filename)
                    source_mime_types.append(source_mime)

                run.source_filename = source_filenames[0] if len(source_filenames) == 1 else f"{len(source_filenames)} source files"
                run.source_file_type = source_mime_types[0] if len(set(source_mime_types)) == 1 else "multiple"
                run.source_file_size_bytes = total_source_bytes
                if normalized_repeat_mode not in {REPEAT_MODE_SOURCE_ROWS, REPEAT_MODE_ALL_SOURCES} and len(uploaded_source_files) > 1:
                    run.total_outputs = len(uploaded_source_files)
            else:
                payload = self._load_extraction_source_payload(
                    db,
                    user_id=user_id,
                    job_id=str(source_job_id),
                    run_id=str(source_run_id),
                    task_id=str(source_task_id) if source_task_id else None,
                    source_scope=normalized_source_scope,
                )
                run.source_filename = "Extraction Results - All Rows" if normalized_source_scope == SOURCE_SCOPE_ALL else "Extraction Results"
                run.source_file_type = "application/json"
                run.source_payload = payload
                run.source_job_id = uuid.UUID(str(source_job_id))
                run.source_run_id = uuid.UUID(str(source_run_id))
                run.source_task_id = uuid.UUID(str(source_task_id)) if source_task_id else None
                if normalized_repeat_mode not in {REPEAT_MODE_SOURCE_ROWS, REPEAT_MODE_ALL_SOURCES}:
                    extraction_task_units = self._extraction_task_units_from_payload(payload)
                    if len(extraction_task_units) > 1:
                        run.total_outputs = len(extraction_task_units)

            if target_file:
                target_bytes = await target_file.read()
                if not target_bytes:
                    raise ValueError("Target file is empty")

                target_filename = target_file.filename or "target"
                target_mime = (target_file.content_type or _guess_mime_type(target_filename)).lower()
                if target_mime not in SUPPORTED_TARGET_MIME_TYPES:
                    raise ValueError("Target file must be a PDF or DOCX")
                target_page_count = await self._count_target_pages_from_bytes(
                    content=target_bytes,
                    filename=target_filename,
                    mime_type=target_mime,
                )

                target_object_name = f"form-fill/{user_id}/runs/{run.id}/target{_safe_ext(target_filename, '.bin')}"
                await self._upload_bytes(target_object_name, target_bytes)

                run.target_filename = target_filename
                run.target_file_type = target_mime
                run.target_gcs_object_name = target_object_name
                run.target_file_size_bytes = len(target_bytes)
                run.target_page_count = target_page_count
                run.allow_docx_table_expansion = bool(allow_docx_table_expansion) if target_mime == DOCX_MIME else False

                if save_template_name and save_template_name.strip():
                    template = FormFillTemplate(
                        user_id=user_id,
                        name=save_template_name.strip(),
                        description=(save_template_description or "").strip() or None,
                        original_filename=target_filename,
                        file_type=target_mime,
                        allow_docx_table_expansion=run.allow_docx_table_expansion,
                        gcs_object_name=f"form-fill/{user_id}/templates/{uuid.uuid4()}/target{_safe_ext(target_filename, '.bin')}",
                        file_size_bytes=len(target_bytes),
                        page_count=target_page_count,
                    )
                    await self._upload_bytes(template.gcs_object_name, target_bytes)
                    db.add(template)
                    db.flush()
            else:
                template = db.query(FormFillTemplate).filter(
                    FormFillTemplate.id == uuid.UUID(str(template_id)),
                    FormFillTemplate.user_id == user_id,
                ).first()
                if not template:
                    raise ValueError("Form Fill template not found")

                run.target_template_id = template.id
                run.target_filename = template.original_filename
                run.target_file_type = template.file_type
                run.target_gcs_object_name = template.gcs_object_name
                run.target_file_size_bytes = template.file_size_bytes
                run.target_page_count = template.page_count
                if template.file_type == DOCX_MIME:
                    if allow_docx_table_expansion is None:
                        run.allow_docx_table_expansion = bool(template.allow_docx_table_expansion)
                    else:
                        run.allow_docx_table_expansion = bool(allow_docx_table_expansion)
                else:
                    run.allow_docx_table_expansion = False

            normalized_output_format = _normalize_output_format(output_format)
            target_default_format = "docx" if run.target_file_type == DOCX_MIME else "pdf"
            if normalized_output_format is None:
                normalized_output_format = target_default_format

            if run.target_file_type == PDF_MIME and normalized_output_format != "pdf":
                raise ValueError("PDF targets currently output PDF only")
            if run.target_file_type == DOCX_MIME and normalized_output_format not in {"docx", "pdf"}:
                raise ValueError("DOCX targets must output DOCX or PDF")

            run.output_format = normalized_output_format
            db.commit()
            db.refresh(run)

            await cloud_run_task_service.enqueue_form_fill_task(str(run.id))
            return self._serialize_run(run)
        except Exception:
            db.rollback()
            raise
        finally:
            db.close()

    def get_run(self, user_id: str, run_id: str) -> FormFillRunResponse:
        db = self._get_session()
        try:
            run = db.query(FormFillRun).filter(
                FormFillRun.id == uuid.UUID(str(run_id)),
                FormFillRun.user_id == user_id,
            ).first()
            if not run:
                raise ValueError("Form Fill run not found")
            return self._serialize_run(run)
        finally:
            db.close()

    def list_runs(
        self,
        user_id: str,
        *,
        limit: int = 25,
        offset: int = 0,
        status: Optional[str] = None,
    ) -> dict[str, Any]:
        db = self._get_session()
        try:
            query = db.query(FormFillRun).filter(FormFillRun.user_id == user_id)
            if status:
                query = query.filter(FormFillRun.status == status)

            total = query.count()
            runs = query.order_by(
                FormFillRun.updated_at.desc(),
                FormFillRun.created_at.desc(),
            ).offset(offset).limit(limit).all()

            return {
                "runs": [self._serialize_run(run) for run in runs],
                "total": total,
                "limit": limit,
                "offset": offset,
            }
        finally:
            db.close()

    def get_run_result_metadata(self, user_id: str, run_id: str) -> FormFillRun:
        db = self._get_session()
        try:
            run = db.query(FormFillRun).filter(
                FormFillRun.id == uuid.UUID(str(run_id)),
                FormFillRun.user_id == user_id,
            ).first()
            if not run:
                raise ValueError("Form Fill run not found")
            if run.status not in {"completed", "completed_with_errors"} or not run.result_gcs_object_name or not run.result_filename:
                raise ValueError("Form Fill output is not ready")
            db.expunge(run)
            return run
        finally:
            db.close()

    def _part_from_uri(self, uri: str, mime_type: str) -> Any:
        return types.Part.from_uri(file_uri=uri, mime_type=mime_type)

    def _get_resp_text(self, response: Any) -> Optional[str]:
        if response is None:
            return None
        text = getattr(response, "text", None)
        if isinstance(text, str) and text.strip():
            return text
        try:
            candidates = getattr(response, "candidates", None)
            if isinstance(candidates, list) and candidates:
                parts = getattr(getattr(candidates[0], "content", None), "parts", None)
                if isinstance(parts, list):
                    joined = "".join(str(getattr(part, "text", "") or "") for part in parts)
                    if joined.strip():
                        return joined
        except Exception:
            pass
        return None

    def _get_finish_reason(self, response: Any) -> Optional[str]:
        try:
            candidates = getattr(response, "candidates", None)
            if isinstance(candidates, list) and candidates:
                finish_reason = getattr(candidates[0], "finish_reason", None)
                return str(finish_reason) if finish_reason is not None else None
        except Exception:
            return None
        return None

    def _get_usage_counts(self, response: Any) -> dict[str, Optional[int]]:
        usage = getattr(response, "usage_metadata", None) or getattr(response, "usage", None)
        counts: dict[str, Optional[int]] = {"prompt_tokens": None, "output_tokens": None, "total_tokens": None}
        if usage is None:
            return counts

        if isinstance(usage, dict):
            for key in ("prompt_token_count", "prompt_tokens", "input_tokens"):
                if isinstance(usage.get(key), int):
                    counts["prompt_tokens"] = usage.get(key)
                    break
            for key in ("candidates_token_count", "output_tokens", "completion_tokens"):
                if isinstance(usage.get(key), int):
                    counts["output_tokens"] = usage.get(key)
                    break
            for key in ("total_token_count", "total_tokens"):
                if isinstance(usage.get(key), int):
                    counts["total_tokens"] = usage.get(key)
                    break
            return counts

        for attr, key in (
            ("prompt_token_count", "prompt_tokens"),
            ("prompt_tokens", "prompt_tokens"),
            ("input_tokens", "prompt_tokens"),
            ("candidates_token_count", "output_tokens"),
            ("output_tokens", "output_tokens"),
            ("completion_tokens", "output_tokens"),
            ("total_token_count", "total_tokens"),
            ("total_tokens", "total_tokens"),
        ):
            value = getattr(usage, attr, None)
            if isinstance(value, int) and counts[key] is None:
                counts[key] = value
        return counts

    def _looks_truncated(
        self,
        response: Any,
        *,
        parsed_ok: bool,
        text: Optional[str],
        parse_exc: Optional[Exception] = None,
    ) -> bool:
        finish_reason = (self._get_finish_reason(response) or "").upper()
        if "MAX_TOKENS" in finish_reason or "LENGTH" in finish_reason:
            return True

        output_tokens = self._get_usage_counts(response).get("output_tokens")
        if isinstance(output_tokens, int) and self.max_output_tokens:
            if output_tokens >= int(self.max_output_tokens * self.near_token_ratio):
                return True

        if not parsed_ok and parse_exc is not None:
            message = str(parse_exc).lower()
            if "unterminated" in message or "expecting" in message or "eof" in message or "end of" in message:
                return True

        if isinstance(text, str) and text.strip() and not parsed_ok:
            if text.rstrip()[-1] not in {"}", "]"}:
                return True
        return False

    def _coerce_parsed_obj(self, obj: Any) -> Any:
        if hasattr(obj, "model_dump"):
            try:
                return obj.model_dump()
            except Exception:
                pass
        if hasattr(obj, "dict"):
            try:
                return obj.dict()
            except Exception:
                pass
        return obj

    def _parse_response_payload(self, response: Any) -> dict[str, Any]:
        parsed = self._coerce_parsed_obj(getattr(response, "parsed", None))
        if isinstance(parsed, dict):
            return parsed
        text = self._get_resp_text(response)
        if not text:
            raise ValueError("Gemini returned an empty response")
        payload = json.loads(text)
        if not isinstance(payload, dict):
            raise ValueError("Gemini returned an unexpected response")
        return payload

    def _generation_config(self, schema: types.Schema, *, temperature: float = 0.1) -> types.GenerateContentConfig:
        return types.GenerateContentConfig(
            response_mime_type="application/json",
            response_schema=schema,
            temperature=temperature,
            max_output_tokens=self.max_output_tokens,
        )

    def _salvage_collection_from_text(self, text: str, collection_key: str) -> list[dict[str, Any]]:
        if not isinstance(text, str) or not text:
            return []

        key_index = text.find(f'"{collection_key}"')
        if key_index == -1:
            key_index = text.find(f"'{collection_key}'")
        if key_index == -1:
            return []

        array_index = text.find("[", key_index)
        if array_index == -1:
            return []

        decoder = json.JSONDecoder()
        items: list[dict[str, Any]] = []
        index = array_index + 1
        while index < len(text):
            while index < len(text) and text[index] in " \t\r\n,":
                index += 1
            if index >= len(text) or text[index] == "]":
                break
            try:
                value, end = decoder.raw_decode(text, index)
            except Exception:
                break
            if isinstance(value, dict):
                items.append(value)
            index = end
        return items

    def _parse_or_salvage_collection(
        self,
        response: Any,
        collection_key: str,
    ) -> tuple[list[dict[str, Any]], list[str], bool, Optional[str], Optional[Exception]]:
        text = self._get_resp_text(response)
        try:
            payload = self._parse_response_payload(response)
            collection = payload.get(collection_key)
            if collection is None:
                collection = []
            if not isinstance(collection, list):
                raise ValueError(f"Gemini response '{collection_key}' is not a list")
            warnings = payload.get("warnings") or []
            return (
                [item for item in collection if isinstance(item, dict)],
                [str(item) for item in warnings if str(item).strip()] if isinstance(warnings, list) else [],
                True,
                text,
                None,
            )
        except Exception as exc:
            salvaged = self._salvage_collection_from_text(text or "", collection_key)
            if salvaged:
                return salvaged, [], False, text, exc
            return [], [], False, text, exc

    def _collection_item_key(self, item: dict[str, Any]) -> str:
        try:
            return json.dumps(item, separators=(",", ":"), ensure_ascii=True, sort_keys=True)
        except Exception:
            return str(item)

    def _compute_collection_overlap(self, existing: list[dict[str, Any]], new_items: list[dict[str, Any]], max_items: int) -> int:
        if not existing or not new_items or max_items <= 0:
            return 0
        max_items = min(max_items, len(existing), len(new_items))
        existing_keys = [self._collection_item_key(item) for item in existing[-max_items:]]
        new_keys = [self._collection_item_key(item) for item in new_items[:max_items]]
        for size in range(max_items, 0, -1):
            if existing_keys[-size:] == new_keys[:size]:
                return size
        return 0

    def _looks_like_collection_restart(self, existing: list[dict[str, Any]], new_items: list[dict[str, Any]]) -> bool:
        if len(existing) < 50 or len(new_items) < 10:
            return False
        return [self._collection_item_key(item) for item in existing[:10]] == [
            self._collection_item_key(item) for item in new_items[:10]
        ]

    def _build_collection_continuation_prompt(
        self,
        prompt: str,
        *,
        collection_key: str,
        prior_items: list[dict[str, Any]],
        total_returned: int,
        max_items: int,
    ) -> str:
        prior_json = json.dumps(prior_items, separators=(",", ":"), ensure_ascii=True)
        max_items_line = ""
        if max_items > 0:
            max_items_line = f"- Return at most {max_items} {collection_key} entries in this response.\n"
        return (
            f"{prompt}\n\n"
            "Next batch:\n"
            f"- You previously returned {total_returned} '{collection_key}' entries for this SAME target and source.\n"
            f"- Continue with the next '{collection_key}' entries that come AFTER the final entry in prior_entries below.\n"
            "- Do not restart from the beginning. Do not repeat any entries from prior_entries.\n"
            "- Do not summarize, collapse, or replace remaining entries with warnings.\n"
            "- Do not write 'see attached', 'too many transactions', or any output-limit workaround.\n"
            "- Do not say only the first N entries were inserted. Do not ask the user to request continuation; this continuation request is already that mechanism.\n"
            "- Do not mention output limits or token limits. Return concrete entries only.\n"
            f"{max_items_line}"
            f"- If there are no more entries, return {{\"{collection_key}\":[],\"warnings\":[]}}.\n\n"
            f"prior_entries (all entries already returned, in order): {prior_json}\n"
        )

    def _build_collection_batch_prompt(self, prompt: str, *, collection_key: str, max_items: int) -> str:
        max_items_line = ""
        if max_items > 0:
            max_items_line = f"- This is batch 1. Return at most {max_items} '{collection_key}' entries in this response.\n"
        return (
            f"{prompt}\n\n"
            "Batching rules:\n"
            f"{max_items_line}"
            "- Return concrete entries only; do not summarize or collapse entries because there are many.\n"
            "- Do not mention output limits or token limits. If more entries remain after this batch, stop cleanly after a complete entry and the system will request the next batch.\n"
            f"- If there are no entries to return, return {{\"{collection_key}\":[],\"warnings\":[]}}.\n"
        )

    def _filter_output_limit_warnings(self, warnings: list[str], *, label: str) -> list[str]:
        kept: list[str] = []
        suppressed: list[str] = []
        for warning in warnings:
            lowered = warning.lower()
            if any(pattern in lowered for pattern in OUTPUT_LIMIT_WARNING_PATTERNS):
                suppressed.append(warning)
            else:
                kept.append(warning)
        if suppressed:
            logger.info("Suppressed Form Fill output-limit warnings (%s): %s", label, suppressed)
        return kept

    def _generate_collection_json_response(
        self,
        contents: list[Any],
        *,
        prompt: str,
        schema: types.Schema,
        collection_key: str,
        label: str,
        continue_on_full_batch: bool = False,
    ) -> dict[str, Any]:
        self._ensure_client()
        batch_mode = bool(
            continue_on_full_batch
            and self.batch_enabled
            and self.batch_items_per_call > 0
        )
        items_per_call = self.batch_items_per_call
        initial_prompt = (
            self._build_collection_batch_prompt(prompt, collection_key=collection_key, max_items=items_per_call)
            if batch_mode
            else prompt
        )
        response = self.client.models.generate_content(
            model=self.model_name,
            contents=contents + [initial_prompt],
            config=self._generation_config(schema),
        )
        items, warnings, parsed_ok, text, parse_exc = self._parse_or_salvage_collection(response, collection_key)
        truncated = self._looks_truncated(response, parsed_ok=parsed_ok, text=text, parse_exc=parse_exc)
        logger.info(
            "Form Fill Gemini response (%s): %s=%s, parsed_ok=%s, truncated=%s, finish_reason=%s, usage=%s",
            label,
            collection_key,
            len(items),
            parsed_ok,
            truncated,
            self._get_finish_reason(response),
            self._get_usage_counts(response),
        )

        if not parsed_ok and not truncated:
            raise parse_exc or ValueError("Failed to parse Gemini response")

        all_items = list(items)
        all_warnings = list(warnings)
        full_batch = (
            continue_on_full_batch
            and items_per_call > 0
            and len(items) >= items_per_call
        )
        if not self.batch_enabled or not all_items or (not truncated and not full_batch):
            if continue_on_full_batch and self.batch_enabled:
                all_warnings = self._filter_output_limit_warnings(all_warnings, label=label)
            return {collection_key: all_items, "warnings": all_warnings}

        no_growth_rounds = 0
        for round_index in range(1, self.batch_max_rounds + 1):
            continuation_prompt = self._build_collection_continuation_prompt(
                prompt,
                collection_key=collection_key,
                prior_items=all_items,
                total_returned=len(all_items),
                max_items=items_per_call,
            )
            continuation_response = self.client.models.generate_content(
                model=self.model_name,
                contents=contents + [continuation_prompt],
                config=self._generation_config(schema, temperature=0.0),
            )
            new_items, new_warnings, parsed_ok2, text2, parse_exc2 = self._parse_or_salvage_collection(
                continuation_response,
                collection_key,
            )
            truncated2 = self._looks_truncated(continuation_response, parsed_ok=parsed_ok2, text=text2, parse_exc=parse_exc2)
            if not parsed_ok2 and not truncated2:
                raise parse_exc2 or ValueError("Failed to parse Gemini continuation response")

            overlap = self._compute_collection_overlap(all_items, new_items, max(1, self.batch_tail_items))
            if overlap == 0 and self._looks_like_collection_restart(all_items, new_items):
                logger.warning("Form Fill Gemini continuation (%s) restarted from the beginning; stopping", label)
                break

            appended = new_items[overlap:] if overlap else new_items
            all_items.extend(appended)
            all_warnings.extend(new_warnings)
            logger.info(
                "Form Fill Gemini continuation (%s) round=%s: returned=%s, added=%s, total=%s, truncated=%s, finish_reason=%s, usage=%s",
                label,
                round_index,
                len(new_items),
                len(appended),
                len(all_items),
                truncated2,
                self._get_finish_reason(continuation_response),
                self._get_usage_counts(continuation_response),
            )

            if not new_items:
                break
            if appended:
                no_growth_rounds = 0
            else:
                no_growth_rounds += 1
            full_batch2 = (
                continue_on_full_batch
                and items_per_call > 0
                and len(new_items) >= items_per_call
            )
            if no_growth_rounds >= 2:
                break
            if not truncated2 and not full_batch2:
                break

        if continue_on_full_batch and self.batch_enabled:
            all_warnings = self._filter_output_limit_warnings(all_warnings, label=label)
        return {collection_key: all_items, "warnings": all_warnings}

    def _extract_pdf_form_fields(self, local_path: str) -> list[str]:
        from PyPDF2 import PdfReader

        reader = PdfReader(local_path)
        fields = reader.get_fields() or {}
        return [str(name) for name in fields.keys()]

    def _extract_pdf_text(self, local_path: str, max_chars: int = 20000) -> str:
        doc = fitz.open(local_path)
        parts: list[str] = []
        try:
            for page_index, page in enumerate(doc, start=1):
                text = (page.get_text("text") or "").strip()
                if text:
                    parts.append(f"Page {page_index}:\n{text}")
                joined = "\n\n".join(parts)
                if len(joined) >= max_chars:
                    return joined[:max_chars]
        finally:
            doc.close()
        return "\n\n".join(parts)[:max_chars]

    def _extract_docx_text(self, local_path: str, max_chars: int = 20000) -> str:
        from docx import Document as DocxDocument

        doc = DocxDocument(local_path)
        parts: list[str] = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in self._docx_row_cells(row) if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        joined = "\n\n".join(parts)
        return joined[:max_chars]

    def _docx_row_cells(self, row: Any) -> list[Any]:
        from docx.table import _Cell

        return [_Cell(tc, row.table) for tc in row._tr.tc_lst]

    def _docx_table_has_merged_cells(self, table: Any) -> bool:
        for row in table.rows:
            for cell in self._docx_row_cells(row):
                if cell.grid_span > 1 or cell._tc.vMerge is not None:
                    return True
        return False

    def _extract_docx_placeholders(self, local_path: str) -> list[str]:
        from docx import Document as DocxDocument

        doc = DocxDocument(local_path)
        found: list[str] = []

        def collect(text: str) -> None:
            for match in DOCX_PLACEHOLDER_RE.findall(text or ""):
                if match not in found:
                    found.append(match)

        for paragraph in doc.paragraphs:
            collect(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in self._docx_row_cells(row):
                    collect(cell.text)
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                collect(paragraph.text)
            for paragraph in section.footer.paragraphs:
                collect(paragraph.text)
        return found

    def _markdown_table(self, columns: list[str], rows: list[list[Any]], max_rows: int) -> str:
        limited_rows = rows[:max_rows]
        safe_columns = [str(col) for col in columns]
        if not safe_columns:
            return ""
        lines = ["| " + " | ".join(safe_columns) + " |", "| " + " | ".join(["---"] * len(safe_columns)) + " |"]
        for row in limited_rows:
            values = []
            for index in range(len(safe_columns)):
                value = row[index] if isinstance(row, list) and index < len(row) else None
                values.append(str(value) if value is not None else "")
            lines.append("| " + " | ".join(values) + " |")
        return "\n".join(lines)

    def _jsonable_cell(self, value: Any) -> Any:
        if value is None or isinstance(value, (str, int, float, bool)):
            return value
        if isinstance(value, datetime):
            return value.isoformat()
        return str(value)

    def _safe_tabular_columns(self, columns: list[Any]) -> list[str]:
        seen: dict[str, int] = {}
        safe_columns: list[str] = []
        for index, column in enumerate(columns):
            base = str(column).strip() if column is not None else ""
            base = base or f"Column {index + 1}"
            count = seen.get(base, 0)
            seen[base] = count + 1
            safe_columns.append(base if count == 0 else f"{base} {count + 1}")
        return safe_columns

    def _tabular_context_from_rows(
        self,
        columns: list[Any],
        rows: list[Any],
        *,
        source_name: str,
        sheet_name: Optional[str] = None,
    ) -> dict[str, Any]:
        safe_columns = self._safe_tabular_columns(columns)
        output_rows: list[dict[str, Any]] = []
        for row_index, row in enumerate(rows, start=1):
            if not isinstance(row, (list, tuple)):
                continue
            payload = {
                column: self._jsonable_cell(row[column_index]) if column_index < len(row) else None
                for column_index, column in enumerate(safe_columns)
            }
            if not any(value is not None and str(value).strip() for value in payload.values()):
                continue
            payload["_row_number"] = row_index + 1
            payload["_source_file"] = source_name
            if sheet_name:
                payload["_sheet_name"] = sheet_name
            output_rows.append(payload)
        return {
            "columns": safe_columns,
            "rows": output_rows,
            "source_files": [source_name] if source_name else [],
            "sheets": [sheet_name] if sheet_name else [],
        }

    def _merge_tabular_contexts(self, contexts: list[dict[str, Any]]) -> Optional[dict[str, Any]]:
        contexts = [context for context in contexts if context and context.get("rows")]
        if not contexts:
            return None
        columns: list[str] = []
        seen_columns: set[str] = set()
        source_files: list[str] = []
        seen_source_files: set[str] = set()
        sheets: list[str] = []
        seen_sheets: set[str] = set()
        rows: list[dict[str, Any]] = []
        for context in contexts:
            for column in context.get("columns") or []:
                column_name = str(column)
                if column_name not in seen_columns:
                    seen_columns.add(column_name)
                    columns.append(column_name)
            for source_file in context.get("source_files") or []:
                source_name = str(source_file)
                if source_name and source_name not in seen_source_files:
                    seen_source_files.add(source_name)
                    source_files.append(source_name)
            for sheet in context.get("sheets") or []:
                sheet_name = str(sheet)
                if sheet_name and sheet_name not in seen_sheets:
                    seen_sheets.add(sheet_name)
                    sheets.append(sheet_name)
            for row in context.get("rows") or []:
                if isinstance(row, dict):
                    rows.append(row)
        return {
            "kind": "tabular_rows",
            "columns": columns,
            "rows": rows,
            "row_count": len(rows),
            "source_files": source_files,
            "sheets": sheets,
        }

    def _load_csv_tabular_context(self, local_path: str, *, source_name: str) -> Optional[dict[str, Any]]:
        with open(local_path, "r", encoding="utf-8-sig", newline="") as handle:
            reader = csv.reader(handle)
            rows = list(reader)
        if len(rows) < 2:
            return None
        return self._tabular_context_from_rows(rows[0], rows[1:], source_name=source_name)

    def _load_xlsx_tabular_context(self, local_path: str, *, source_name: str) -> Optional[dict[str, Any]]:
        workbook = load_workbook(local_path, read_only=True, data_only=True)
        contexts: list[dict[str, Any]] = []
        try:
            for worksheet in workbook.worksheets:
                rows = list(worksheet.iter_rows(values_only=True))
                if len(rows) < 2:
                    continue
                context = self._tabular_context_from_rows(
                    list(rows[0]),
                    [list(row) for row in rows[1:]],
                    source_name=source_name,
                    sheet_name=worksheet.title,
                )
                if context.get("rows"):
                    contexts.append(context)
        finally:
            workbook.close()
        return self._merge_tabular_contexts(contexts)

    def _extraction_tabular_context(self, payload: dict[str, Any], *, source_name: str = "extraction results") -> Optional[dict[str, Any]]:
        columns = list(payload.get("columns") or [])
        rows = list(payload.get("rows") or [])
        context = self._tabular_context_from_rows(columns, rows, source_name=source_name)
        source_files = [str(item) for item in (payload.get("source_files") or []) if str(item).strip()]
        if context and source_files:
            context["source_files"] = source_files
        return self._merge_tabular_contexts([context])

    def _tabular_source_summary(self, tabular_context: dict[str, Any]) -> str:
        columns = [str(item) for item in (tabular_context.get("columns") or [])]
        source_files = [str(item) for item in (tabular_context.get("source_files") or [])]
        sample_rows = list(tabular_context.get("rows") or [])[: self.tabular_code_sample_rows]
        return (
            "Tabular source data is available to generated code as rows.\n"
            f"Rows: {int(tabular_context.get('row_count') or len(tabular_context.get('rows') or []))}\n"
            f"Columns: {', '.join(columns) if columns else '(none)'}\n"
            f"Source files: {', '.join(source_files) if source_files else '(not specified)'}\n"
            f"Sample rows: {json.dumps(sample_rows, ensure_ascii=True, default=str)}"
        )

    def _load_csv_text(self, local_path: str) -> str:
        with open(local_path, "r", encoding="utf-8-sig", newline="") as handle:
            reader = csv.reader(handle)
            rows = list(reader)
        if not rows:
            return ""
        columns = [str(item) for item in rows[0]]
        data_rows = [list(row) for row in rows[1:]]
        return self._markdown_table(columns, data_rows, len(data_rows))

    def _load_xlsx_text(self, local_path: str) -> str:
        workbook = load_workbook(local_path, read_only=True, data_only=True)
        parts: list[str] = []
        try:
            for worksheet in workbook.worksheets:
                rows = list(worksheet.iter_rows(values_only=True))
                if not rows:
                    continue
                columns = [str(item) if item is not None else "" for item in rows[0]]
                data_rows = [list(row) for row in rows[1:]]
                rendered = self._markdown_table(columns, data_rows, len(data_rows))
                if rendered:
                    parts.append(f"Sheet: {worksheet.title}\n{rendered}")
        finally:
            workbook.close()
        return "\n\n".join(parts)

    def _record_label(self, payload: dict[str, Any], index: int) -> str:
        lowered = {str(key).strip().lower(): value for key, value in payload.items()}
        for column in REPEAT_LABEL_COLUMNS:
            value = lowered.get(column)
            if value is not None and str(value).strip():
                return str(value).strip()
        for value in payload.values():
            if value is not None and str(value).strip():
                return str(value).strip()
        return f"row {index + 1}"

    def _records_from_rows(self, columns: list[Any], rows: list[Any]) -> list[dict[str, Any]]:
        safe_columns = [str(column).strip() or f"Column {index + 1}" for index, column in enumerate(columns)]
        records: list[dict[str, Any]] = []
        for row in rows:
            if not isinstance(row, (list, tuple)):
                continue
            payload = {
                column: row[column_index] if column_index < len(row) else None
                for column_index, column in enumerate(safe_columns)
            }
            if not any(value is not None and str(value).strip() for value in payload.values()):
                continue
            records.append(
                {
                    "record_index": len(records),
                    "record_label": self._record_label(payload, len(records)),
                    "record_payload": payload,
                }
            )
        return records

    def _load_csv_records(self, local_path: str) -> list[dict[str, Any]]:
        with open(local_path, "r", encoding="utf-8-sig", newline="") as handle:
            reader = csv.reader(handle)
            rows = list(reader)
        if len(rows) < 2:
            return []
        return self._records_from_rows([str(item) for item in rows[0]], rows[1:])

    def _load_xlsx_records(self, local_path: str) -> list[dict[str, Any]]:
        workbook = load_workbook(local_path, read_only=True, data_only=True)
        try:
            for worksheet in workbook.worksheets:
                rows = list(worksheet.iter_rows(values_only=True))
                if len(rows) >= 2:
                    records = self._records_from_rows(list(rows[0]), [list(row) for row in rows[1:]])
                    if records:
                        return records
        finally:
            workbook.close()
        return []

    async def _extract_repeat_records(self, run: FormFillRun) -> list[dict[str, Any]]:
        if isinstance(run.source_payload, dict) and run.source_payload.get("kind") == "extraction_result":
            return self._records_from_rows(list(run.source_payload.get("columns") or []), list(run.source_payload.get("rows") or []))

        uploaded_source_files = list(run.source_files or [])
        if len(uploaded_source_files) != 1:
            raise ValueError("Repeat mode currently supports one CSV or XLSX source file")

        source_file = uploaded_source_files[0]
        if source_file.file_type not in TABULAR_SOURCE_MIME_TYPES:
            raise ValueError("Repeat mode requires a CSV or XLSX source file")

        with tempfile.TemporaryDirectory(prefix="form_fill_records_") as temp_dir:
            local_path = os.path.join(temp_dir, f"source{_safe_ext(source_file.original_filename, '.bin')}")
            await self._download_to_local(source_file.gcs_object_name, local_path)
            if source_file.file_type in {"text/csv", "application/vnd.ms-excel"}:
                return self._load_csv_records(local_path)
            if source_file.file_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
                return self._load_xlsx_records(local_path)
        return []

    async def _build_tabular_source_context(
        self,
        run: FormFillRun,
        *,
        record_payload: Optional[dict[str, Any]] = None,
    ) -> Optional[dict[str, Any]]:
        if not self.tabular_code_enabled:
            return None

        payload = record_payload if isinstance(record_payload, dict) else None
        if payload and payload.get("kind") == "extraction_task":
            return self._extraction_tabular_context(payload, source_name="extraction task")

        if payload and payload.get("kind") == "source_file":
            file_type = str(payload.get("file_type") or "").lower()
            if file_type not in TABULAR_SOURCE_MIME_TYPES:
                return None
            filename = str(payload.get("original_filename") or "source")
            object_name = str(payload.get("gcs_object_name") or "")
            if not object_name:
                return None
            with tempfile.TemporaryDirectory(prefix="form_fill_tabular_") as temp_dir:
                local_path = os.path.join(temp_dir, f"source{_safe_ext(filename, '.bin')}")
                await self._download_to_local(object_name, local_path)
                if file_type in {"text/csv", "application/vnd.ms-excel"}:
                    return self._load_csv_tabular_context(local_path, source_name=filename)
                if file_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
                    return self._load_xlsx_tabular_context(local_path, source_name=filename)
            return None

        source_payload = getattr(run, "source_payload", None)
        if isinstance(source_payload, dict) and source_payload.get("kind") == "extraction_result":
            return self._extraction_tabular_context(source_payload, source_name="extraction results")

        uploaded_source_files = list(getattr(run, "source_files", None) or [])
        contexts: list[dict[str, Any]] = []
        for source_file in uploaded_source_files:
            file_type = str(source_file.file_type or "").lower()
            if file_type not in TABULAR_SOURCE_MIME_TYPES:
                return None
            with tempfile.TemporaryDirectory(prefix="form_fill_tabular_") as temp_dir:
                local_path = os.path.join(temp_dir, f"source{_safe_ext(source_file.original_filename, '.bin')}")
                await self._download_to_local(source_file.gcs_object_name, local_path)
                if file_type in {"text/csv", "application/vnd.ms-excel"}:
                    context = self._load_csv_tabular_context(local_path, source_name=source_file.original_filename)
                elif file_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
                    context = self._load_xlsx_tabular_context(local_path, source_name=source_file.original_filename)
                else:
                    context = None
                if context:
                    contexts.append(context)

        if contexts:
            return self._merge_tabular_contexts(contexts)

        source_gcs_object_name = getattr(run, "source_gcs_object_name", None)
        source_file_type = str(getattr(run, "source_file_type", None) or "").lower()
        if source_gcs_object_name and source_file_type in TABULAR_SOURCE_MIME_TYPES:
            file_type = source_file_type
            filename = getattr(run, "source_filename", None) or "source"
            with tempfile.TemporaryDirectory(prefix="form_fill_tabular_") as temp_dir:
                local_path = os.path.join(temp_dir, f"source{_safe_ext(filename, '.bin')}")
                await self._download_to_local(source_gcs_object_name, local_path)
                if file_type in {"text/csv", "application/vnd.ms-excel"}:
                    return self._load_csv_tabular_context(local_path, source_name=filename)
                if file_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
                    return self._load_xlsx_tabular_context(local_path, source_name=filename)

        return None

    def _source_file_units(self, run: FormFillRun) -> list[dict[str, Any]]:
        uploaded_source_files = list(run.source_files or [])
        if len(uploaded_source_files) <= 1:
            return []

        units: list[dict[str, Any]] = []
        for index, source_file in enumerate(uploaded_source_files):
            label = _safe_filename_part(Path(source_file.original_filename or "").stem, f"source-{index + 1}")
            units.append(
                {
                    "record_index": index,
                    "record_label": label,
                    "record_payload": {
                        "kind": "source_file",
                        "source_file_id": str(source_file.id),
                        "original_filename": source_file.original_filename,
                        "file_type": source_file.file_type,
                        "gcs_object_name": source_file.gcs_object_name,
                        "display_order": int(source_file.display_order if source_file.display_order is not None else index),
                    },
                }
            )
        return units

    def _extraction_task_units_from_payload(self, payload: Any) -> list[dict[str, Any]]:
        if not isinstance(payload, dict) or payload.get("kind") != "extraction_result" or payload.get("scope") != SOURCE_SCOPE_ALL:
            return []

        task_groups = payload.get("task_groups")
        if not isinstance(task_groups, list):
            return []

        units: list[dict[str, Any]] = []
        for group in task_groups:
            if not isinstance(group, dict):
                continue
            source_files = [str(item) for item in (group.get("source_files") or [])]
            label_source = Path(source_files[0]).stem if source_files else str(group.get("task_id") or "")[:8]
            label = _safe_filename_part(label_source, f"task-{len(units) + 1}")
            units.append(
                {
                    "record_index": len(units),
                    "record_label": label,
                    "record_payload": {
                        "kind": "extraction_task",
                        "task_id": str(group.get("task_id") or ""),
                        "source_files": source_files,
                        "columns": list(group.get("columns") or []),
                        "rows": list(group.get("rows") or []),
                    },
                }
            )
        return units

    def _source_units_for_run(self, run: FormFillRun) -> list[dict[str, Any]]:
        source_file_units = self._source_file_units(run)
        if source_file_units:
            return source_file_units
        return self._extraction_task_units_from_payload(run.source_payload)

    async def _build_output_source_context(
        self,
        *,
        run: FormFillRun,
        record_payload: Any,
        record_index: int,
    ) -> tuple[list[Any], str]:
        payload = record_payload if isinstance(record_payload, dict) else {}
        source_parts: list[Any] = []
        source_text_sections: list[str] = []

        if payload.get("kind") == "source_file":
            await self._append_uploaded_source_context(
                run=run,
                filename=str(payload.get("original_filename") or f"source-{record_index + 1}"),
                mime_type=str(payload.get("file_type") or ""),
                object_name=str(payload.get("gcs_object_name") or ""),
                source_key=str(payload.get("source_file_id") or f"source-{record_index + 1}"),
                display_order=int(payload.get("display_order") if payload.get("display_order") is not None else record_index),
                source_parts=source_parts,
                source_text_sections=source_text_sections,
            )
            return source_parts, "\n\n".join(item for item in source_text_sections if item)

        if payload.get("kind") == "extraction_task":
            source_files = [str(item) for item in (payload.get("source_files") or [])]
            source_text_sections.append("Extraction result source files: " + (", ".join(source_files) or "(manual rows)"))
            source_text_sections.append(
                self._markdown_table(list(payload.get("columns") or []), list(payload.get("rows") or []), len(payload.get("rows") or []))
            )
            return source_parts, "\n\n".join(item for item in source_text_sections if item)

        return source_parts, self._record_source_text({"record_payload": payload})

    def _record_source_text(self, record: dict[str, Any]) -> str:
        payload = record.get("record_payload") if isinstance(record.get("record_payload"), dict) else {}
        lines = ["Fill the target using this single source record only:"]
        for key, value in payload.items():
            lines.append(f"- {key}: {'' if value is None else value}")
        return "\n".join(lines)

    def _filled_filename(self, target_filename: str, suffix: str, extension: str) -> str:
        stem = _safe_filename_part(Path(target_filename).stem, "filled")
        safe_suffix = _safe_filename_part(suffix, "") if suffix else ""
        name = f"{stem}_{safe_suffix}_filled" if safe_suffix else f"{stem}_filled"
        return f"{name}.{extension}"

    def _ensure_client(self) -> None:
        if not self.client:
            raise RuntimeError("Form Fill AI is not configured")

    def _field_mapping_schema(self) -> types.Schema:
        return types.Schema(
            type="OBJECT",
            properties={
                "items": types.Schema(
                    type="ARRAY",
                    items=types.Schema(
                        type="OBJECT",
                        properties={
                            "name": types.Schema(type="STRING"),
                            "value": types.Schema(type="STRING", nullable=True),
                        },
                        required=["name", "value"],
                    ),
                ),
                "warnings": types.Schema(type="ARRAY", items=types.Schema(type="STRING")),
            },
            required=["items", "warnings"],
        )

    def _document_schema(self) -> types.Schema:
        return types.Schema(
            type="OBJECT",
            properties={
                "title": types.Schema(type="STRING", nullable=True),
                "content": types.Schema(type="STRING"),
                "warnings": types.Schema(type="ARRAY", items=types.Schema(type="STRING")),
            },
            required=["title", "content", "warnings"],
        )

    def _generated_code_schema(self) -> types.Schema:
        return types.Schema(
            type="OBJECT",
            properties={
                "language": types.Schema(type="STRING"),
                "code": types.Schema(type="STRING"),
                "warnings": types.Schema(type="ARRAY", items=types.Schema(type="STRING")),
                "assumptions": types.Schema(type="ARRAY", items=types.Schema(type="STRING")),
            },
            required=["language", "code", "warnings"],
        )

    def _strip_code_fence(self, code: str) -> str:
        text_value = str(code or "").strip()
        if text_value.startswith("```"):
            text_value = re.sub(r"^```(?:python)?\s*", "", text_value, flags=re.IGNORECASE)
            text_value = re.sub(r"\s*```$", "", text_value)
        return text_value.strip()

    def _build_generated_code_prompt(
        self,
        *,
        tabular_context: dict[str, Any],
        target_kind: str,
        target_context: dict[str, Any],
        output_contract: str,
        previous_code: Optional[str] = None,
        previous_error: Optional[str] = None,
    ) -> str:
        columns = list(tabular_context.get("columns") or [])
        rows = list(tabular_context.get("rows") or [])
        sample_rows = rows[: self.tabular_code_sample_rows]
        repair_block = ""
        if previous_code or previous_error:
            repair_block = f"""

Repair request:
- The previous generated code failed with this error: {previous_error or 'unknown error'}
- Return corrected complete code only in the code field.
- Previous code:
{previous_code or ''}
"""
        return f"""Generate Python code to fill a {target_kind} from tabular source rows.

Source row metadata:
- Total rows: {len(rows)}
- Columns: {json.dumps(columns, ensure_ascii=True)}
- Sample rows: {json.dumps(sample_rows, ensure_ascii=True, default=str)}

Target context:
{json.dumps(target_context, ensure_ascii=True, default=str)}

Required code shape:
- Define exactly one callable function named transform(rows, context).
- rows is a list of dictionaries containing all source rows, not just the sample above.
- context is the target context shown above.
- Return a JSON-compatible dict.
- Do not import modules.
- Do not read or write files.
- Do not use eval, exec, open, globals, locals, getattr, setattr, or dunder attributes.
- Use parse_number(value) for numeric parsing when summing currency or amount columns.
- Use as_text(value) to convert nullable values to strings.

Output contract:
{output_contract}

Rules:
- The code must process every relevant row algorithmically.
- Do not hard-code only the sample rows.
- Keep operations deterministic and minimal.
- Add ambiguous or missing data issues to warnings.
{repair_block}
"""

    def _generate_transform_code(self, contents: list[Any], *, prompt: str, label: str) -> tuple[str, list[str]]:
        payload = self._generate_json_response(contents + [prompt], self._generated_code_schema())
        language = str(payload.get("language") or GENERATED_CODE_LANGUAGE).strip().lower()
        if language not in {GENERATED_CODE_LANGUAGE, "py"}:
            raise ValueError(f"Gemini returned unsupported generated-code language '{language}'")
        code = self._strip_code_fence(str(payload.get("code") or ""))
        _validate_generated_transform_code(code)
        warnings = [str(item) for item in (payload.get("warnings") or []) if str(item).strip()]
        assumptions = [str(item) for item in (payload.get("assumptions") or []) if str(item).strip()]
        logger.info("Form Fill generated transform code (%s): chars=%s warnings=%s", label, len(code), len(warnings))
        return code, warnings + assumptions

    def _execute_generated_transform(
        self,
        code: str,
        *,
        rows: list[dict[str, Any]],
        context: dict[str, Any],
    ) -> dict[str, Any]:
        _validate_generated_transform_code(code)
        start_method = "fork" if "fork" in multiprocessing.get_all_start_methods() else "spawn"
        ctx = multiprocessing.get_context(start_method)
        result_queue = ctx.Queue(maxsize=1)
        process = ctx.Process(target=_run_generated_transform_worker, args=(code, rows, context, result_queue))
        process.start()
        try:
            message = result_queue.get(timeout=self.tabular_code_timeout_seconds)
        except queue.Empty:
            process.join(0)
            if process.is_alive():
                process.terminate()
                process.join(5)
                raise TimeoutError(f"Generated Form Fill code exceeded {self.tabular_code_timeout_seconds} seconds")
            raise ValueError("Generated Form Fill code exited without returning a result")

        process.join(5)
        if process.is_alive():
            process.terminate()
            process.join(5)
            raise TimeoutError("Generated Form Fill code returned a result but did not exit cleanly")

        if not isinstance(message, dict) or not message.get("ok"):
            raise ValueError(str((message or {}).get("error") or "Generated Form Fill code failed"))
        result = message.get("result")
        if not isinstance(result, dict):
            raise ValueError("Generated Form Fill code must return a dict")
        try:
            encoded = json.dumps(result, ensure_ascii=True, default=str)
        except Exception as exc:
            raise ValueError(f"Generated Form Fill code returned non-JSON data: {exc}") from exc
        if len(encoded.encode("utf-8")) > self.tabular_code_max_output_bytes:
            raise ValueError("Generated Form Fill code returned too much data")
        return result

    def _generated_operation_text_value(self, value: Any) -> str:
        if value is None:
            return ""
        if isinstance(value, dict):
            if "text" in value:
                return "" if value.get("text") is None else str(value.get("text"))
            if "value" in value:
                return "" if value.get("value") is None else str(value.get("value"))
            return str(value)
        if isinstance(value, (list, tuple)):
            parts = [self._generated_operation_text_value(item) for item in value]
            return " ".join(part for part in parts if part)
        return str(value)

    def _normalize_docx_edit_operation(self, operation: Any) -> Optional[dict[str, Any]]:
        if not isinstance(operation, dict):
            return None
        normalized = dict(operation)
        if not str(normalized.get("action") or "").strip() and str(normalized.get("operation") or "").strip():
            normalized["action"] = str(normalized.get("operation") or "").strip()
        if "cells" in normalized and isinstance(normalized.get("cells"), list):
            normalized["cells"] = [self._generated_operation_text_value(item) for item in normalized.get("cells") or []]
        return normalized

    def _validate_generated_transform_result(self, result: dict[str, Any], *, expected_key: str) -> dict[str, Any]:
        warnings = result.get("warnings") or []
        if not isinstance(warnings, list):
            warnings = [str(warnings)]
        normalized = {"warnings": [str(item) for item in warnings if str(item).strip()]}

        if expected_key in {"items", "operations"}:
            values = result.get(expected_key) or []
            if not isinstance(values, list):
                raise ValueError(f"Generated Form Fill code returned '{expected_key}' as a non-list")
            if expected_key == "operations":
                normalized[expected_key] = [item for item in (self._normalize_docx_edit_operation(item) for item in values) if item]
            else:
                normalized[expected_key] = [item for item in values if isinstance(item, dict)]
            return normalized

        values = result.get(expected_key) or {}
        if not isinstance(values, dict):
            raise ValueError(f"Generated Form Fill code returned '{expected_key}' as a non-object")
        normalized[expected_key] = {str(key): "" if value is None else str(value) for key, value in values.items()}
        return normalized

    def _generate_and_execute_tabular_transform(
        self,
        contents: list[Any],
        *,
        prompt: str,
        rows: list[dict[str, Any]],
        context: dict[str, Any],
        expected_key: str,
        label: str,
    ) -> dict[str, Any]:
        code = ""
        code_warnings: list[str] = []
        try:
            code, code_warnings = self._generate_transform_code(contents, prompt=prompt, label=label)
            result = self._execute_generated_transform(code, rows=rows, context=context)
            normalized = self._validate_generated_transform_result(result, expected_key=expected_key)
        except Exception as first_exc:
            logger.warning("Form Fill generated transform failed (%s), requesting repair: %s", label, first_exc)
            repair_prompt = self._build_generated_code_prompt(
                tabular_context={"columns": context.get("source_columns") or [], "rows": rows},
                target_kind=str(context.get("target_kind") or "target"),
                target_context=context,
                output_contract=str(context.get("output_contract") or "Return the requested output."),
                previous_code=code,
                previous_error=str(first_exc),
            )
            code, repair_warnings = self._generate_transform_code(contents, prompt=repair_prompt, label=f"{label}:repair")
            code_warnings.extend(repair_warnings)
            result = self._execute_generated_transform(code, rows=rows, context=context)
            normalized = self._validate_generated_transform_result(result, expected_key=expected_key)

        normalized["warnings"] = code_warnings + list(normalized.get("warnings") or [])
        normalized["code_hash"] = hashlib.sha256(code.encode("utf-8")).hexdigest()
        return normalized

    def _generate_json_response(self, contents: list[Any], schema: types.Schema) -> dict[str, Any]:
        self._ensure_client()
        response = self.client.models.generate_content(
            model=self.model_name,
            contents=contents,
            config=self._generation_config(schema),
        )
        return self._parse_response_payload(response)

    async def _append_uploaded_source_context(
        self,
        *,
        run: FormFillRun,
        filename: str,
        mime_type: str,
        object_name: str,
        source_key: str,
        display_order: int,
        source_parts: list[Any],
        source_text_sections: list[str],
    ) -> None:
        label = f"Source file {display_order + 1}: {filename}"
        normalized_mime = mime_type.lower()
        if normalized_mime in {PDF_MIME, DOCX_MIME}:
            source_object_name = object_name
            part_mime = normalized_mime
            if normalized_mime == DOCX_MIME:
                preview_object = f"form-fill/{run.user_id}/runs/{run.id}/source-previews/{source_key}.pdf"
                converter = get_document_conversion_service()
                await converter.convert_docx_gcs_to_pdf_gcs(self.storage_service, source_object_name, preview_object)
                source_object_name = preview_object
                part_mime = PDF_MIME
                source_text_sections.append(f"{label} (DOCX converted to PDF for Gemini input).")
            else:
                source_text_sections.append(f"{label} (PDF attached).")
            source_parts.append(self._part_from_uri(self.storage_service.construct_gcs_uri_for_object(source_object_name), part_mime))
            return

        with tempfile.TemporaryDirectory(prefix="form_fill_source_") as temp_dir:
            local_path = os.path.join(temp_dir, f"source-{source_key}{_safe_ext(filename, '.bin')}")
            await self._download_to_local(object_name, local_path)
            if normalized_mime in {"text/csv", "application/vnd.ms-excel"}:
                rendered = self._load_csv_text(local_path)
            elif normalized_mime == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
                rendered = self._load_xlsx_text(local_path)
            else:
                raise ValueError("Unsupported source file type")
        source_text_sections.append(f"{label}\n{rendered}".strip())

    async def _build_source_context(
        self,
        *,
        run: FormFillRun,
        source_parts: list[Any],
        source_text_sections: list[str],
    ) -> tuple[list[Any], str]:
        if isinstance(run.source_payload, dict) and run.source_payload.get("kind") == "extraction_result":
            columns = list(run.source_payload.get("columns") or [])
            rows = list(run.source_payload.get("rows") or [])
            source_files = list(run.source_payload.get("source_files") or [])
            source_text_sections.append("Extraction result source files: " + ", ".join(source_files))
            source_text_sections.append(self._markdown_table(columns, rows, len(rows)))
            return source_parts, "\n\n".join(item for item in source_text_sections if item)

        uploaded_source_files = list(run.source_files or [])
        if uploaded_source_files:
            for index, source_file in enumerate(uploaded_source_files):
                await self._append_uploaded_source_context(
                    run=run,
                    filename=source_file.original_filename,
                    mime_type=source_file.file_type,
                    object_name=source_file.gcs_object_name,
                    source_key=str(source_file.id),
                    display_order=int(source_file.display_order if source_file.display_order is not None else index),
                    source_parts=source_parts,
                    source_text_sections=source_text_sections,
                )
            return source_parts, "\n\n".join(item for item in source_text_sections if item)

        if not run.source_gcs_object_name or not run.source_file_type:
            raise ValueError("Form Fill source is missing")

        await self._append_uploaded_source_context(
            run=run,
            filename=run.source_filename or "source",
            mime_type=run.source_file_type,
            object_name=run.source_gcs_object_name,
            source_key="legacy-source",
            display_order=0,
            source_parts=source_parts,
            source_text_sections=source_text_sections,
        )
        return source_parts, "\n\n".join(item for item in source_text_sections if item)

    def _replace_text_in_paragraph(self, paragraph: Any, replacements: dict[str, str]) -> None:
        text = paragraph.text
        if not text:
            return
        updated = text
        for placeholder, value in replacements.items():
            updated = updated.replace(placeholder, value)
        if updated != text:
            paragraph.text = updated

    def _apply_docx_placeholders(self, local_target_path: str, replacements: dict[str, str], output_path: str) -> None:
        from docx import Document as DocxDocument

        doc = DocxDocument(local_target_path)
        for paragraph in doc.paragraphs:
            self._replace_text_in_paragraph(paragraph, replacements)
        for table in doc.tables:
            for row in table.rows:
                for cell in self._docx_row_cells(row):
                    for paragraph in cell.paragraphs:
                        self._replace_text_in_paragraph(paragraph, replacements)
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                self._replace_text_in_paragraph(paragraph, replacements)
            for paragraph in section.footer.paragraphs:
                self._replace_text_in_paragraph(paragraph, replacements)
        doc.save(output_path)

    def _apply_fillable_pdf(self, local_target_path: str, field_values: dict[str, str], output_path: str) -> None:
        from PyPDF2 import PdfReader, PdfWriter
        from PyPDF2.generic import BooleanObject, NameObject

        reader = PdfReader(local_target_path)
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)

        for page in writer.pages:
            writer.update_page_form_field_values(page, field_values)

        try:
            if "/AcroForm" in reader.trailer["/Root"]:
                writer._root_object.update({NameObject("/AcroForm"): reader.trailer["/Root"]["/AcroForm"]})
                writer._root_object[NameObject("/AcroForm")].update({NameObject("/NeedAppearances"): BooleanObject(True)})
        except Exception as exc:
            logger.warning("Failed to copy AcroForm metadata: %s", exc)

        with open(output_path, "wb") as handle:
            writer.write(handle)

    def _pdf_overlay_schema(self) -> types.Schema:
        return types.Schema(
            type="OBJECT",
            properties={
                "items": types.Schema(
                    type="ARRAY",
                    items=types.Schema(
                        type="OBJECT",
                        properties={
                            "page_number": types.Schema(type="INTEGER", nullable=True),
                            "anchor_text": types.Schema(type="STRING", nullable=True),
                            "anchor_before": types.Schema(type="STRING", nullable=True),
                            "anchor_after": types.Schema(type="STRING", nullable=True),
                            "overlay_text": types.Schema(type="STRING"),
                            "placement_hint": types.Schema(type="STRING", nullable=True),
                            "cover_anchor": types.Schema(type="BOOLEAN", nullable=True),
                            "font_size": types.Schema(type="NUMBER", nullable=True),
                        },
                        required=["page_number", "overlay_text"],
                    ),
                ),
                "warnings": types.Schema(type="ARRAY", items=types.Schema(type="STRING")),
            },
            required=["items", "warnings"],
        )

    def _docx_edit_schema(self) -> types.Schema:
        return types.Schema(
            type="OBJECT",
            properties={
                "operations": types.Schema(
                    type="ARRAY",
                    items=types.Schema(
                        type="OBJECT",
                        properties={
                            "action": types.Schema(type="STRING"),
                            "block_id": types.Schema(type="STRING", nullable=True),
                            "table_id": types.Schema(type="STRING", nullable=True),
                            "find_text": types.Schema(type="STRING", nullable=True),
                            "text": types.Schema(type="STRING", nullable=True),
                            "row_index": types.Schema(type="INTEGER", nullable=True),
                            "column_index": types.Schema(type="INTEGER", nullable=True),
                            "cells": types.Schema(type="ARRAY", items=types.Schema(type="STRING"), nullable=True),
                        },
                        required=["action"],
                    ),
                ),
                "warnings": types.Schema(type="ARRAY", items=types.Schema(type="STRING")),
            },
            required=["operations", "warnings"],
        )

    def _build_pdf_overlay_prompt(self, *, source_text: str, target_preview_text: str) -> str:
        source_block = source_text.strip() or "The source is provided as attached document(s)."
        target_block = target_preview_text.strip() or "No preview text was available from the target PDF."
        return f"""You are filling the provided PDF by overlaying text onto the original pages.

Source material summary:
{source_block}

Target PDF text preview:
{target_block}

Instructions:
- Return overlay items only for content that should be added to the original PDF.
- Use the original page numbers from the target PDF.
- Prefer anchor_text that appears visibly in the PDF near where the overlay should go.
- Use placement_hint values such as replace_anchor, right_of, below, or near_blank.
- Set cover_anchor to true when replacing placeholder text, blanks, or underscores already present in the PDF.
- overlay_text must be concise and ready to render.
- Do not summarize, collapse, or omit entries because there are many; continuation will request more entries when needed.
- Do not write "see attached", "too many transactions", or any output-limit workaround.
- Do not warn that only the first N entries were inserted. Do not ask the user to request continuation.
- Do not mention output limits or token limits.
- Add ambiguities or missing values to warnings.
"""

    def _build_docx_edit_prompt(
        self,
        *,
        source_text: str,
        block_summary: str,
        table_summary: str,
        target_preview_text: str,
        allow_table_expansion: bool,
        restrict_to_table_expansion: bool = False,
    ) -> str:
        source_block = source_text.strip() or "The source is provided as attached document(s)."
        target_block = target_preview_text.strip() or "No preview text was available from the target DOCX."
        block_section = block_summary.strip() or "No editable paragraph blocks were found."
        table_section = table_summary.strip() or "No editable tables were found."
        table_instruction = (
            "You may use insert_table_row_after or insert_table_column_after when the existing table needs more space.\n"
            "- For insert_table_row_after, provide table_id, row_index, and cells with one value per existing column.\n"
            "- For insert_table_column_after, provide table_id, column_index, and cells with one value per existing row.\n"
            "- Prefer adding rows or columns only when the source clearly contains more data than the current table can hold."
            if allow_table_expansion
            else "Do not add rows or columns to any table. If more space is needed, explain that in warnings."
        )
        scope_instruction = (
            "Return only table expansion operations. Do not use paragraph edit actions in this pass."
            if restrict_to_table_expansion
            else "Return only operations against the provided block_id and table_id values."
        )
        return f"""You are editing the provided DOCX in place.

Source material summary:
{source_block}

Target DOCX preview:
{target_block}

Editable blocks:
{block_section}

Editable tables:
{table_section}

Instructions:
- {scope_instruction}
- Prefer replace_text_in_block when a specific phrase inside a block should change.
- Put the operation name in the action field. Do not use operation as the key name.
- Use replace_block_text when the entire block should be rewritten.
- Use insert_after_block or insert_before_block for new paragraphs adjacent to an existing block.
- Use append_to_block only for short additions to an existing block.
- {table_instruction}
- For tabular targets, emit one insert_table_row_after operation per source row that belongs in the table.
- Calculate totals, subtotals, recapitulations, and other numeric rollups from the provided source rows when numeric amounts are present.
- Do not summarize, collapse, or omit source rows because there are many; continuation will request more operations when needed.
- Do not write "see attached", "too many transactions", or any output-limit workaround.
- Do not insert only the first N rows and then warn that continuation is needed. Return the next concrete operations for this request; the system will continue automatically.
- Do not claim totals cannot be calculated due to transaction volume.
- Do not mention output limits or token limits.
- Keep operations minimal and deterministic.
- Add any ambiguities or low-confidence choices to warnings.
"""

    def _page_texts_with_numbers(self, local_path: str, max_chars: int = 24000) -> str:
        doc = fitz.open(local_path)
        pages: list[str] = []
        try:
            for page_number, page in enumerate(doc, start=1):
                text = (page.get_text("text") or "").strip()
                pages.append(f"Page {page_number}:\n{text or '[no searchable text]'}")
                combined = "\n\n".join(pages)
                if len(combined) >= max_chars:
                    return combined[:max_chars]
        finally:
            doc.close()
        return "\n\n".join(pages)[:max_chars]

    def _collect_docx_blocks(self, doc: Any) -> tuple[list[dict[str, Any]], dict[str, Any]]:
        blocks: list[dict[str, Any]] = []
        block_map: dict[str, Any] = {}

        def add_block(block_id: str, paragraph: Any, location: str) -> None:
            text = (paragraph.text or "").strip()
            display_text = text if text else "[blank]"
            block = {
                "block_id": block_id,
                "location": location,
                "text": display_text,
                "paragraph": paragraph,
            }
            blocks.append(block)
            block_map[block_id] = paragraph

        for paragraph_index, paragraph in enumerate(doc.paragraphs):
            add_block(f"body.paragraph.{paragraph_index}", paragraph, f"body paragraph {paragraph_index}")

        for table_index, table in enumerate(doc.tables):
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(self._docx_row_cells(row)):
                    for paragraph_index, paragraph in enumerate(cell.paragraphs):
                        add_block(
                            f"table.{table_index}.row.{row_index}.cell.{cell_index}.paragraph.{paragraph_index}",
                            paragraph,
                            f"table {table_index} row {row_index} cell {cell_index} paragraph {paragraph_index}",
                        )

        for section_index, section in enumerate(doc.sections):
            for paragraph_index, paragraph in enumerate(section.header.paragraphs):
                add_block(
                    f"header.section.{section_index}.paragraph.{paragraph_index}",
                    paragraph,
                    f"header section {section_index} paragraph {paragraph_index}",
                )
            for paragraph_index, paragraph in enumerate(section.footer.paragraphs):
                add_block(
                    f"footer.section.{section_index}.paragraph.{paragraph_index}",
                    paragraph,
                    f"footer section {section_index} paragraph {paragraph_index}",
                )

        return blocks, block_map

    def _collect_docx_tables(self, doc: Any) -> tuple[list[dict[str, Any]], dict[str, Any]]:
        tables: list[dict[str, Any]] = []
        table_map: dict[str, Any] = {}

        for table_index, table in enumerate(doc.tables):
            table_id = f"table.{table_index}"
            row_count = len(table.rows)
            column_count = max((len(self._docx_row_cells(row)) for row in table.rows), default=0)
            preview_rows: list[str] = []
            for row_index, row in enumerate(table.rows[:8]):
                values = [cell.text.strip() for cell in self._docx_row_cells(row)[:8]]
                rendered = " | ".join(value if value else "[blank]" for value in values)
                preview_rows.append(f"row {row_index}: {rendered or '[blank]'}")
            table_info = {
                "table_id": table_id,
                "rows": row_count,
                "columns": column_count,
                "preview_rows": preview_rows,
                "table": table,
            }
            tables.append(table_info)
            table_map[table_id] = table

        return tables, table_map

    def _summarize_docx_blocks(self, blocks: list[dict[str, Any]]) -> str:
        lines: list[str] = []
        for block in blocks:
            lines.append(f"{block['block_id']} | {block['location']} | {block['text']}")
            combined = "\n".join(lines)
            if len(combined) >= DOCX_BLOCK_TEXT_LIMIT:
                return combined[:DOCX_BLOCK_TEXT_LIMIT]
        return "\n".join(lines)[:DOCX_BLOCK_TEXT_LIMIT]

    def _summarize_docx_tables(self, tables: list[dict[str, Any]]) -> str:
        lines: list[str] = []
        for table in tables:
            lines.append(f"{table['table_id']} | rows={table['rows']} | columns={table['columns']}")
            for preview_row in table.get("preview_rows") or []:
                lines.append(f"  {preview_row}")
            combined = "\n".join(lines)
            if len(combined) >= DOCX_BLOCK_TEXT_LIMIT:
                return combined[:DOCX_BLOCK_TEXT_LIMIT]
        return "\n".join(lines)[:DOCX_BLOCK_TEXT_LIMIT]

    def _public_docx_blocks(self, blocks: list[dict[str, Any]]) -> list[dict[str, Any]]:
        return [
            {
                "block_id": str(block.get("block_id") or ""),
                "location": str(block.get("location") or ""),
                "text": str(block.get("text") or ""),
            }
            for block in blocks
        ]

    def _public_docx_tables(self, tables: list[dict[str, Any]]) -> list[dict[str, Any]]:
        return [
            {
                "table_id": str(table.get("table_id") or ""),
                "rows": int(table.get("rows") or 0),
                "columns": int(table.get("columns") or 0),
                "preview_rows": [str(item) for item in (table.get("preview_rows") or [])],
            }
            for table in tables
        ]

    def _resolve_pdf_anchor_rect(self, page: fitz.Page, item: dict[str, Any]) -> fitz.Rect | None:
        anchor_text = str(item.get("anchor_text") or "").strip()
        anchor_before = str(item.get("anchor_before") or "").strip()
        anchor_after = str(item.get("anchor_after") or "").strip()

        if anchor_text:
            matches = page.search_for(anchor_text)
            if matches:
                return matches[0]

        before_rect = None
        after_rect = None
        if anchor_before:
            before_matches = page.search_for(anchor_before)
            if before_matches:
                before_rect = before_matches[0]
        if anchor_after:
            after_matches = page.search_for(anchor_after)
            if after_matches:
                after_rect = after_matches[0]

        if before_rect and after_rect:
            x0 = min(before_rect.x1, after_rect.x0)
            x1 = max(before_rect.x1, after_rect.x0)
            y0 = min(before_rect.y0, after_rect.y0)
            y1 = max(before_rect.y1, after_rect.y1)
            return fitz.Rect(x0, y0, x1, y1)

        return before_rect or after_rect

    def _target_rect_from_anchor(
        self,
        page: fitz.Page,
        anchor_rect: fitz.Rect | None,
        item: dict[str, Any],
        fallback_index: int,
    ) -> fitz.Rect:
        placement_hint = str(item.get("placement_hint") or "near_blank").strip().lower()
        font_size = float(item.get("font_size") or 10.0)
        margin = 24.0
        page_rect = page.rect
        line_height = max(font_size * 1.8, 18.0)

        if anchor_rect is None:
            y0 = margin + (fallback_index * (line_height + 6.0))
            return fitz.Rect(margin, y0, page_rect.width - margin, y0 + line_height)

        expanded = fitz.Rect(anchor_rect.x0 - 1.5, anchor_rect.y0 - 1.5, anchor_rect.x1 + 1.5, anchor_rect.y1 + 1.5)
        if placement_hint == "replace_anchor":
            return fitz.Rect(
                max(margin, expanded.x0 - 4.0),
                max(margin, expanded.y0 - 2.0),
                min(page_rect.width - margin, max(expanded.x1 + 120.0, expanded.x0 + 180.0)),
                min(page_rect.height - margin, expanded.y1 + 4.0),
            )
        if placement_hint == "below":
            y0 = expanded.y1 + 4.0
            return fitz.Rect(
                max(margin, expanded.x0),
                y0,
                min(page_rect.width - margin, max(expanded.x1 + 140.0, expanded.x0 + 220.0)),
                min(page_rect.height - margin, y0 + line_height + 8.0),
            )

        x0 = min(page_rect.width - margin - 10.0, expanded.x1 + 6.0)
        x1 = min(page_rect.width - margin, max(x0 + 180.0, expanded.x1 + 220.0))
        if x1 <= x0 + 30.0:
            x0 = max(margin, expanded.x0)
            x1 = min(page_rect.width - margin, max(expanded.x1 + 140.0, expanded.x0 + 220.0))
            y0 = expanded.y1 + 4.0
            return fitz.Rect(x0, y0, x1, min(page_rect.height - margin, y0 + line_height + 8.0))
        return fitz.Rect(x0, max(margin, expanded.y0 - 2.0), x1, min(page_rect.height - margin, expanded.y1 + 4.0))

    def _apply_pdf_overlay_plan(
        self,
        local_target_path: str,
        overlay_items: list[dict[str, Any]],
        output_path: str,
    ) -> list[str]:
        warnings: list[str] = []
        doc = fitz.open(local_target_path)
        try:
            fallback_slots: dict[int, int] = {}
            for item in overlay_items:
                overlay_text = str(item.get("overlay_text") or "").strip()
                if not overlay_text:
                    continue

                page_number = int(item.get("page_number") or 1)
                if page_number < 1 or page_number > len(doc):
                    warnings.append(f"Overlay item referenced missing page {page_number}.")
                    continue

                page = doc[page_number - 1]
                anchor_rect = self._resolve_pdf_anchor_rect(page, item)
                if anchor_rect is None:
                    warnings.append(
                        f"Could not locate anchor for overlay '{overlay_text[:40]}' on page {page_number}; placed it in a fallback position."
                    )
                fallback_index = fallback_slots.get(page_number, 0)
                target_rect = self._target_rect_from_anchor(page, anchor_rect, item, fallback_index)
                fallback_slots[page_number] = fallback_index + 1

                placement_hint = str(item.get("placement_hint") or "near_blank").strip().lower()
                cover_anchor = bool(item.get("cover_anchor")) or placement_hint == "replace_anchor"
                if cover_anchor and anchor_rect is not None:
                    page.draw_rect(
                        fitz.Rect(anchor_rect.x0 - 1.5, anchor_rect.y0 - 1.5, anchor_rect.x1 + 1.5, anchor_rect.y1 + 1.5),
                        color=(1, 1, 1),
                        fill=(1, 1, 1),
                        overlay=True,
                    )

                font_size = float(item.get("font_size") or 10.0)
                page.insert_textbox(
                    target_rect,
                    overlay_text,
                    fontsize=font_size,
                    fontname="helv",
                    color=(0, 0, 0),
                    align=fitz.TEXT_ALIGN_LEFT,
                    overlay=True,
                )

            doc.save(output_path)
        finally:
            doc.close()
        return warnings

    def _replace_text_in_docx_block(self, paragraph: Any, find_text: str, replacement_text: str) -> bool:
        needle = (find_text or "").strip()
        if not needle:
            return False
        for run in paragraph.runs:
            if needle in run.text:
                run.text = run.text.replace(needle, replacement_text, 1)
                return True

        full_text = paragraph.text or ""
        if needle not in full_text:
            return False

        paragraph.text = full_text.replace(needle, replacement_text, 1)
        return True

    def _insert_paragraph_after(self, paragraph: Any, text: str) -> Any:
        from docx.oxml import OxmlElement
        from docx.text.paragraph import Paragraph

        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_paragraph = Paragraph(new_p, paragraph._parent)
        if getattr(paragraph, "style", None) is not None:
            new_paragraph.style = paragraph.style
        new_paragraph.text = text
        return new_paragraph

    def _clear_docx_cell(self, cell: Any) -> None:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.text = ""
            if not paragraph.runs:
                paragraph.text = ""

    def _insert_table_row_after(self, table: Any, row_index: int, values: list[str]) -> list[str]:
        warnings: list[str] = []
        if row_index < 0 or row_index >= len(table.rows):
            return [f"Could not insert table row after index {row_index}; the row does not exist."]

        template_row = table.rows[row_index]
        new_tr = deepcopy(template_row._tr)
        template_row._tr.addnext(new_tr)

        inserted_row = table.rows[row_index + 1]
        logical_cells = self._docx_row_cells(inserted_row)
        for cell in logical_cells:
            self._clear_docx_cell(cell)

        if len(values) > len(logical_cells):
            warnings.append(
                f"Inserted row after {row_index} received {len(values)} values, but the table only has {len(logical_cells)} columns."
            )

        for cell_index, cell in enumerate(logical_cells):
            cell.text = values[cell_index] if cell_index < len(values) else ""
        return warnings

    def _insert_table_column_after(self, table: Any, column_index: int, values: list[str]) -> list[str]:
        warnings: list[str] = []
        if not table.rows:
            return ["Could not insert a table column because the table has no rows."]
        if self._docx_table_has_merged_cells(table):
            return ["Could not insert a table column because the table contains merged cells."]
        if column_index < 0 or column_index >= len(table.rows[0].cells):
            return [f"Could not insert table column after index {column_index}; the column does not exist."]

        if len(values) > len(table.rows):
            warnings.append(
                f"Inserted column after {column_index} received {len(values)} values, but the table only has {len(table.rows)} rows."
            )

        for row in table.rows:
            template_cell = row.cells[column_index]
            new_tc = deepcopy(template_cell._tc)
            template_cell._tc.addnext(new_tc)

        for row_index, row in enumerate(table.rows):
            inserted_cell = row.cells[column_index + 1]
            self._clear_docx_cell(inserted_cell)
            inserted_cell.text = values[row_index] if row_index < len(values) else ""

        return warnings

    def _apply_docx_edit_plan(
        self,
        local_target_path: str,
        operations: list[dict[str, Any]],
        output_path: str,
        allow_table_expansion: bool = False,
    ) -> list[str]:
        from docx import Document as DocxDocument

        warnings: list[str] = []
        doc = DocxDocument(local_target_path)
        _blocks, block_map = self._collect_docx_blocks(doc)
        _tables, table_map = self._collect_docx_tables(doc)

        for raw_operation in operations:
            operation = self._normalize_docx_edit_operation(raw_operation)
            if not operation:
                continue
            action = str(operation.get("action") or "").strip()
            block_id = str(operation.get("block_id") or "").strip()
            text = str(operation.get("text") or "")
            if action in {"insert_table_row_after", "insert_table_column_after"}:
                if not allow_table_expansion:
                    warnings.append(f"Ignored DOCX table expansion action '{action}' because it was not permitted for this run.")
                    continue
                table_id = str(operation.get("table_id") or "").strip()
                table = table_map.get(table_id)
                if table is None:
                    warnings.append(f"Could not resolve DOCX table '{table_id}'.")
                    continue

                cells = operation.get("cells") or []
                values = [str(item) for item in cells if item is not None]
                if action == "insert_table_row_after":
                    row_index = operation.get("row_index")
                    if not isinstance(row_index, int):
                        warnings.append(f"DOCX table row insertion for '{table_id}' is missing a valid row_index.")
                        continue
                    warnings.extend(self._insert_table_row_after(table, row_index, values))
                else:
                    column_index = operation.get("column_index")
                    if not isinstance(column_index, int):
                        warnings.append(f"DOCX table column insertion for '{table_id}' is missing a valid column_index.")
                        continue
                    warnings.extend(self._insert_table_column_after(table, column_index, values))
                continue

            target = block_map.get(block_id)
            if target is None:
                warnings.append(f"Could not resolve DOCX block '{block_id}'.")
                continue

            if action == "replace_text_in_block":
                find_text = str(operation.get("find_text") or "")
                if not self._replace_text_in_docx_block(target, find_text, text):
                    warnings.append(f"Could not find '{find_text}' inside DOCX block '{block_id}'.")
            elif action == "replace_block_text":
                target.text = text
            elif action == "append_to_block":
                existing = target.text or ""
                target.text = f"{existing} {text}".strip() if existing else text
            elif action == "insert_before_block":
                inserted = target.insert_paragraph_before(text)
                if getattr(target, "style", None) is not None:
                    inserted.style = target.style
            elif action == "insert_after_block":
                self._insert_paragraph_after(target, text)
            else:
                warnings.append(f"Unsupported DOCX edit action '{action}'.")

        doc.save(output_path)
        return warnings

    def _build_mapping_prompt(
        self,
        *,
        source_text: str,
        mapping_items: list[str],
        mapping_label: str,
        target_hint: str,
    ) -> str:
        source_block = source_text.strip() or "The source is provided as attached document(s)."
        names = "\n".join(f"- {item}" for item in mapping_items)
        return f"""You are filling a {target_hint} using the provided source material.

Source material summary:
{source_block}

{mapping_label}:
{names}

Instructions:
- Return one item for every provided name.
- Keep the original name exactly as given.
- Use null when the source does not clearly provide a value.
- Return concise values suitable for direct insertion into the target document.
- Add any important caveats to warnings.
"""

    def _generate_mapping_payload(
        self,
        contents: list[Any],
        *,
        source_text: str,
        mapping_items: list[str],
        mapping_label: str,
        target_hint: str,
        label: str,
    ) -> dict[str, Any]:
        all_items: list[dict[str, Any]] = []
        all_warnings: list[str] = []
        chunk_size = max(1, int(self.mapping_chunk_size or self.batch_items_per_call or 1000))

        for start in range(0, len(mapping_items), chunk_size):
            chunk = mapping_items[start : start + chunk_size]
            prompt = self._build_mapping_prompt(
                source_text=source_text,
                mapping_items=chunk,
                mapping_label=mapping_label,
                target_hint=target_hint,
            )
            payload = self._generate_collection_json_response(
                contents,
                prompt=prompt,
                schema=self._field_mapping_schema(),
                collection_key="items",
                label=f"{label}:chunk_{(start // chunk_size) + 1}",
            )
            allowed_names = {str(item) for item in chunk}
            all_items.extend(
                item
                for item in (payload.get("items") or [])
                if isinstance(item, dict) and str(item.get("name") or "") in allowed_names
            )
            all_warnings.extend([str(item) for item in (payload.get("warnings") or []) if str(item).strip()])

        return {"items": all_items, "warnings": all_warnings}

    def _compact_fill_plan(self, fill_plan: dict[str, Any]) -> dict[str, Any]:
        compact: dict[str, Any] = {}
        for key, value in (fill_plan or {}).items():
            if key in {"items", "operations", "table_operations"} and isinstance(value, list):
                compact[f"{key}_count"] = len(value)
                compact[f"{key}_sample"] = value[:20]
            elif key in {"field_values", "replacements"} and isinstance(value, dict):
                compact[f"{key}_count"] = len(value)
                compact[f"{key}_sample"] = dict(list(value.items())[:50])
            else:
                compact[key] = value
        return compact

    async def _generate_filled_document(
        self,
        *,
        run: FormFillRun,
        temp_dir: str,
        target_local_path: str,
        source_parts: list[Any],
        source_text: str,
        tabular_context: Optional[dict[str, Any]] = None,
        filename_suffix: str = "",
    ) -> dict[str, Any]:
        target_parts: list[Any] = []
        warnings: list[str] = []
        fill_plan: dict[str, Any] = {}
        tabular_rows = list((tabular_context or {}).get("rows") or [])

        if run.target_file_type == PDF_MIME:
            pdf_fields = self._extract_pdf_form_fields(target_local_path)
            target_parts.append(self._part_from_uri(self.storage_service.construct_gcs_uri_for_object(run.target_gcs_object_name), PDF_MIME))
            if pdf_fields:
                processing_strategy = "fillable_pdf"
                if tabular_rows:
                    processing_strategy = "fillable_pdf_generated_code"
                    output_contract = (
                        "Return {'field_values': {field_name: value, ...}, 'warnings': [...]} using only the provided field names. "
                        "Values must be strings ready for direct insertion."
                    )
                    code_context = {
                        "target_kind": "fillable PDF",
                        "field_names": pdf_fields,
                        "source_columns": list((tabular_context or {}).get("columns") or []),
                        "source_files": list((tabular_context or {}).get("source_files") or []),
                        "output_contract": output_contract,
                    }
                    code_prompt = self._build_generated_code_prompt(
                        tabular_context=tabular_context or {},
                        target_kind="fillable PDF",
                        target_context=code_context,
                        output_contract=output_contract,
                    )
                    mapping_payload = self._generate_and_execute_tabular_transform(
                        source_parts + target_parts,
                        prompt=code_prompt,
                        rows=tabular_rows,
                        context=code_context,
                        expected_key="field_values",
                        label="fillable_pdf_generated_code",
                    )
                    allowed_fields = set(pdf_fields)
                    field_values = {
                        str(name): str(value)
                        for name, value in (mapping_payload.get("field_values") or {}).items()
                        if str(name) in allowed_fields
                    }
                    warnings.extend([str(item) for item in (mapping_payload.get("warnings") or []) if str(item).strip()])
                    fill_plan = {
                        "strategy": processing_strategy,
                        "field_values": field_values,
                        "source_rows": len(tabular_rows),
                        "code_hash": mapping_payload.get("code_hash"),
                    }
                else:
                    mapping_payload = self._generate_mapping_payload(
                        source_parts + target_parts,
                        source_text=source_text,
                        mapping_items=pdf_fields,
                        mapping_label="Fillable PDF field names",
                        target_hint="fillable PDF form",
                        label="fillable_pdf_mapping",
                    )
                    field_values = {
                        str(item.get("name")): "" if item.get("value") is None else str(item.get("value"))
                        for item in mapping_payload.get("items") or []
                        if isinstance(item, dict) and item.get("name")
                    }
                    warnings.extend([str(item) for item in (mapping_payload.get("warnings") or []) if str(item).strip()])
                    fill_plan = {"strategy": processing_strategy, "field_values": field_values}
                output_pdf_path = os.path.join(temp_dir, f"filled-{uuid.uuid4()}.pdf")
                self._apply_fillable_pdf(target_local_path, field_values, output_pdf_path)
                return {
                    "local_path": output_pdf_path,
                    "filename": self._filled_filename(run.target_filename, filename_suffix, "pdf"),
                    "mime_type": PDF_MIME,
                    "strategy": processing_strategy,
                    "warnings": warnings,
                    "fill_plan": self._compact_fill_plan(fill_plan),
                }

            processing_strategy = "pdf_overlay"
            target_preview_text = self._page_texts_with_numbers(target_local_path)
            if tabular_rows:
                processing_strategy = "pdf_overlay_generated_code"
                output_contract = (
                    "Return {'items': [overlay_item, ...], 'warnings': [...]}. Each overlay_item must include "
                    "page_number and overlay_text, and may include anchor_text, anchor_before, anchor_after, "
                    "placement_hint, cover_anchor, and font_size."
                )
                code_context = {
                    "target_kind": "PDF overlay",
                    "target_preview_text": target_preview_text,
                    "source_columns": list((tabular_context or {}).get("columns") or []),
                    "source_files": list((tabular_context or {}).get("source_files") or []),
                    "output_contract": output_contract,
                }
                overlay_prompt = self._build_generated_code_prompt(
                    tabular_context=tabular_context or {},
                    target_kind="PDF overlay",
                    target_context=code_context,
                    output_contract=output_contract,
                )
                overlay_payload = self._generate_and_execute_tabular_transform(
                    source_parts + target_parts,
                    prompt=overlay_prompt,
                    rows=tabular_rows,
                    context=code_context,
                    expected_key="items",
                    label="pdf_overlay_generated_code",
                )
            else:
                overlay_prompt = self._build_pdf_overlay_prompt(
                    source_text=source_text,
                    target_preview_text=target_preview_text,
                )
                overlay_payload = self._generate_collection_json_response(
                    source_parts + target_parts,
                    prompt=overlay_prompt,
                    schema=self._pdf_overlay_schema(),
                    collection_key="items",
                    label="pdf_overlay",
                    continue_on_full_batch=True,
                )
            overlay_items = [
                item for item in (overlay_payload.get("items") or []) if isinstance(item, dict) and str(item.get("overlay_text") or "").strip()
            ]
            warnings.extend([str(item) for item in (overlay_payload.get("warnings") or []) if str(item).strip()])
            fill_plan = {
                "strategy": processing_strategy,
                "items": overlay_items,
                "source_rows": len(tabular_rows) if tabular_rows else None,
                "code_hash": overlay_payload.get("code_hash"),
            }
            final_local_path = os.path.join(temp_dir, f"filled-{uuid.uuid4()}.pdf")
            warnings.extend(self._apply_pdf_overlay_plan(target_local_path, overlay_items, final_local_path))
            warnings.append("The target PDF was not fillable, so Form Fill applied text overlays onto the original PDF.")
            return {
                "local_path": final_local_path,
                "filename": self._filled_filename(run.target_filename, filename_suffix, "pdf"),
                "mime_type": PDF_MIME,
                "strategy": processing_strategy,
                "warnings": warnings,
                "fill_plan": self._compact_fill_plan(fill_plan),
            }

        if run.target_file_type == DOCX_MIME:
            converter = get_document_conversion_service()
            preview_object = f"form-fill/{run.user_id}/runs/{run.id}/target-preview.pdf"
            await converter.convert_docx_gcs_to_pdf_gcs(self.storage_service, run.target_gcs_object_name, preview_object)
            target_parts.append(self._part_from_uri(self.storage_service.construct_gcs_uri_for_object(preview_object), PDF_MIME))
            placeholders = self._extract_docx_placeholders(target_local_path)
            if placeholders:
                processing_strategy = "docx_placeholders"
                if tabular_rows:
                    processing_strategy = "docx_placeholders_generated_code"
                    output_contract = (
                        "Return {'replacements': {placeholder: value, ...}, 'warnings': [...]} using only the provided placeholders. "
                        "Values must be strings ready for direct insertion."
                    )
                    code_context = {
                        "target_kind": "DOCX placeholders",
                        "placeholders": placeholders,
                        "target_preview_text": self._extract_docx_text(target_local_path),
                        "source_columns": list((tabular_context or {}).get("columns") or []),
                        "source_files": list((tabular_context or {}).get("source_files") or []),
                        "output_contract": output_contract,
                    }
                    code_prompt = self._build_generated_code_prompt(
                        tabular_context=tabular_context or {},
                        target_kind="DOCX placeholders",
                        target_context=code_context,
                        output_contract=output_contract,
                    )
                    mapping_payload = self._generate_and_execute_tabular_transform(
                        source_parts + target_parts,
                        prompt=code_prompt,
                        rows=tabular_rows,
                        context=code_context,
                        expected_key="replacements",
                        label="docx_placeholders_generated_code",
                    )
                    allowed_placeholders = set(placeholders)
                    replacements = {
                        str(name): str(value)
                        for name, value in (mapping_payload.get("replacements") or {}).items()
                        if str(name) in allowed_placeholders
                    }
                    warnings.extend([str(item) for item in (mapping_payload.get("warnings") or []) if str(item).strip()])
                else:
                    mapping_payload = self._generate_mapping_payload(
                        source_parts + target_parts,
                        source_text=source_text,
                        mapping_items=placeholders,
                        mapping_label="DOCX placeholders",
                        target_hint="DOCX template",
                        label="docx_placeholder_mapping",
                    )
                    replacements = {
                        str(item.get("name")): "" if item.get("value") is None else str(item.get("value"))
                        for item in mapping_payload.get("items") or []
                        if isinstance(item, dict) and item.get("name")
                    }
                    warnings.extend([str(item) for item in (mapping_payload.get("warnings") or []) if str(item).strip()])
                fill_plan = {
                    "strategy": processing_strategy,
                    "replacements": replacements,
                    "allow_docx_table_expansion": bool(run.allow_docx_table_expansion),
                    "source_rows": len(tabular_rows) if tabular_rows else None,
                    "code_hash": mapping_payload.get("code_hash") if tabular_rows else None,
                }
                output_docx_path = os.path.join(temp_dir, f"filled-{uuid.uuid4()}.docx")
                self._apply_docx_placeholders(target_local_path, replacements, output_docx_path)
                final_local_path = output_docx_path

                if run.allow_docx_table_expansion:
                    from docx import Document as DocxDocument

                    placeholder_doc = DocxDocument(output_docx_path)
                    docx_blocks, _block_map = self._collect_docx_blocks(placeholder_doc)
                    docx_tables, _table_map = self._collect_docx_tables(placeholder_doc)
                    table_operations: list[dict[str, Any]] = []
                    if docx_tables:
                        edit_prompt = self._build_docx_edit_prompt(
                            source_text=source_text,
                            block_summary=self._summarize_docx_blocks(docx_blocks),
                            table_summary=self._summarize_docx_tables(docx_tables),
                            target_preview_text=self._extract_docx_text(output_docx_path),
                            allow_table_expansion=True,
                            restrict_to_table_expansion=True,
                        )
                        if tabular_rows:
                            output_contract = (
                                "Return {'operations': [operation, ...], 'warnings': [...]}. Operations must use only "
                                "insert_table_row_after or insert_table_column_after and valid table_id/row_index/column_index values. "
                                "Each operation object must put the operation name in the 'action' key and must provide cells as strings."
                            )
                            code_context = {
                                "target_kind": "DOCX table expansion",
                                "target_preview_text": self._extract_docx_text(output_docx_path),
                                "editable_blocks": self._public_docx_blocks(docx_blocks),
                                "editable_tables": self._public_docx_tables(docx_tables),
                                "allow_table_expansion": True,
                                "source_columns": list((tabular_context or {}).get("columns") or []),
                                "source_files": list((tabular_context or {}).get("source_files") or []),
                                "output_contract": output_contract,
                            }
                            code_prompt = self._build_generated_code_prompt(
                                tabular_context=tabular_context or {},
                                target_kind="DOCX table expansion",
                                target_context=code_context,
                                output_contract=output_contract,
                            )
                            edit_payload = self._generate_and_execute_tabular_transform(
                                source_parts + target_parts,
                                prompt=code_prompt,
                                rows=tabular_rows,
                                context=code_context,
                                expected_key="operations",
                                label="docx_table_expansion_generated_code",
                            )
                            fill_plan["table_code_hash"] = edit_payload.get("code_hash")
                        else:
                            edit_payload = self._generate_collection_json_response(
                                source_parts + target_parts,
                                prompt=edit_prompt,
                                schema=self._docx_edit_schema(),
                                collection_key="operations",
                                label="docx_table_expansion",
                                continue_on_full_batch=True,
                            )
                        operations = [
                            item
                            for item in (self._normalize_docx_edit_operation(item) for item in (edit_payload.get("operations") or []))
                            if item and str(item.get("action") or "").strip()
                        ]
                        table_operations = [
                            item
                            for item in operations
                            if str(item.get("action") or "").strip() in {"insert_table_row_after", "insert_table_column_after"}
                        ]
                        warnings.extend([str(item) for item in (edit_payload.get("warnings") or []) if str(item).strip()])

                    fill_plan["table_operations"] = table_operations
                    expanded_docx_path = os.path.join(temp_dir, f"filled-expanded-{uuid.uuid4()}.docx")
                    warnings.extend(
                        self._apply_docx_edit_plan(
                            output_docx_path,
                            table_operations,
                            expanded_docx_path,
                            allow_table_expansion=True,
                        )
                    )
                    final_local_path = expanded_docx_path
            else:
                processing_strategy = "docx_edit_in_place"
                target_preview_text = self._extract_docx_text(target_local_path)
                from docx import Document as DocxDocument

                target_doc = DocxDocument(target_local_path)
                docx_blocks, _block_map = self._collect_docx_blocks(target_doc)
                docx_tables, _table_map = self._collect_docx_tables(target_doc)
                edit_prompt = self._build_docx_edit_prompt(
                    source_text=source_text,
                    block_summary=self._summarize_docx_blocks(docx_blocks),
                    table_summary=self._summarize_docx_tables(docx_tables),
                    target_preview_text=target_preview_text,
                    allow_table_expansion=bool(run.allow_docx_table_expansion),
                )
                if tabular_rows:
                    processing_strategy = "docx_edit_generated_code"
                    output_contract = (
                        "Return {'operations': [operation, ...], 'warnings': [...]}. Each operation must use one of "
                        "replace_text_in_block, replace_block_text, append_to_block, insert_before_block, insert_after_block, "
                        "insert_table_row_after, or insert_table_column_after with valid provided block_id/table_id values. "
                        "Each operation object must put the operation name in the 'action' key and must provide cells as strings."
                    )
                    code_context = {
                        "target_kind": "DOCX edit in place",
                        "target_preview_text": target_preview_text,
                        "editable_blocks": self._public_docx_blocks(docx_blocks),
                        "editable_tables": self._public_docx_tables(docx_tables),
                        "allow_table_expansion": bool(run.allow_docx_table_expansion),
                        "source_columns": list((tabular_context or {}).get("columns") or []),
                        "source_files": list((tabular_context or {}).get("source_files") or []),
                        "output_contract": output_contract,
                    }
                    code_prompt = self._build_generated_code_prompt(
                        tabular_context=tabular_context or {},
                        target_kind="DOCX edit in place",
                        target_context=code_context,
                        output_contract=output_contract,
                    )
                    edit_payload = self._generate_and_execute_tabular_transform(
                        source_parts + target_parts,
                        prompt=code_prompt,
                        rows=tabular_rows,
                        context=code_context,
                        expected_key="operations",
                        label="docx_edit_generated_code",
                    )
                else:
                    edit_payload = self._generate_collection_json_response(
                        source_parts + target_parts,
                        prompt=edit_prompt,
                        schema=self._docx_edit_schema(),
                        collection_key="operations",
                        label="docx_edit_in_place",
                        continue_on_full_batch=True,
                    )
                operations = [
                    item
                    for item in (self._normalize_docx_edit_operation(item) for item in (edit_payload.get("operations") or []))
                    if item and str(item.get("action") or "").strip()
                ]
                warnings.extend([str(item) for item in (edit_payload.get("warnings") or []) if str(item).strip()])
                fill_plan = {
                    "strategy": processing_strategy,
                    "operations": operations,
                    "allow_docx_table_expansion": bool(run.allow_docx_table_expansion),
                    "source_rows": len(tabular_rows) if tabular_rows else None,
                    "code_hash": edit_payload.get("code_hash") if tabular_rows else None,
                }
                final_local_path = os.path.join(temp_dir, f"filled-{uuid.uuid4()}.docx")
                warnings.extend(
                    self._apply_docx_edit_plan(
                        target_local_path,
                        operations,
                        final_local_path,
                        allow_table_expansion=bool(run.allow_docx_table_expansion),
                    )
                )
                warnings.append("The DOCX target had no placeholders, so Form Fill edited the original DOCX in place.")

            final_extension = "docx"
            final_mime_type = DOCX_MIME
            if run.output_format == "pdf":
                final_local_path = await converter.convert_docx_local_to_pdf_local(final_local_path, out_dir=temp_dir)
                final_extension = "pdf"
                final_mime_type = PDF_MIME
            return {
                "local_path": final_local_path,
                "filename": self._filled_filename(run.target_filename, filename_suffix, final_extension),
                "mime_type": final_mime_type,
                "strategy": processing_strategy,
                "warnings": warnings,
                "fill_plan": self._compact_fill_plan(fill_plan),
            }

        raise ValueError("Unsupported target file type")

    def _sync_run_output_counts(self, db: Session, run: FormFillRun) -> None:
        completed = db.query(FormFillOutput).filter(
            FormFillOutput.run_id == run.id,
            FormFillOutput.status == "completed",
        ).count()
        failed = db.query(FormFillOutput).filter(
            FormFillOutput.run_id == run.id,
            FormFillOutput.status == "failed",
        ).count()
        run.completed_outputs = int(completed or 0)
        run.failed_outputs = int(failed or 0)

    async def _enqueue_output_units(
        self,
        *,
        db: Session,
        run: FormFillRun,
        target_page_count: int,
        units: list[dict[str, Any]],
        strategy: str,
    ) -> dict[str, Any]:
        if not units:
            raise ValueError("Form Fill could not find any source records to fill")

        self._check_usage_limit_or_raise(db, user_id=run.user_id, page_count=target_page_count * len(units))

        outputs = list(run.outputs or [])
        if not outputs:
            run.total_outputs = len(units)
            run.completed_outputs = 0
            run.failed_outputs = 0
            run.usage_basis = "target_pages_per_output"
            run.fill_plan = {"strategy": strategy, "outputs": len(units)}
            for unit in units:
                output = FormFillOutput(
                    run_id=run.id,
                    record_index=int(unit["record_index"]),
                    record_label=str(unit["record_label"]),
                    record_payload=unit["record_payload"],
                    status="pending",
                )
                db.add(output)
            db.commit()
        else:
            run.total_outputs = max(int(run.total_outputs or 0), len(outputs))
            if not isinstance(run.fill_plan, dict) or run.fill_plan.get("strategy") != strategy:
                run.fill_plan = {"strategy": strategy, "outputs": int(run.total_outputs or len(outputs))}
            self._sync_run_output_counts(db, run)
            db.commit()

        pending_outputs = db.query(FormFillOutput).filter(
            FormFillOutput.run_id == run.id,
            FormFillOutput.status == "pending",
        ).order_by(FormFillOutput.record_index.asc()).all()

        enqueued = 0
        failed_to_enqueue = 0
        if pending_outputs:
            last_delay = cloud_run_task_service.calculate_stagger_delay(
                len(pending_outputs) - 1,
                batch_size_env="FORM_FILL_OUTPUT_ENQUEUE_BATCH_SIZE",
                batch_delay_env="FORM_FILL_OUTPUT_ENQUEUE_BATCH_DELAY_SECONDS",
                max_delay_env="FORM_FILL_OUTPUT_ENQUEUE_MAX_DELAY_SECONDS",
                jitter_env="FORM_FILL_OUTPUT_ENQUEUE_JITTER_SECONDS",
                jitter_seed=f"{run.id}:{pending_outputs[-1].id}",
            )
            logger.info(
                "Staggering %s Form Fill output tasks for run %s across delays 0s-%ss",
                len(pending_outputs),
                run.id,
                last_delay,
            )

        for index, output in enumerate(pending_outputs):
            try:
                delay_seconds = cloud_run_task_service.calculate_stagger_delay(
                    index,
                    batch_size_env="FORM_FILL_OUTPUT_ENQUEUE_BATCH_SIZE",
                    batch_delay_env="FORM_FILL_OUTPUT_ENQUEUE_BATCH_DELAY_SECONDS",
                    max_delay_env="FORM_FILL_OUTPUT_ENQUEUE_MAX_DELAY_SECONDS",
                    jitter_env="FORM_FILL_OUTPUT_ENQUEUE_JITTER_SECONDS",
                    jitter_seed=f"{run.id}:{output.id}",
                )
                await cloud_run_task_service.enqueue_form_fill_output_task(
                    str(run.id),
                    str(output.id),
                    delay_seconds=delay_seconds,
                )
                enqueued += 1
            except Exception as enqueue_exc:
                logger.exception("Failed to enqueue Form Fill output %s", output.id)
                output.status = "failed"
                output.error_message = f"Failed to enqueue output task: {enqueue_exc}"
                output.completed_at = datetime.now(timezone.utc)
                failed_to_enqueue += 1

        self._sync_run_output_counts(db, run)
        db.commit()
        await self._finalize_run_if_ready(str(run.id))
        return {
            "success": True,
            "run_id": str(run.id),
            "outputs": int(run.total_outputs or len(units)),
            "enqueued_outputs": enqueued,
            "failed_to_enqueue": failed_to_enqueue,
        }

    async def process_output(
        self,
        run_id: str,
        output_id: str,
        *,
        task_retry_count: Optional[int] = None,
        task_execution_count: Optional[int] = None,
        task_name: Optional[str] = None,
        task_queue_name: Optional[str] = None,
    ) -> dict[str, Any]:
        db = self._get_session()
        temp_dir = tempfile.mkdtemp(prefix="form_fill_output_")
        lock_acquired = False
        try:
            lock_acquired = self._try_advisory_lock(db, output_id)
            if not lock_acquired:
                raise RuntimeError(f"Form Fill output {output_id} is already being processed")

            run = db.query(FormFillRun).filter(FormFillRun.id == uuid.UUID(str(run_id))).first()
            if not run:
                raise ValueError("Form Fill run not found")
            output = db.query(FormFillOutput).filter(
                FormFillOutput.id == uuid.UUID(str(output_id)),
                FormFillOutput.run_id == run.id,
            ).first()
            if not output:
                raise ValueError("Form Fill output not found")

            if output.status == "completed" and output.result_gcs_object_name:
                await self._finalize_run_if_ready(str(run.id))
                return {"success": True, "run_id": str(run.id), "output_id": str(output.id), "skipped": True}
            if output.status == "failed":
                await self._finalize_run_if_ready(str(run.id))
                return {"success": False, "run_id": str(run.id), "output_id": str(output.id), "skipped": True, "terminal": True}

            output.status = "in_progress"
            output.error_message = None
            db.commit()

            try:
                target_local_path = os.path.join(temp_dir, f"target{_safe_ext(run.target_filename, '.bin')}")
                await self._download_to_local(run.target_gcs_object_name, target_local_path)
                await self._ensure_run_target_page_count(db, run, target_local_path, temp_dir)
                tabular_context = await self._build_tabular_source_context(
                    run,
                    record_payload=output.record_payload if isinstance(output.record_payload, dict) else {},
                )
                if tabular_context:
                    source_parts = []
                    source_text = self._tabular_source_summary(tabular_context)
                else:
                    source_parts, source_text = await self._build_output_source_context(
                        run=run,
                        record_payload=output.record_payload or {},
                        record_index=int(output.record_index or 0),
                    )
                generated = await self._generate_filled_document(
                    run=run,
                    temp_dir=temp_dir,
                    target_local_path=target_local_path,
                    source_parts=source_parts,
                    source_text=source_text,
                    tabular_context=tabular_context,
                    filename_suffix=f"{output.record_index + 1:03d}_{output.record_label}",
                )
                result_object_name = (
                    f"form-fill/{run.user_id}/runs/{run.id}/outputs/"
                    f"{output.record_index + 1:03d}-{uuid.uuid4()}{_safe_ext(generated['filename'], '.bin')}"
                )
                await self.storage_service.upload_file(generated["local_path"], result_object_name)
                output.status = "completed"
                output.result_gcs_object_name = result_object_name
                output.result_filename = generated["filename"]
                output.result_file_type = generated["mime_type"]
                output.fill_plan = generated["fill_plan"]
                output.warnings = generated["warnings"]
                output.completed_at = datetime.now(timezone.utc)
            except Exception as output_exc:
                logger.exception("Form Fill output %s failed", output_id)
                db.rollback()
                output = db.query(FormFillOutput).filter(FormFillOutput.id == uuid.UUID(str(output_id))).first()
                run = db.query(FormFillRun).filter(FormFillRun.id == uuid.UUID(str(run_id))).first()
                if not output or not run:
                    raise

                if self._should_retry_output_error(
                    output_exc,
                    task_retry_count=task_retry_count,
                    task_execution_count=task_execution_count,
                ):
                    attempt = self._form_fill_output_attempt(task_retry_count, task_execution_count) or 1
                    logger.warning(
                        "Retrying Form Fill output %s for run %s after attempt %s/%s failed: %s (task=%s queue=%s)",
                        output_id,
                        run_id,
                        attempt,
                        self.output_max_attempts,
                        output_exc,
                        task_name,
                        task_queue_name,
                    )
                    output.status = "pending"
                    output.error_message = f"Retrying after attempt {attempt} failed: {output_exc}"
                    output.completed_at = None
                    self._sync_run_output_counts(db, run)
                    db.commit()
                    raise

                output.status = "failed"
                output.error_message = str(output_exc)
                output.completed_at = datetime.now(timezone.utc)

            self._sync_run_output_counts(db, run)
            db.commit()
            await self._finalize_run_if_ready(str(run.id))
            return {"success": output.status == "completed", "run_id": str(run.id), "output_id": str(output.id), "status": output.status}
        finally:
            if lock_acquired:
                self._advisory_unlock(db, output_id)
            db.close()
            shutil.rmtree(temp_dir, ignore_errors=True)

    async def _finalize_run_if_ready(self, run_id: str) -> dict[str, Any]:
        db = self._get_session()
        temp_dir = tempfile.mkdtemp(prefix="form_fill_finalize_")
        lock_acquired = False
        try:
            lock_acquired = self._try_advisory_lock(db, run_id)
            if not lock_acquired:
                return {"success": True, "run_id": run_id, "finalized": False, "locked": True}

            run = db.query(FormFillRun).filter(FormFillRun.id == uuid.UUID(str(run_id))).first()
            if not run:
                raise ValueError("Form Fill run not found")
            if run.status in FORM_FILL_RUN_TERMINAL_STATUSES and run.result_gcs_object_name:
                return {"success": True, "run_id": str(run.id), "finalized": False, "skipped": True}

            outputs = db.query(FormFillOutput).filter(FormFillOutput.run_id == run.id).order_by(FormFillOutput.record_index.asc()).all()
            if not outputs:
                return {"success": True, "run_id": str(run.id), "finalized": False, "outputs": 0}
            if any(output.status not in FORM_FILL_OUTPUT_TERMINAL_STATUSES for output in outputs):
                self._sync_run_output_counts(db, run)
                db.commit()
                return {"success": True, "run_id": str(run.id), "finalized": False, "outputs": len(outputs)}

            completed_outputs = [
                output for output in outputs
                if output.status == "completed" and output.result_gcs_object_name and output.result_filename
            ]
            aggregate_warnings: list[str] = []
            for output in outputs:
                if isinstance(output.warnings, list):
                    aggregate_warnings.extend(f"{output.record_label}: {warning}" for warning in output.warnings)
                if output.status == "failed" and output.error_message:
                    aggregate_warnings.append(f"{output.record_label}: {output.error_message}")

            self._sync_run_output_counts(db, run)
            run.total_outputs = len(outputs)
            run.usage_basis = "target_pages_per_output"
            run.usage_pages = int(run.target_page_count or 0) * int(run.completed_outputs or 0)
            run.warnings = aggregate_warnings
            run.completed_at = datetime.now(timezone.utc)

            if not completed_outputs:
                run.status = "failed"
                run.error_message = run.error_message or "Form Fill failed to generate any output documents"
                db.commit()
                return {"success": False, "run_id": str(run.id), "finalized": True, "outputs": 0}

            zip_filename = self._filled_filename(run.target_filename, "batch", "zip")
            zip_local_path = os.path.join(temp_dir, zip_filename)
            used_names: set[str] = set()
            with zipfile.ZipFile(zip_local_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
                for output in completed_outputs:
                    suffix = _safe_ext(output.result_filename or "output.bin", ".bin")
                    local_path = os.path.join(temp_dir, f"output-{output.record_index}-{uuid.uuid4()}{suffix}")
                    await self._download_to_local(output.result_gcs_object_name, local_path)
                    archive_name = output.result_filename or f"output-{output.record_index + 1}{suffix}"
                    if archive_name in used_names:
                        archive_name = f"{Path(archive_name).stem}_{len(used_names) + 1}{Path(archive_name).suffix}"
                    used_names.add(archive_name)
                    archive.write(local_path, arcname=archive_name)

            result_object_name = f"form-fill/{run.user_id}/runs/{run.id}/result.zip"
            await self.storage_service.upload_file(zip_local_path, result_object_name)
            run.status = "completed" if int(run.failed_outputs or 0) == 0 else "completed_with_errors"
            last_completed = completed_outputs[-1]
            if isinstance(last_completed.fill_plan, dict):
                run.processing_strategy = last_completed.fill_plan.get("strategy") or run.processing_strategy
            run.result_gcs_object_name = result_object_name
            run.result_filename = zip_filename
            run.result_file_type = "application/zip"
            db.commit()
            self._record_usage_for_run(db, run)
            return {
                "success": True,
                "run_id": str(run.id),
                "finalized": True,
                "outputs": len(completed_outputs),
                "failed_outputs": int(run.failed_outputs or 0),
            }
        finally:
            if lock_acquired:
                self._advisory_unlock(db, run_id)
            db.close()
            shutil.rmtree(temp_dir, ignore_errors=True)

    async def process_run(self, run_id: str) -> dict[str, Any]:
        db = self._get_session()
        temp_dir = tempfile.mkdtemp(prefix="form_fill_run_")
        try:
            run = db.query(FormFillRun).filter(FormFillRun.id == uuid.UUID(str(run_id))).first()
            if not run:
                raise ValueError("Form Fill run not found")
            if run.status in {"completed", "completed_with_errors"} and run.result_gcs_object_name:
                self._record_usage_for_run(db, run)
                return {"success": True, "run_id": str(run.id), "skipped": True}

            run.status = "in_progress"
            run.error_message = None
            db.commit()

            target_local_path = os.path.join(temp_dir, f"target{_safe_ext(run.target_filename, '.bin')}")
            await self._download_to_local(run.target_gcs_object_name, target_local_path)
            target_page_count = await self._ensure_run_target_page_count(db, run, target_local_path, temp_dir)

            if (run.repeat_mode or REPEAT_MODE_SINGLE) == REPEAT_MODE_SOURCE_ROWS:
                records = await self._extract_repeat_records(run)
                if not records:
                    raise ValueError("Repeat mode could not find any source rows to fill")
                return await self._enqueue_output_units(
                    db=db,
                    run=run,
                    target_page_count=target_page_count,
                    units=records,
                    strategy="repeat_source_rows",
                )

            if (run.repeat_mode or REPEAT_MODE_SINGLE) != REPEAT_MODE_ALL_SOURCES:
                source_units = self._source_units_for_run(run)
                if len(source_units) > 1:
                    unit_kind = source_units[0].get("record_payload", {}).get("kind") if isinstance(source_units[0].get("record_payload"), dict) else None
                    return await self._enqueue_output_units(
                        db=db,
                        run=run,
                        target_page_count=target_page_count,
                        units=source_units,
                        strategy="extraction_tasks" if unit_kind == "extraction_task" else "source_files",
                    )

            self._check_usage_limit_or_raise(db, user_id=run.user_id, page_count=target_page_count)
            tabular_context = await self._build_tabular_source_context(run)
            if tabular_context:
                source_parts = []
                source_text = self._tabular_source_summary(tabular_context)
            else:
                source_parts = []
                source_text_sections: list[str] = []
                source_parts, source_text = await self._build_source_context(
                    run=run,
                    source_parts=source_parts,
                    source_text_sections=source_text_sections,
                )
            generated = await self._generate_filled_document(
                run=run,
                temp_dir=temp_dir,
                target_local_path=target_local_path,
                source_parts=source_parts,
                source_text=source_text,
                tabular_context=tabular_context,
            )
            result_object_name = f"form-fill/{run.user_id}/runs/{run.id}/result{_safe_ext(generated['filename'], '.bin')}"
            await self.storage_service.upload_file(generated["local_path"], result_object_name)

            run.status = "completed"
            run.processing_strategy = generated["strategy"]
            run.result_gcs_object_name = result_object_name
            run.result_filename = generated["filename"]
            run.result_file_type = generated["mime_type"]
            run.fill_plan = generated["fill_plan"]
            run.warnings = generated["warnings"]
            run.total_outputs = 1
            run.completed_outputs = 1
            run.failed_outputs = 0
            run.usage_basis = "target_pages_per_output"
            run.usage_pages = target_page_count
            run.completed_at = datetime.now(timezone.utc)
            db.commit()
            self._record_usage_for_run(db, run)
            return {"success": True, "run_id": str(run.id), "strategy": run.processing_strategy}
        except Exception as exc:
            logger.exception("Form Fill run %s failed", run_id)
            try:
                db.rollback()
                run = db.query(FormFillRun).filter(FormFillRun.id == uuid.UUID(str(run_id))).first()
                if run:
                    run.status = "failed"
                    run.error_message = str(exc)
                    run.completed_at = datetime.now(timezone.utc)
                    db.commit()
            except Exception:
                db.rollback()
            raise
        finally:
            db.close()
            shutil.rmtree(temp_dir, ignore_errors=True)


form_fill_service = FormFillService()
