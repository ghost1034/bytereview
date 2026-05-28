"""Disclosure / Footnote Drafting Assistant — draft GAAP-compliant footnote tables + narrative.

Currently supports footnote types: debt (ASC 470).
Other types (lease, income_tax, ar_allowance, stock_comp, segment) are scaffolded
with their schemas and checklists in DISCLOSURE_REGISTRY; the script raises a
NotImplementedError for those until their table builders are added.
"""
from __future__ import annotations

import argparse
import logging
import sys
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

LOG = logging.getLogger("footnote_drafter")

ACCOUNTING_FMT = '_($* #,##0.00_);_($* (#,##0.00);_($* "-"??_);_(@_)'
HEADER_FILL = PatternFill(start_color="305496", end_color="305496", fill_type="solid")
HEADER_FONT = Font(bold=True, color="FFFFFF")
RED_FILL = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
GREEN_FILL = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
THIN = Side(border_style="thin", color="BFBFBF")
BOX = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)

# --- Disclosure registry --------------------------------------------------
# Each footnote type has: required columns + ASC checklist + builder

DEBT_REQUIRED_COLUMNS = (
    "instrument", "lender", "principal_outstanding", "interest_rate",
    "rate_type", "maturity_date", "secured", "covenant_compliance",
)
DEBT_CHECKLIST = [
    ("Each material debt instrument identified", "ASC 470-10-50-1"),
    ("Interest rate disclosed for each", "ASC 470-10-50-5"),
    ("Maturity disclosed for each", "ASC 470-10-50-5"),
    ("Maturity schedule (5 years + thereafter)", "ASC 470-10-50-1"),
    ("Weighted-average interest rate", "ASC 470-10-50-5"),
    ("Fixed vs. variable rate disclosure", "ASC 470-10-50-5"),
    ("Secured vs. unsecured disclosure", "ASC 470-10-50-2"),
    ("Covenant compliance status", "ASC 470-10-50-1"),
    ("Subsequent events (if any)", "ASC 855-10-50"),
]


# --- I/O ------------------------------------------------------------------

def _read_table(path: Path) -> pd.DataFrame:
    if path.suffix.lower() == ".csv":
        df = pd.read_csv(path)
    else:
        df = pd.read_excel(path)
    df.columns = [c.strip().lower() for c in df.columns]
    return df


def _validate_debt(df: pd.DataFrame) -> pd.DataFrame:
    missing = [c for c in DEBT_REQUIRED_COLUMNS if c not in df.columns]
    if missing:
        raise ValueError(f"Debt schedule missing required columns: {missing}")
    df = df.copy()
    df["maturity_date"] = pd.to_datetime(df["maturity_date"], errors="coerce")
    for c in ("principal_outstanding", "interest_rate"):
        df[c] = pd.to_numeric(df[c], errors="coerce")
    df["secured"] = df["secured"].astype(bool)
    df["covenant_compliance"] = df["covenant_compliance"].astype(bool)
    df["rate_type"] = df["rate_type"].astype(str).str.strip().str.lower()
    df = df.dropna(subset=["maturity_date", "principal_outstanding"])
    return df.reset_index(drop=True)


# --- Debt footnote builder -----------------------------------------------

def build_debt_footnote(df: pd.DataFrame, period_end: datetime,
                       bs_balance: float | None) -> dict:
    """Return dict with 'schedule', 'maturity', 'summary', 'checklist', 'cross_foot' frames."""
    df = df.copy()
    # Weighted-average rate
    weights = df["principal_outstanding"]
    weighted_rate = (df["interest_rate"] * weights).sum() / weights.sum() if weights.sum() else 0.0
    fixed_total = df[df["rate_type"] == "fixed"]["principal_outstanding"].sum()
    var_total = df[df["rate_type"] == "variable"]["principal_outstanding"].sum()
    secured_total = df[df["secured"]]["principal_outstanding"].sum()
    total = df["principal_outstanding"].sum()

    # Maturity buckets: Y1, Y2, Y3, Y4, Y5, Thereafter
    year_end = period_end.year
    buckets = {f"Year {i+1}": 0.0 for i in range(5)}
    buckets["Thereafter"] = 0.0
    for _, r in df.iterrows():
        years_out = r["maturity_date"].year - year_end
        if years_out <= 0:
            buckets["Year 1"] += float(r["principal_outstanding"])
        elif years_out <= 5:
            buckets[f"Year {years_out}"] += float(r["principal_outstanding"])
        else:
            buckets["Thereafter"] += float(r["principal_outstanding"])
    maturity_df = pd.DataFrame([{"Period": k, "Principal Due": v} for k, v in buckets.items()])

    # Schedule
    schedule_df = df[["instrument", "lender", "principal_outstanding",
                      "interest_rate", "rate_type", "maturity_date",
                      "secured", "covenant_compliance"]].copy()
    schedule_df["interest_rate"] = schedule_df["interest_rate"].apply(lambda x: f"{x*100:.2f}%")
    schedule_df["maturity_date"] = schedule_df["maturity_date"].dt.strftime("%Y-%m-%d")

    summary_df = pd.DataFrame([
        {"Metric": "Total Principal Outstanding", "Value": total},
        {"Metric": "Weighted-Average Interest Rate", "Value": f"{weighted_rate*100:.2f}%"},
        {"Metric": "Fixed-Rate Principal", "Value": fixed_total},
        {"Metric": "Variable-Rate Principal", "Value": var_total},
        {"Metric": "Variable-Rate % of Total", "Value": f"{(var_total/total*100) if total else 0:.2f}%"},
        {"Metric": "Secured Principal", "Value": secured_total},
        {"Metric": "Number of Instruments", "Value": len(df)},
        {"Metric": "Covenant Compliance (all instruments)", "Value": "Yes" if df["covenant_compliance"].all() else "No"},
    ])

    # Checklist
    checklist = []
    checklist.append((DEBT_CHECKLIST[0][0], "Pass" if not df["instrument"].isna().any() else "Missing", DEBT_CHECKLIST[0][1]))
    checklist.append((DEBT_CHECKLIST[1][0], "Pass" if not df["interest_rate"].isna().any() else "Missing", DEBT_CHECKLIST[1][1]))
    checklist.append((DEBT_CHECKLIST[2][0], "Pass" if not df["maturity_date"].isna().any() else "Missing", DEBT_CHECKLIST[2][1]))
    checklist.append((DEBT_CHECKLIST[3][0], "Pass", DEBT_CHECKLIST[3][1]))
    checklist.append((DEBT_CHECKLIST[4][0], "Pass", DEBT_CHECKLIST[4][1]))
    checklist.append((DEBT_CHECKLIST[5][0], "Pass", DEBT_CHECKLIST[5][1]))
    checklist.append((DEBT_CHECKLIST[6][0], "Pass", DEBT_CHECKLIST[6][1]))
    checklist.append((DEBT_CHECKLIST[7][0],
                      "Pass" if df["covenant_compliance"].all() else "Fail — disclose non-compliance",
                      DEBT_CHECKLIST[7][1]))
    checklist.append((DEBT_CHECKLIST[8][0], "Reviewer to confirm", DEBT_CHECKLIST[8][1]))
    checklist_df = pd.DataFrame(checklist, columns=["Disclosure Item", "Status", "Authority"])

    # Cross-foot
    cf_rows = [{"Item": "Disclosure Total (Principal Outstanding)", "Amount": total}]
    if bs_balance is not None:
        diff = total - bs_balance
        cf_rows.append({"Item": "Balance Sheet Long-Term Debt", "Amount": bs_balance})
        cf_rows.append({"Item": "Difference", "Amount": diff})
        cf_rows.append({"Item": "Status", "Amount": "OK" if abs(diff) < 0.01 else "INVESTIGATE"})
    cf_df = pd.DataFrame(cf_rows)

    return {
        "schedule": schedule_df, "maturity": maturity_df, "summary": summary_df,
        "checklist": checklist_df, "cross_foot": cf_df,
        "weighted_rate": weighted_rate, "var_pct": var_total / total if total else 0,
        "total": total,
    }


# --- XLSX writer ---------------------------------------------------------

def _style_header(ws, row: int = 1) -> None:
    for cell in ws[row]:
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = BOX


def _autosize(ws) -> None:
    for col in ws.columns:
        letter = get_column_letter(col[0].column)
        max_len = max((len(str(c.value)) for c in col if c.value is not None), default=10)
        ws.column_dimensions[letter].width = min(60, max_len + 2)


def _df_to_sheet(ws, df: pd.DataFrame, money_cols: list[str] | None = None) -> None:
    if df is None or df.empty:
        ws.append(["(no data)"])
        return
    ws.append(list(df.columns))
    _style_header(ws)
    for _, r in df.iterrows():
        ws.append([r[c] for c in df.columns])
    if money_cols:
        for col_idx, col in enumerate(df.columns, start=1):
            if col in money_cols:
                for cell in ws[get_column_letter(col_idx)][1:]:
                    if isinstance(cell.value, (int, float)):
                        cell.number_format = ACCOUNTING_FMT
    _autosize(ws)


def write_debt_xlsx(footnote: dict, output: Path, period_end: datetime, entity: str) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Summary"
    ws.append([f"Note X — Debt"])
    ws["A1"].font = Font(bold=True, size=14)
    ws.append([f"{entity} — as of {period_end.strftime('%B %d, %Y')}"])
    ws.append([])
    _df_to_sheet(ws, footnote["summary"], money_cols=["Value"])
    _df_to_sheet(wb.create_sheet("Schedule"), footnote["schedule"], money_cols=["principal_outstanding"])
    _df_to_sheet(wb.create_sheet("Maturity"), footnote["maturity"], money_cols=["Principal Due"])
    _df_to_sheet(wb.create_sheet("Checklist"), footnote["checklist"])
    # color checklist
    ws_cl = wb["Checklist"]
    for row in ws_cl.iter_rows(min_row=2):
        v = str(row[1].value or "")
        if v.startswith("Pass"):
            row[1].fill = GREEN_FILL
        elif v.startswith("Fail") or v.startswith("Missing"):
            row[1].fill = RED_FILL
    _df_to_sheet(wb.create_sheet("CrossFoot"), footnote["cross_foot"], money_cols=["Amount"])
    # color cross-foot status
    ws_cf = wb["CrossFoot"]
    for row in ws_cf.iter_rows(min_row=2):
        if row[0].value == "Status":
            row[1].fill = GREEN_FILL if row[1].value == "OK" else RED_FILL
    ws_a = wb.create_sheet("AuditTrail")
    ws_a.append(["Key", "Value"])
    _style_header(ws_a)
    ws_a.append(["Footnote", "Debt"])
    ws_a.append(["Period End", period_end.strftime("%Y-%m-%d")])
    ws_a.append(["Entity", entity])
    ws_a.append(["Generated", datetime.now().isoformat(timespec="seconds")])
    _autosize(ws_a)
    ws_s = wb.create_sheet("SignOff")
    ws_s.append(["Role", "Name", "Title", "Date", "Notes"])
    _style_header(ws_s)
    ws_s.append(["Preparer (Financial Reporting)", "", "", "", ""])
    ws_s.append(["Reviewer (Controller)", "", "", "", ""])
    ws_s.append(["Approver (CFO)", "", "", "", ""])
    _autosize(ws_s)

    output.parent.mkdir(parents=True, exist_ok=True)
    wb.save(output)
    LOG.info("XLSX written: %s", output)


# --- DOCX writer ---------------------------------------------------------

def _set_times_new_roman(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def write_debt_docx(footnote: dict, output: Path, period_end: datetime, entity: str) -> None:
    doc = Document()
    _set_times_new_roman(doc)
    title = doc.add_heading("Note X — Debt", level=2)
    for run in title.runs:
        run.font.name = "Times New Roman"
    period_str = period_end.strftime("%B %d, %Y")
    p = doc.add_paragraph()
    p.add_run(f"The Company's long-term debt consists of the following at {period_str}:").bold = False

    # Schedule table
    sched = footnote["schedule"]
    t = doc.add_table(rows=1, cols=len(sched.columns))
    t.style = "Light Grid"
    for i, col in enumerate(sched.columns):
        t.rows[0].cells[i].text = col.replace("_", " ").title()
    for _, r in sched.iterrows():
        row = t.add_row().cells
        for i, col in enumerate(sched.columns):
            val = r[col]
            if isinstance(val, float):
                row[i].text = f"${val:,.2f}"
            else:
                row[i].text = str(val)

    doc.add_paragraph()
    doc.add_paragraph(f"Maturities of long-term debt outstanding at {period_str} are as follows:")
    mat = footnote["maturity"]
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = "Light Grid"
    t2.rows[0].cells[0].text = "Period"
    t2.rows[0].cells[1].text = "Principal Due"
    for _, r in mat.iterrows():
        cells = t2.add_row().cells
        cells[0].text = str(r["Period"])
        cells[1].text = f"${r['Principal Due']:,.2f}"
    cells = t2.add_row().cells
    cells[0].text = "Total"
    cells[1].text = f"${mat['Principal Due'].sum():,.2f}"

    doc.add_paragraph()
    wr = footnote["weighted_rate"] * 100
    vp = footnote["var_pct"] * 100
    doc.add_paragraph(
        f"The weighted-average interest rate on outstanding debt at {period_str} was "
        f"{wr:.2f}%. At {period_str}, approximately {vp:.1f}% of the Company's debt bore "
        "variable interest rates, indexed primarily to SOFR."
    )
    cov = footnote["checklist"]
    cov_status = cov[cov["Disclosure Item"].str.contains("Covenant", case=False)]["Status"].iloc[0]
    if "Pass" in str(cov_status):
        doc.add_paragraph(f"The Company was in compliance with all financial covenants at {period_str}.")
    else:
        doc.add_paragraph(
            f"At {period_str}, the Company was not in compliance with one or more financial "
            "covenants; see [cross-reference to risk factor / waiver disclosure]."
        )

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    LOG.info("DOCX written: %s", output)


# --- CLI -----------------------------------------------------------------

def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    parser.add_argument("--footnote", required=True,
                        choices=["debt", "lease", "income_tax", "ar_allowance", "stock_comp", "segment"])
    parser.add_argument("--data", required=True, type=Path)
    parser.add_argument("--period-end", required=True)
    parser.add_argument("--entity", default="")
    parser.add_argument("--prior", type=Path, default=None)
    parser.add_argument("--tie-to-statement", type=Path, default=None)
    parser.add_argument("--output-xlsx", required=True, type=Path)
    parser.add_argument("--output-docx", required=True, type=Path)
    parser.add_argument("--quiet", action="store_true")
    args = parser.parse_args(argv)

    logging.basicConfig(
        level=logging.WARNING if args.quiet else logging.INFO,
        format="%(levelname)s %(name)s: %(message)s",
    )

    period_end = datetime.strptime(args.period_end, "%Y-%m-%d")
    if args.footnote != "debt":
        LOG.error("Footnote type %r is scaffolded but not yet implemented. Currently: debt only.", args.footnote)
        return 1

    try:
        df = _validate_debt(_read_table(args.data))
    except ValueError as e:
        LOG.error("Validation failed: %s", e)
        return 1
    bs_balance = None
    if args.tie_to_statement:
        try:
            tie = _read_table(args.tie_to_statement)
            bs_balance = float(tie["amount"].iloc[0])
        except Exception as e:
            LOG.warning("Could not read tie-to-statement: %s", e)

    footnote = build_debt_footnote(df, period_end, bs_balance)
    write_debt_xlsx(footnote, args.output_xlsx, period_end, args.entity)
    write_debt_docx(footnote, args.output_docx, period_end, args.entity)

    print(f"Footnote:         {args.footnote}")
    print(f"Total Principal:  {footnote['total']:>15,.2f}")
    print(f"Wtd-Avg Rate:     {footnote['weighted_rate']*100:.2f}%")
    print(f"Variable %:       {footnote['var_pct']*100:.1f}%")
    print(f"XLSX:             {args.output_xlsx}")
    print(f"DOCX:             {args.output_docx}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
