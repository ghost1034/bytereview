"""Tax Research Orchestrator — IRAC memo builder with structured IRC/ASC query framework.

Shared pattern for tax research skills. Domain-specific citations passed via --domain flag.
"""
from __future__ import annotations

import argparse
import logging
import os
import sys
from datetime import datetime
from pathlib import Path
from typing import Any

import yaml
from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Side, Border

LOG = logging.getLogger("tax_research")

HEADER_FILL = PatternFill(start_color="305496", end_color="305496", fill_type="solid")
HEADER_FONT = Font(bold=True, color="FFFFFF")
THIN = Side(border_style="thin", color="BFBFBF")
BOX = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)

DOMAIN_CONFIG = {
    "sales-tax": {
        "title": "State and Local Tax (SALT) Research Memo",
        "primary_sources": ["State DOR statutes", "Wayfair economic nexus", "Streamlined Sales Tax"],
        "search_prefix": "state sales tax nexus taxability",
    },
    "transfer-pricing": {
        "title": "Transfer Pricing Research Memo",
        "primary_sources": ["IRC §482", "OECD Transfer Pricing Guidelines", "Treas. Reg. §1.482"],
        "search_prefix": "transfer pricing arm's length OECD",
    },
    "individual": {
        "title": "Individual Income Tax Research Memo",
        "primary_sources": ["IRC Subtitle A", "Treas. Reg. §1", "IRS Publications"],
        "search_prefix": "individual income tax IRC",
    },
    "corporate": {
        "title": "Corporate Income Tax Research Memo",
        "primary_sources": ["IRC §11", "IRC §243", "IRC §382", "ASC 740"],
        "search_prefix": "corporate income tax IRC",
    },
    "partnership": {
        "title": "Partnership Tax Research Memo",
        "primary_sources": ["IRC Subchapter K", "IRC §704", "IRC §707", "IRC §751"],
        "search_prefix": "partnership tax IRC subchapter K",
    },
}


def _load_query(path: Path) -> dict[str, Any]:
    with path.open(encoding="utf-8") as fh:
        data = yaml.safe_load(fh)
    if not isinstance(data, dict):
        raise ValueError("Query file must be a YAML mapping")
    return data


def _try_search(query: str) -> list[str]:
    api_key = os.environ.get("TAVILY_API_KEY")
    if not api_key:
        LOG.info("TAVILY_API_KEY not set — skipping live search")
        return []
    try:
        from tavily import TavilyClient
        client = TavilyClient(api_key=api_key)
        resp = client.search(query=query, max_results=3)
        return [r.get("content", "")[:500] for r in resp.get("results", [])]
    except Exception as exc:
        LOG.warning("Search failed: %s", exc)
        return []


def build_research_plan(query: dict[str, Any], domain: str) -> list[dict[str, str]]:
    cfg = DOMAIN_CONFIG[domain]
    questions = query.get("research_questions") or [query.get("question", "Primary research question")]
    plan = []
    for i, q in enumerate(questions, start=1):
        search_q = f"{cfg['search_prefix']} {q}"
        snippets = _try_search(search_q)
        plan.append({
            "step": str(i),
            "question": str(q),
            "search_query": search_q,
            "sources": "; ".join(cfg["primary_sources"]),
            "findings": snippets[0] if snippets else "Manual research required — review primary sources.",
        })
    return plan


def write_docx(query: dict[str, Any], plan: list[dict[str, str]], domain: str, output: Path) -> None:
    cfg = DOMAIN_CONFIG[domain]
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    doc.add_heading(cfg["title"], level=1)
    doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph(f"Matter: {query.get('matter_id', 'N/A')}")

    doc.add_heading("Issue", level=2)
    doc.add_paragraph(str(query.get("issue", query.get("question", ""))))

    doc.add_heading("Rule", level=2)
    for src in cfg["primary_sources"]:
        doc.add_paragraph(src)

    doc.add_heading("Analysis", level=2)
    for step in plan:
        doc.add_paragraph(f"Question {step['step']}: {step['question']}")
        doc.add_paragraph(f"Research query: {step['search_query']}")
        doc.add_paragraph(f"Findings: {step['findings']}")

    doc.add_heading("Conclusion", level=2)
    doc.add_paragraph(str(query.get("preliminary_conclusion",
                                     "Preliminary conclusion — subject to Tax Manager review.")))

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    LOG.info("Memo written: %s", output)


def write_xlsx(plan: list[dict[str, str]], query: dict[str, Any], output: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Summary"
    ws.append(["Tax Research Summary"])
    ws["A1"].font = Font(bold=True, size=14)
    ws.append([])
    ws.append(["Matter ID", query.get("matter_id", "")])
    ws.append(["Steps", len(plan)])
    ws.append(["Generated", datetime.now().isoformat(timespec="seconds")])

    ws_p = wb.create_sheet("ResearchPlan")
    ws_p.append(["Step", "Question", "Search Query", "Sources", "Findings"])
    for cell in ws_p[1]:
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
    for step in plan:
        ws_p.append([step["step"], step["question"], step["search_query"], step["sources"], step["findings"]])

    ws_at = wb.create_sheet("AuditTrail")
    ws_at.append(["Key", "Value"])
    ws_at.append(["Generated", datetime.now().isoformat(timespec="seconds")])

    output.parent.mkdir(parents=True, exist_ok=True)
    wb.save(output)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    parser.add_argument("--query", required=True, type=Path)
    parser.add_argument("--domain", required=True, choices=list(DOMAIN_CONFIG))
    parser.add_argument("--output-docx", required=True, type=Path)
    parser.add_argument("--output-xlsx", required=True, type=Path)
    parser.add_argument("--quiet", action="store_true")
    args = parser.parse_args(argv)

    logging.basicConfig(
        level=logging.WARNING if args.quiet else logging.INFO,
        format="%(levelname)s %(name)s: %(message)s",
    )

    try:
        query = _load_query(args.query)
    except (OSError, yaml.YAMLError, ValueError) as exc:
        LOG.error("%s", exc)
        return 1

    plan = build_research_plan(query, args.domain)
    write_docx(query, plan, args.domain, args.output_docx)
    write_xlsx(plan, query, args.output_xlsx)

    print(f"Matter:          {query.get('matter_id', 'N/A')}")
    print(f"Research steps:  {len(plan)}")
    print(f"Memo:            {args.output_docx}")
    print(f"Workpaper:       {args.output_xlsx}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
