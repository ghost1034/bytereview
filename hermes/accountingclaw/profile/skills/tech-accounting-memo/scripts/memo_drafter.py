"""Technical Accounting Memo Drafter — IRAC-format DOCX from transaction facts YAML.

Implements the workflow in ../SKILL.md.
"""
from __future__ import annotations

import argparse
import logging
import sys
from datetime import datetime
from pathlib import Path
from typing import Any

import yaml
from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Side, Border, Alignment
from openpyxl.utils import get_column_letter

LOG = logging.getLogger("tech_accounting_memo")

HEADER_FILL = PatternFill(start_color="305496", end_color="305496", fill_type="solid")
HEADER_FONT = Font(bold=True, color="FFFFFF")
THIN = Side(border_style="thin", color="BFBFBF")
BOX = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)

ASC_CITATIONS = {
    "ASC 842": "Leases — identification, classification, measurement, and modification (ASC 842-10-25)",
    "ASC 606": "Revenue from Contracts with Customers — five-step model (ASC 606-10)",
    "ASC 805": "Business Combinations — acquisition method (ASC 805-20)",
    "ASC 480": "Distinguishing Liabilities from Equity (ASC 480-10)",
    "ASC 815": "Derivatives and Hedging — embedded derivatives (ASC 815-15)",
}


def _load_facts(path: Path) -> dict[str, Any]:
    with path.open(encoding="utf-8") as fh:
        data = yaml.safe_load(fh)
    if not isinstance(data, dict):
        raise ValueError("Transaction facts must be a YAML mapping")
    return data


def _apply_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def draft_memo(facts: dict[str, Any], output: Path) -> None:
    doc = Document()
    _apply_font(doc)
    title = facts.get("title", "Technical Accounting Memo")
    standard = str(facts.get("standard", "ASC"))
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Prepared: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph(f"Transaction ID: {facts.get('transaction_id', 'N/A')}")
    doc.add_paragraph(f"Issuer: {facts.get('issuer', 'N/A')}")

    doc.add_heading("Issue", level=2)
    issues = facts.get("issues") or []
    if issues:
        for issue in issues:
            doc.add_paragraph(str(issue))
    else:
        doc.add_paragraph(f"What is the appropriate accounting treatment under {standard}?")

    doc.add_heading("Rule", level=2)
    citation = ASC_CITATIONS.get(standard, f"Relevant guidance under {standard}.")
    doc.add_paragraph(citation)
    doc.add_paragraph(
        "The analysis follows U.S. GAAP authoritative literature and considers "
        "substance over form. All conclusions require human review before posting."
    )

    doc.add_heading("Analysis", level=2)
    doc.add_paragraph("Facts:")
    for fact in facts.get("facts") or []:
        doc.add_paragraph(str(fact), style="List Bullet")

    doc.add_paragraph(
        f"Applying {standard}, the entity should evaluate whether the transaction "
        "creates a separate unit of account or modifies an existing arrangement. "
        "Key judgments include timing of recognition, measurement basis, and "
        "presentation/classification on the balance sheet."
    )
    if standard == "ASC 842":
        doc.add_paragraph(
            "Under ASC 842-10-25-8, a lease modification not accounted for as a "
            "separate contract requires remeasurement of the lease liability using "
            "the modified terms and the discount rate at the effective date of the "
            "modification."
        )

    doc.add_heading("Conclusion", level=2)
    requested = facts.get("conclusion_requested", "Draft conclusion pending reviewer sign-off.")
    doc.add_paragraph(str(requested))
    doc.add_paragraph(
        "Draft journal entries (illustrative — not for posting without review):"
    )
    doc.add_paragraph("Debit  ROU Asset (adjustment)     [amount TBD]")
    doc.add_paragraph("Credit Lease Liability (adjustment)  [amount TBD]")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    LOG.info("Memo written: %s", output)


def write_summary_xlsx(facts: dict[str, Any], output: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Summary"
    ws.append(["Technical Accounting Memo Summary"])
    ws["A1"].font = Font(bold=True, size=14)
    ws.append([])
    ws.append(["Transaction ID", facts.get("transaction_id", "")])
    ws.append(["Standard", facts.get("standard", "")])
    ws.append(["Title", facts.get("title", "")])
    ws.append(["Issues Count", len(facts.get("issues") or [])])
    ws.append(["Generated", datetime.now().isoformat(timespec="seconds")])

    ws_at = wb.create_sheet("AuditTrail")
    ws_at.append(["Key", "Value"])
    for cell in ws_at[1]:
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
    ws_at.append(["Memo Standard", facts.get("standard", "")])
    _autosize = lambda w: None  # noqa: E731
    output.parent.mkdir(parents=True, exist_ok=True)
    wb.save(output)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    parser.add_argument("--facts", required=True, type=Path)
    parser.add_argument("--output-docx", required=True, type=Path)
    parser.add_argument("--output-xlsx", type=Path, default=None)
    parser.add_argument("--quiet", action="store_true")
    args = parser.parse_args(argv)

    logging.basicConfig(
        level=logging.WARNING if args.quiet else logging.INFO,
        format="%(levelname)s %(name)s: %(message)s",
    )

    try:
        facts = _load_facts(args.facts)
    except (OSError, yaml.YAMLError, ValueError) as exc:
        LOG.error("%s", exc)
        return 1

    draft_memo(facts, args.output_docx)
    if args.output_xlsx:
        write_summary_xlsx(facts, args.output_xlsx)

    print(f"Transaction:     {facts.get('transaction_id', 'N/A')}")
    print(f"Standard:        {facts.get('standard', 'N/A')}")
    print(f"Memo:            {args.output_docx}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
