"""Debt vs. Equity Classification Reviewer — ASC 480 / 815 / 470 waterfall.

Implements the workflow in ../SKILL.md.
"""
from __future__ import annotations

import argparse
import logging
import sys
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any

import yaml
from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

LOG = logging.getLogger("debt_equity_reviewer")

ACCOUNTING_FMT = '_($* #,##0.00_);_($* (#,##0.00);_($* "-"??_);_(@_)'
HEADER_FILL = PatternFill(start_color="305496", end_color="305496", fill_type="solid")
HEADER_FONT = Font(bold=True, color="FFFFFF")
THIN = Side(border_style="thin", color="BFBFBF")
BOX = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)


@dataclass
class StepResult:
    step: int
    gate: str
    triggered: bool
    citation: str
    rationale: str
    classification_if_stop: str | None = None


@dataclass
class ClassificationResult:
    instrument_id: str
    description: str
    classification: str
    steps: list[StepResult] = field(default_factory=list)
    embedded_features: list[dict[str, Any]] = field(default_factory=list)
    proceeds: float = 0.0
    mezzanine_accretion: list[dict[str, Any]] = field(default_factory=list)


def _load_yaml(path: Path) -> dict[str, Any]:
    with path.open(encoding="utf-8") as fh:
        data = yaml.safe_load(fh)
    if not isinstance(data, dict):
        raise ValueError("Instrument file must be a YAML mapping")
    return data


def _style_header(ws, row: int = 1) -> None:
    for cell in ws[row]:
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = BOX


def _autosize(ws) -> None:
    for col in ws.columns:
        letter = get_column_letter(col[0].column)
        max_len = max((len(str(c.value)) for c in col if c.value is not None), default=10)
        ws.column_dimensions[letter].width = min(45, max_len + 2)


def run_waterfall(inst: dict[str, Any], sec_registrant: bool) -> ClassificationResult:
    instrument_id = str(inst.get("instrument_id", "UNKNOWN"))
    description = str(inst.get("description", ""))
    proceeds = float(inst.get("proceeds", 0))
    steps: list[StepResult] = []

    def add(step: int, gate: str, triggered: bool, citation: str, rationale: str,
            stop_class: str | None = None) -> None:
        steps.append(StepResult(step, gate, triggered, citation, rationale, stop_class))

    # Step 1 — mandatorily redeemable
    mandatory = bool(inst.get("mandatorily_redeemable", False))
    add(1, "Mandatorily redeemable share?", mandatory,
        "ASC 480-10-25-4",
        "Unconditional obligation to redeem at specified or determinable date.",
        "Liability" if mandatory else None)
    if mandatory:
        return ClassificationResult(instrument_id, description, "Liability", steps, proceeds=proceeds)

    # Step 2 — variable shares for fixed monetary amount
    variable = bool(inst.get("variable_share_count_for_fixed_monetary_amount", False))
    add(2, "Obligation to issue variable shares for fixed monetary amount?", variable,
        "ASC 480-10-25-14",
        "Settlement amount varies with share price while monetary obligation is fixed.",
        "Liability" if variable else None)
    if variable:
        return ClassificationResult(instrument_id, description, "Liability", steps, proceeds=proceeds)

    # Step 3 — written put / forward purchase
    written_put = bool(inst.get("written_put_or_forward_purchase_on_own_shares", False))
    add(3, "Written put or forward purchase on own shares?", written_put,
        "ASC 480-10-25-8",
        "Entity may be obligated to repurchase shares for cash or other assets.",
        "Liability" if written_put else None)
    if written_put:
        return ClassificationResult(instrument_id, description, "Liability", steps, proceeds=proceeds)

    # Step 4 — embedded derivative bifurcation screen
    embedded = inst.get("embedded_features") or []
    needs_bifurcation = any(
        not bool(f.get("clearly_and_closely_related", True)) for f in embedded
    )
    add(4, "Embedded feature not clearly and closely related?", needs_bifurcation,
        "ASC 815-15",
        "Hybrid instrument may require separate derivative accounting if not C&CR.",
        None)

    # Step 5 — ASC 815-40 indexed + equity settled
    fixed_for_fixed = bool(inst.get("fixed_for_fixed", True))
    settlement = str(inst.get("settlement_form", "shares")).lower()
    cash_settlement = settlement in ("cash", "choice_of_either")
    indexed = fixed_for_fixed and not cash_settlement
    add(5, "Indexed to own equity AND equity-settled (815-40)?", indexed,
        "ASC 815-40-25",
        "Fixed-for-fixed pricing and no cash settlement outside issuer control.",
        "Liability (mark-to-market)" if not indexed else None)
    if not indexed:
        return ClassificationResult(
            instrument_id, description, "Liability (Derivative / Hybrid)",
            steps, embedded_features=list(embedded), proceeds=proceeds,
        )

    # Step 6 — SEC mezzanine (480-10-S99)
    holder_put = bool(inst.get("redeemable_at_holder_option", False))
    outside_control = bool(inst.get("holder_redemption_outside_issuer_control", False))
    contingency = bool(inst.get("redeemable_on_contingency", False))
    mezzanine = sec_registrant and (holder_put or outside_control or contingency)
    add(6, "Redemption outside issuer control (SEC registrant)?", mezzanine,
        "ASC 480-10-S99 / SEC ASR 268",
        "Holder put or contingent redemption not solely within issuer control.",
        "Mezzanine (Temporary Equity)" if mezzanine else None)

    accretion: list[dict[str, Any]] = []
    if mezzanine:
        redemption_price = float(inst.get("redemption_price", 0))
        shares = float(inst.get("shares_issued", 0))
        redemption_value = redemption_price * shares if redemption_price and shares else proceeds
        issue_date = str(inst.get("issue_date", ""))
        redemption_date = str(inst.get("redemption_date", ""))
        accretion.append({
            "carrying_value_at_issuance": proceeds,
            "redemption_value": redemption_value,
            "accretion_to_redemption": round(redemption_value - proceeds, 2),
            "issue_date": issue_date,
            "redemption_date": redemption_date,
        })
        return ClassificationResult(
            instrument_id, description, "Mezzanine (Temporary Equity)",
            steps, embedded_features=list(embedded), proceeds=proceeds,
            mezzanine_accretion=accretion,
        )

    # Step 7 — permanent equity
    add(7, "Permanent equity classification", True,
        "ASC 480 / Subtopic 505",
        "No liability triggers; not redeemable outside issuer control for SEC registrants.",
        "Permanent Equity")
    return ClassificationResult(
        instrument_id, description, "Permanent Equity",
        steps, embedded_features=list(embedded), proceeds=proceeds,
    )


def build_je_shell(result: ClassificationResult) -> list[dict[str, Any]]:
    p = result.proceeds
    cls = result.classification
    if "Liability" in cls:
        return [
            {"account": "1000 Cash", "debit": p, "credit": 0, "memo": "Proceeds — liability instrument"},
            {"account": "2500 Debt / Liability", "debit": 0, "credit": p, "memo": "Initial recognition"},
        ]
    if "Mezzanine" in cls:
        return [
            {"account": "1000 Cash", "debit": p, "credit": 0, "memo": "Proceeds — mezzanine preferred"},
            {"account": "3100 Temporary Equity (Mezzanine)", "debit": 0, "credit": p, "memo": "Initial recognition"},
        ]
    return [
        {"account": "1000 Cash", "debit": p, "credit": 0, "memo": "Proceeds — equity instrument"},
        {"account": "3000 Additional Paid-In Capital", "debit": 0, "credit": p, "memo": "Initial recognition"},
    ]


def write_workpaper(result: ClassificationResult, inst: dict[str, Any],
                    sec_registrant: bool, output: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Summary"
    ws.append(["Debt vs. Equity Classification Workpaper"])
    ws["A1"].font = Font(bold=True, size=14)
    ws.append([])
    ws.append(["Instrument ID", result.instrument_id])
    ws.append(["Description", result.description])
    ws.append(["Issuer", inst.get("issuer", "")])
    ws.append(["Proceeds", result.proceeds])
    ws.append(["SEC Registrant", "Yes" if sec_registrant else "No"])
    ws.append(["Classification", result.classification])
    ws["B5"].number_format = ACCOUNTING_FMT
    _autosize(ws)

    ws_dt = wb.create_sheet("DecisionTree")
    ws_dt.append(["Step", "Gate", "Triggered?", "Citation", "Rationale", "Stop Classification"])
    _style_header(ws_dt)
    for s in result.steps:
        ws_dt.append([
            s.step, s.gate, "Y" if s.triggered else "N",
            s.citation, s.rationale, s.classification_if_stop or "",
        ])
    _autosize(ws_dt)

    ws_ef = wb.create_sheet("EmbeddedFeatures")
    ws_ef.append(["Type", "Strike", "Clearly & Closely Related?", "Bifurcation Note"])
    _style_header(ws_ef)
    if result.embedded_features:
        for f in result.embedded_features:
            ws_ef.append([
                f.get("type", ""), f.get("strike", ""),
                "Y" if f.get("clearly_and_closely_related", True) else "N",
                "Bifurcate if not C&CR" if not f.get("clearly_and_closely_related", True) else "No bifurcation",
            ])
    else:
        ws_ef.append(["(none)"])
    _autosize(ws_ef)

    ws_je = wb.create_sheet("JEShell")
    ws_je.append(["Account", "Debit", "Credit", "Memo"])
    _style_header(ws_je)
    for line in build_je_shell(result):
        ws_je.append([line["account"], line["debit"], line["credit"], line["memo"]])
    for letter in ("B", "C"):
        for cell in ws_je[letter][1:]:
            cell.number_format = ACCOUNTING_FMT
    _autosize(ws_je)

    ws_m = wb.create_sheet("Mezzanine_Accretion")
    ws_m.append(["Carrying Value", "Redemption Value", "Accretion Required", "Issue Date", "Redemption Date"])
    _style_header(ws_m)
    if result.mezzanine_accretion:
        for row in result.mezzanine_accretion:
            ws_m.append([
                row["carrying_value_at_issuance"], row["redemption_value"],
                row["accretion_to_redemption"], row["issue_date"], row["redemption_date"],
            ])
        for letter in ("A", "B", "C"):
            for cell in ws_m[letter][1:]:
                cell.number_format = ACCOUNTING_FMT
    else:
        ws_m.append(["(not applicable)"])
    _autosize(ws_m)

    ws_ref = wb.create_sheet("References")
    ws_ref.append(["Standard", "Topic"])
    _style_header(ws_ref)
    for ref in (
        ("ASC 480", "Distinguishing Liabilities from Equity"),
        ("ASC 815-40", "Contracts in Entity's Own Equity"),
        ("ASC 815-15", "Embedded Derivatives"),
        ("ASC 480-10-S99", "SEC Mezzanine Guidance"),
        ("ASC 470-20", "Debt with Conversion Options (ASU 2020-06)"),
    ):
        ws_ref.append(ref)
    _autosize(ws_ref)

    ws_at = wb.create_sheet("AuditTrail")
    ws_at.append(["Key", "Value"])
    _style_header(ws_at)
    ws_at.append(["Generated", datetime.now().isoformat(timespec="seconds")])
    ws_at.append(["Classification", result.classification])
    _autosize(ws_at)

    ws_so = wb.create_sheet("SignOff")
    ws_so.append(["Role", "Name", "Title", "Date", "Notes"])
    _style_header(ws_so)
    for role in ("Preparer (Senior Accountant)", "Reviewer (Technical Accounting)", "Approver (Controller)"):
        ws_so.append([role, "", "", "", ""])
    _autosize(ws_so)

    output.parent.mkdir(parents=True, exist_ok=True)
    wb.save(output)
    LOG.info("Workpaper written: %s", output)


def write_docx_memo(result: ClassificationResult, inst: dict[str, Any], output: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    doc.add_heading("Technical Accounting Memo — Debt vs. Equity Classification", level=1)
    doc.add_paragraph(f"Instrument: {result.instrument_id}")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")

    doc.add_heading("Issue", level=2)
    doc.add_paragraph(
        f"How should {result.description} issued by {inst.get('issuer', 'the issuer')} "
        f"be classified under U.S. GAAP?"
    )

    doc.add_heading("Rule", level=2)
    doc.add_paragraph(
        "Classification follows ASC 480 (mandatorily redeemable instruments and certain "
        "settlement obligations), ASC 815-40 (indexed to and settled in own equity), and "
        "ASC 480-10-S99 for SEC registrants with redemption outside issuer control."
    )

    doc.add_heading("Analysis", level=2)
    for s in result.steps:
        trigger = "Yes" if s.triggered else "No"
        doc.add_paragraph(
            f"Step {s.step} — {s.gate} [{trigger}]. {s.rationale} ({s.citation})."
        )
        if s.classification_if_stop:
            doc.add_paragraph(f"Conclusion at this step: {s.classification_if_stop}.")

    doc.add_heading("Conclusion", level=2)
    doc.add_paragraph(
        f"Based on the instrument terms, the appropriate classification is "
        f"{result.classification}."
    )
    doc.add_paragraph("Draft initial recognition journal entry:")
    for line in build_je_shell(result):
        side = "Debit" if line["debit"] else "Credit"
        amount = line["debit"] or line["credit"]
        doc.add_paragraph(f"{side} {line['account']} — ${amount:,.2f}. {line['memo']}")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    LOG.info("Memo written: %s", output)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    parser.add_argument("--instrument", required=True, type=Path)
    parser.add_argument("--output-xlsx", required=True, type=Path)
    parser.add_argument("--output-docx", type=Path, default=None)
    parser.add_argument("--issuer-is-sec-registrant", action="store_true")
    parser.add_argument("--quiet", action="store_true")
    args = parser.parse_args(argv)

    logging.basicConfig(
        level=logging.WARNING if args.quiet else logging.INFO,
        format="%(levelname)s %(name)s: %(message)s",
    )

    try:
        inst = _load_yaml(args.instrument)
    except (OSError, yaml.YAMLError, ValueError) as exc:
        LOG.error("Failed to load instrument: %s", exc)
        return 1

    result = run_waterfall(inst, args.issuer_is_sec_registrant)
    write_workpaper(result, inst, args.issuer_is_sec_registrant, args.output_xlsx)
    if args.output_docx:
        write_docx_memo(result, inst, args.output_docx)

    print(f"Instrument:      {result.instrument_id}")
    print(f"Classification:  {result.classification}")
    print(f"Workpaper:       {args.output_xlsx}")
    if args.output_docx:
        print(f"Memo:            {args.output_docx}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
